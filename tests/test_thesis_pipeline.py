"""论文模式：严格清洗、语义角色、分页约束和前缀样式回归测试。"""

import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor

from core.apply import apply_format
from core.label_roles import detect_document_profile, regex_role
from core.schema import SpecValidationError, validate_spec
from core.text_integrity import check_text_integrity


def _rule(**overrides):
    value = {
        "font_eastasia": "宋体", "font_ascii": "Times New Roman",
        "size_pt": 12, "bold": False, "italic": False,
        "underline": False, "color": "000000", "alignment": "justify",
    }
    value.update(overrides)
    return value


def _direct_run_has(run, tag):
    rpr = run._element.find(qn("w:rPr"))
    return rpr is not None and rpr.find(qn(tag)) is not None


def test_thesis_strict_cleanup_and_semantic_styles(tmp_path):
    source = tmp_path / "thesis_messy.docx"
    output = tmp_path / "thesis_clean.docx"
    doc = Document()
    chapter = doc.add_paragraph("一、绪论")
    chapter.runs[0].italic = True
    chapter.runs[0].underline = True
    chapter.runs[0].font.color.rgb = RGBColor(192, 0, 0)
    body = doc.add_paragraph("正文不应继承绿色斜体。")
    body.runs[0].italic = True
    body.runs[0].font.color.rgb = RGBColor(0, 112, 0)
    doc.add_paragraph("摘要：本文研究文档格式迁移。")
    doc.add_paragraph("关键词：文档格式；样式迁移")
    doc.add_paragraph("参考文献")
    doc.add_paragraph("[1] Author. A deliberately long bibliography entry for hanging indent.")
    doc.save(source)

    spec = {
        "profile": "thesis",
        "cleanup": {"mode": "strict"},
        "roles": {
            "body": _rule(),
            "chapter_heading": _rule(
                bold=True, alignment="center", keep_with_next=True,
                keep_together=True, page_break_before=True, outline_level=0),
            "abstract_body": _rule(
                first_line_indent_chars=0,
                label_prefix={"text": ["摘要：", "摘要:"], "bold": True}),
            "keywords": _rule(
                first_line_indent_chars=0,
                label_prefix={"text": ["关键词：", "关键词:"], "bold": True}),
            "bibliography_heading": _rule(
                bold=True, alignment="center", keep_with_next=True,
                page_break_before=True, outline_level=0),
            "bibliography_entry": _rule(
                alignment="left", left_indent_chars=2,
                hanging_indent_chars=2, widow_control=True),
        },
    }
    validate_spec(spec)
    rolemap = {
        0: "chapter_heading", 1: "body", 2: "abstract_body",
        3: "keywords", 4: "bibliography_heading", 5: "bibliography_entry",
    }
    apply_format(source, spec, rolemap, output)
    result = Document(output)

    assert [p.text for p in result.paragraphs] == [
        "一、绪论", "正文不应继承绿色斜体。", "摘要：本文研究文档格式迁移。",
        "关键词：文档格式；样式迁移", "参考文献",
        "[1] Author. A deliberately long bibliography entry for hanging indent.",
    ]
    for paragraph in result.paragraphs[:2]:
        assert not _direct_run_has(paragraph.runs[0], "w:i")
        assert not _direct_run_has(paragraph.runs[0], "w:u")
        assert not _direct_run_has(paragraph.runs[0], "w:color")

    chapter_pf = result.paragraphs[0].style.paragraph_format
    assert chapter_pf.keep_with_next is True
    assert chapter_pf.keep_together is True
    assert chapter_pf.page_break_before is True

    for index, prefix in ((2, "摘要："), (3, "关键词：")):
        paragraph = result.paragraphs[index]
        assert len(paragraph.runs) >= 2
        assert paragraph.runs[0].text == prefix
        assert paragraph.runs[0].bold is True
        assert paragraph.runs[1].bold is None

    ref_ind = result.paragraphs[5].style.element.find(qn("w:pPr")).find(qn("w:ind"))
    assert ref_ind.get(qn("w:leftChars")) == "200"
    assert ref_ind.get(qn("w:hangingChars")) == "200"


def test_preserve_emphasis_removes_noise_but_keeps_bold_italic(tmp_path):
    source = tmp_path / "emphasis.docx"
    output = tmp_path / "emphasis_out.docx"
    doc = Document()
    run = doc.add_paragraph("需要保留的强调").runs[0]
    run.bold = True
    run.italic = True
    run.font.color.rgb = RGBColor(255, 0, 0)
    doc.save(source)
    spec = {
        "cleanup": {"mode": "preserve_emphasis"},
        "roles": {"body": _rule()},
    }
    apply_format(source, spec, {0: "body"}, output)
    result_run = Document(output).paragraphs[0].runs[0]
    assert _direct_run_has(result_run, "w:b")
    assert _direct_run_has(result_run, "w:i")
    assert not _direct_run_has(result_run, "w:color")


def test_thesis_role_detection_and_schema_errors():
    records = [
        {"text": "摘要：内容", "in_table": False},
        {"text": "关键词：格式", "in_table": False},
        {"text": "参考文献", "in_table": False},
    ]
    assert detect_document_profile(records) == "thesis"
    assert regex_role("摘要：内容", profile="thesis") == "abstract_body"
    assert regex_role("关键词：格式", profile="thesis") == "keywords"
    assert regex_role("参考文献", profile="thesis") == "bibliography_heading"
    assert regex_role("[12] Author. Paper.", profile="thesis") == "bibliography_entry"
    assert regex_role("一、绪论", profile="thesis") == "chapter_heading"

    bad = {"cleanup": {"mode": "erase_everything"}, "roles": {"body": _rule()}}
    try:
        validate_spec(bad)
    except SpecValidationError as exc:
        assert "cleanup.mode" in str(exc)
    else:
        raise AssertionError("非法 cleanup.mode 应被 schema 拒绝")


def test_highlight_is_serialized_before_underline(tmp_path):
    """CT_RPr 对节点顺序敏感；顺序错误会触发 Word“发现不可读内容”。"""
    source = tmp_path / "highlight.docx"
    output = tmp_path / "highlight_out.docx"
    doc = Document()
    doc.add_paragraph("高亮顺序回归")
    doc.save(source)
    spec = {
        "roles": {
            "body": _rule(underline=True, highlight="yellow"),
        },
    }
    apply_format(source, spec, {0: "body"}, output)
    style = Document(output).paragraphs[0].style
    tags = [child.tag for child in style.element.find(qn("w:rPr"))]
    assert tags.index(qn("w:highlight")) < tags.index(qn("w:u"))


def test_thesis_structure_builds_cover_toc_sections_and_fields(tmp_path):
    source = tmp_path / "thesis_source.docx"
    output = tmp_path / "thesis_structured.docx"
    doc = Document()
    doc.add_paragraph("文档格式迁移研究")
    doc.add_paragraph("摘要：本文研究文档格式迁移。")
    doc.add_paragraph("关键词：文档格式；样式迁移")
    doc.add_paragraph("一、绪论")
    doc.add_paragraph("这是正文。")
    doc.add_paragraph("参考文献")
    doc.add_paragraph("[1] Author. Paper.")
    doc.save(source)

    spec = {
        "profile": "thesis",
        "cleanup": {"mode": "strict"},
        "page": {
            "margin": {
                "top_mm": 35, "bottom_mm": 40,
                "left_mm": 25, "right_mm": 25,
            },
        },
        "toc": {"enabled": True, "levels": [1, 2]},
        "roles": {
            "body": _rule(),
            "title": _rule(size_pt=22, bold=True, alignment="center"),
            "abstract_heading": _rule(
                size_pt=16, bold=True, alignment="center", outline_level=0),
            "abstract_body": _rule(first_line_indent_chars=0),
            "keywords": _rule(first_line_indent_chars=0),
            "chapter_heading": _rule(
                size_pt=16, bold=True, alignment="center", outline_level=0,
                keep_with_next=True, page_break_before=True),
            "bibliography_heading": _rule(
                size_pt=16, bold=True, alignment="center", outline_level=0,
                keep_with_next=True, page_break_before=True),
            "bibliography_entry": _rule(
                alignment="left", left_indent_chars=2,
                hanging_indent_chars=2),
        },
        "structure": {
            "enabled": True,
            "mode": "thesis",
            "cover": {
                "enabled": True,
                "logo": False,
                "institution": "测试大学硕士学位论文",
                "type_label": "测试大学硕士学位论文",
                "metadata": {"姓名": "（待填写）", "学号": "（待填写）"},
                "date_text": "日期：____年__月__日",
            },
            "front_matter": {
                "abstract": True, "toc": True, "declarations": False,
            },
            "page_numbering": {
                "front_format": "upperRoman", "front_start": 1,
                "body_format": "decimal", "body_start": 1,
            },
            "running_header": {
                "left_text": "测试大学硕士学位论文",
                "header_distance_mm": 15,
                "footer_distance_mm": 17.5,
            },
        },
    }
    validate_spec(spec)
    rolemap = {
        0: "title", 1: "abstract_body", 2: "keywords",
        3: "chapter_heading", 4: "body",
        5: "bibliography_heading", 6: "bibliography_entry",
    }
    changelog = apply_format(source, spec, rolemap, output)
    result = Document(output)

    assert len(result.sections) == 5
    texts = [p.text for p in result.paragraphs]
    assert texts.count("文档格式迁移研究") == 1
    assert "摘　要" in texts
    assert "目　录" in texts
    assert "摘要：本文研究文档格式迁移。" not in texts
    assert "本文研究文档格式迁移。" in texts

    body_fields = " ".join(
        (element.text or "")
        for element in result.element.body.iter(qn("w:instrText")))
    assert 'TOC \\o "1-2"' in body_fields
    all_story_fields = []
    for section in result.sections:
        for story in (section.header, section.footer):
            all_story_fields.extend(
                (element.text or "")
                for element in story._element.iter(qn("w:instrText")))
    assert any("STYLEREF" in field for field in all_story_fields)
    assert sum(field.strip() == "PAGE" for field in all_story_fields) >= 4

    formats = []
    starts = []
    for section in result.sections:
        page_type = section._sectPr.find(qn("w:pgNumType"))
        formats.append(page_type.get(qn("w:fmt")) if page_type is not None else None)
        starts.append(page_type.get(qn("w:start")) if page_type is not None else None)
    assert formats == [None, "upperRoman", "upperRoman", "decimal", "decimal"]
    assert starts == [None, "1", None, "1", None]
    assert result.settings.element.find(qn("w:updateFields")).get(qn("w:val")) == "true"

    structure_change = next(c for c in changelog if c["role"] == "structure")
    integrity = check_text_integrity(
        source, output,
        allowed_additions=structure_change["allowed_additions"],
        expected_stripped_prefixes=structure_change["stripped_prefixes"],
        allowed_story_changes={"headers", "footers"},
    )
    assert integrity["ok"], integrity
