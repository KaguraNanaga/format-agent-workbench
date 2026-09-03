# -*- coding: utf-8 -*-
"""公文、技术手册、法律 brief 与旧格式输入适配回归测试。"""

import os
import json
import subprocess
import sys

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from core.apply import apply_format
from core.agent import Agent
from core.field_refresh import _update_safe_fields
from core.input_conversion import InputConversionError, converted_input
from core.label_roles import detect_document_profile, regex_role
from core.preflight import scan_docx
from core.schema import SpecValidationError
from core.style_packs import get_style_pack
from core.text_integrity import check_text_integrity


def _append_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)
    run = paragraph.add_run()
    code = OxmlElement("w:instrText")
    code.set(qn("xml:space"), "preserve")
    code.text = f" {instruction} "
    run._r.append(code)
    run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


def test_official_profile_roles_and_outside_page_number(tmp_path):
    records = [
        {"text": "某某局通知"}, {"text": "某局〔2026〕3号"},
        {"text": "各处室："}, {"text": "特此通知。"},
    ]
    assert detect_document_profile(records) == "official_cn"
    assert regex_role("某局〔2026〕3号", profile="official_cn") == "document_number"
    assert regex_role("各处室：", profile="official_cn") == "recipient"
    assert regex_role("特此通知。", profile="official_cn") == "closing"
    assert regex_role("（1）具体事项", profile="official_cn") == "heading_4"

    source = tmp_path / "official.docx"
    output = tmp_path / "official_out.docx"
    doc = Document()
    for item in records:
        doc.add_paragraph(item["text"])
    doc.save(source)
    spec = get_style_pack("official-cn-gbt9704")
    apply_format(source, spec, {
        0: "title", 1: "document_number", 2: "recipient", 3: "closing",
    }, output)
    result = Document(output)
    section = result.sections[0]
    assert round(section.top_margin.mm) == 37
    assert round(section.footer_distance.mm) == 7
    assert result.settings.odd_and_even_pages_header_footer is True
    for footer in (section.footer, section.even_page_footer):
        assert any(
            (node.text or "").strip() == "PAGE"
            for node in footer._element.iter(qn("w:instrText"))
        )
        assert "—" in footer.paragraphs[0].text
    assert check_text_integrity(
        source, output, allowed_story_changes={"footers"})["ok"]


def test_technical_callouts_and_figure_caption_binding(tmp_path):
    records = [
        {"text": "Installation"}, {"text": "WARNING: Disconnect power."},
        {"text": "Step 1 Install the package"}, {"text": "GET /v1/users"},
    ]
    assert detect_document_profile(records) == "english_technical"
    assert regex_role(records[1]["text"], profile="english_technical") == "warning_box"
    assert regex_role(records[2]["text"], profile="english_technical") == "procedure_step"
    assert regex_role(records[3]["text"], profile="english_technical") == "command"

    source = tmp_path / "manual.docx"
    output = tmp_path / "manual_out.docx"
    doc = Document()
    doc.add_paragraph("WARNING: Disconnect power before servicing.")
    drawing_paragraph = doc.add_paragraph()
    drawing_run = drawing_paragraph.add_run()
    drawing_run._r.append(OxmlElement("w:drawing"))
    doc.add_paragraph("Figure 1. Power connector")
    doc.save(source)
    changelog = apply_format(
        source, get_style_pack("technical-manual"),
        {0: "warning_box", 1: "body", 2: "figure_caption"}, output)
    result = Document(output)
    warning_style = result.styles["警告框"]
    ppr = warning_style.element.get_or_add_pPr()
    assert ppr.find(qn("w:shd")).get(qn("w:fill")) == "FEE2E2"
    assert ppr.find(qn("w:pBdr")).find(qn("w:left")) is not None
    assert result.paragraphs[1]._p.pPr.find(qn("w:keepNext")) is not None
    technical = next(item for item in changelog if item["role"] == "technical_features")
    assert "figure_caption_bindings_1" in technical["changed_fields"]
    assert check_text_integrity(source, output)["ok"]


def test_legal_brief_ta_toa_fields_are_explicit_and_idempotent(tmp_path):
    records = [
        {"text": "IN THE SUPREME COURT OF EXAMPLE"},
        {"text": "No. 24-123"}, {"text": "TABLE OF AUTHORITIES"},
        {"text": "BRIEF OF APPELLANT"},
    ]
    assert detect_document_profile(records) == "english_legal_brief"
    assert regex_role(records[0]["text"], profile="english_legal_brief") == "court_caption"
    assert regex_role(records[1]["text"], profile="english_legal_brief") == "case_number"
    assert regex_role(records[2]["text"], profile="english_legal_brief") == "table_of_authorities_heading"
    assert regex_role(records[3]["text"], profile="english_legal_brief") == "brief_title"

    source = tmp_path / "brief.docx"
    output = tmp_path / "brief_out.docx"
    second = tmp_path / "brief_out_2.docx"
    doc = Document()
    doc.add_paragraph("TABLE OF AUTHORITIES")
    doc.add_paragraph("The Court followed Example v. Sample in resolving the issue.")
    doc.save(source)
    options = {
        "insert_toa": True,
        "citation_marks": [{
            "text": "Example v. Sample",
            "long": "Example v. Sample, 123 F.3d 456 (2024)",
            "short": "Example", "category": 1,
        }],
    }
    spec = get_style_pack("us-legal-brief", **options)
    rolemap = {0: "table_of_authorities_heading", 1: "body"}
    apply_format(source, spec, rolemap, output)
    scan = scan_docx(output)
    assert scan["field_inventory"]["TA"] == 1
    assert scan["field_inventory"]["TOA"] == 1
    assert check_text_integrity(source, output)["ok"]
    result = Document(output)
    ta_instr = next(
        node for node in result.element.iter(qn("w:instrText"))
        if (node.text or "").strip().startswith("TA ")
    )
    rpr = ta_instr.getparent().find(qn("w:rPr"))
    assert rpr is not None and rpr.find(qn("w:vanish")) is not None

    apply_format(output, spec, rolemap, second)
    repeated = scan_docx(second)
    assert repeated["field_inventory"]["TA"] == 1
    assert repeated["field_inventory"]["TOA"] == 1

    with pytest.raises(SpecValidationError):
        get_style_pack("us-legal-brief", citation_marks=[
            {"text": "Duplicate", "long": "First", "category": 1},
            {"text": "Duplicate", "long": "Second", "category": 1},
        ])


def test_input_conversion_passthrough_pdf_rejection_and_odt_adapter(tmp_path, monkeypatch):
    native = tmp_path / "native.docx"
    doc = Document()
    doc.add_paragraph("Native DOCX")
    doc.save(native)
    with converted_input(native) as converted:
        assert converted.converter == "native-docx"
        assert converted.lossy is False
        assert converted.docx_path == str(native.resolve())

    pdf = tmp_path / "scan.pdf"
    pdf.write_bytes(b"%PDF-1.4")
    with pytest.raises(InputConversionError, match="PDF"):
        with converted_input(pdf):
            pass

    odt = tmp_path / "legacy.odt"
    odt.write_bytes(b"placeholder")
    import core.input_conversion as conversion

    def fake_libreoffice(source, destination_dir):
        generated = destination_dir / f"{source.stem}.docx"
        converted_doc = Document()
        converted_doc.add_paragraph("Converted ODT")
        converted_doc.save(generated)
        return generated

    monkeypatch.setattr(conversion, "_convert_with_libreoffice", fake_libreoffice)
    with converted_input(odt) as converted:
        temporary = converted.docx_path
        assert converted.converter == "LibreOffice"
        assert converted.lossy is True
        assert Document(temporary).paragraphs[0].text == "Converted ODT"
    assert not os.path.exists(temporary)


def test_preflight_ignores_disabled_wps_document_protection_marker(tmp_path):
    disabled = tmp_path / "wps_disabled_protection.docx"
    enabled = tmp_path / "enabled_protection.docx"

    doc = Document()
    doc.add_paragraph("WPS placeholder")
    protection = OxmlElement("w:documentProtection")
    protection.set(qn("w:enforcement"), "0")
    doc.settings._element.append(protection)
    doc.save(disabled)
    assert scan_docx(disabled)["has_document_protection"] is False

    doc = Document()
    doc.add_paragraph("Actually protected")
    protection = OxmlElement("w:documentProtection")
    protection.set(qn("w:enforcement"), "1")
    doc.settings._element.append(protection)
    doc.save(enabled)
    assert scan_docx(enabled)["has_document_protection"] is True


def test_field_refresh_worker_timeout_is_bounded(tmp_path, monkeypatch):
    import core.field_refresh as refresh
    import core.input_conversion as conversion

    source = tmp_path / "refresh.docx"
    Document().save(source)
    monkeypatch.setattr(refresh.sys, "platform", "win32")
    monkeypatch.setattr(conversion, "_office_pids", lambda _names: set())
    monkeypatch.setattr(
        conversion, "_terminate_new_office_processes",
        lambda _names, _before: [1234],
    )
    monkeypatch.setattr(conversion, "_com_timeout_seconds", lambda: 5.0)

    def timeout(*_args, **_kwargs):
        raise subprocess.TimeoutExpired(["python", "worker"], 5)

    monkeypatch.setattr(refresh.subprocess, "run", timeout)
    with pytest.raises(RuntimeError, match="超时.*1234"):
        refresh.refresh_fields_word(source)


def test_preflight_reports_external_fields_and_refresh_uses_allowlist(tmp_path):
    source = tmp_path / "external_fields.docx"
    doc = Document()
    _append_field(doc.add_paragraph(), 'INCLUDETEXT "https://example.invalid/data.docx"')
    _append_field(doc.add_paragraph(), "PAGE")
    doc.part.relate_to("https://example.invalid/", RT.HYPERLINK, is_external=True)
    doc.save(source)

    scan = scan_docx(source)
    assert scan["unsafe_field_inventory"] == {"INCLUDETEXT": 1}
    assert scan["external_relationship_count"] == 1
    from core.preflight import preflight_docx
    codes = {item["code"] for item in preflight_docx(source)["warnings"]}
    assert {"UNSAFE_EXTERNAL_FIELDS", "EXTERNAL_RELATIONSHIPS"} <= codes

    class FakeField:
        def __init__(self, code):
            self.Code = type("Code", (), {"Text": code})()
            self.updated = 0

        def Update(self):
            self.updated += 1

    class FakeFields:
        def __init__(self, fields):
            self.fields = fields
            self.Count = len(fields)

        def __call__(self, index):
            return self.fields[index - 1]

    page = FakeField(" PAGE ")
    external = FakeField(' INCLUDETEXT "https://example.invalid/data.docx" ')
    result = _update_safe_fields(FakeFields([page, external]))
    assert page.updated == 1
    assert external.updated == 0
    assert result == {
        "total": 2, "updated": 1, "skipped": {"INCLUDETEXT": 1},
    }


def test_cli_exposes_explicit_legal_ta_toa_options(tmp_path, monkeypatch):
    source = tmp_path / "cli_brief.docx"
    output = tmp_path / "cli_brief_out.docx"
    citations = tmp_path / "citations.json"
    rolemap = tmp_path / "rolemap.json"
    doc = Document()
    doc.add_paragraph("TABLE OF AUTHORITIES")
    doc.add_paragraph("The Court followed Example v. Sample.")
    doc.save(source)
    citations.write_text(json.dumps([{
        "text": "Example v. Sample", "long": "Example v. Sample, 123 F.3d 456",
        "category": 1,
    }]), encoding="utf-8")
    rolemap.write_text(json.dumps({"0": "table_of_authorities_heading", "1": "body"}),
                       encoding="utf-8")

    import main as cli
    monkeypatch.setattr(sys, "argv", [
        "main.py", "--style-pack", "us-legal-brief",
        "--legal-citations-json", str(citations), "--insert-toa",
        "--rolemap-json", str(rolemap), "--target", str(source),
        "--out", str(output),
    ])
    assert cli.main() == 0
    inventory = scan_docx(output)["field_inventory"]
    assert inventory["TA"] == 1
    assert inventory["TOA"] == 1
    preflight = json.loads(
        (tmp_path / "cli_brief_out_preflight.json").read_text(encoding="utf-8"))
    assert preflight["input_conversion"]["target"]["converter"] == "native-docx"


def test_external_rolemap_rejects_stale_indices(tmp_path):
    source = tmp_path / "rolemap_source.docx"
    output = tmp_path / "rolemap_output.docx"
    doc = Document()
    doc.add_paragraph("Body")
    doc.save(source)
    with pytest.raises(ValueError, match="不存在的段落 idx"):
        Agent().run(
            source, output, style_pack="mla9",
            rolemap={0: "body", 999: "body"},
        )
    assert not output.exists()
