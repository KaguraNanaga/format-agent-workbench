# -*- coding: utf-8 -*-
"""命名样式流水线测试：模板取样 → FormatSpec → 目标文档样式创建与绑定。"""

import json
import os
import sys
import tempfile

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from core.apply import apply_format
from core.rules_from_template import extract_rules_from_template
from core.schema import validate_spec
from core.style_set import resolve_target_body_style, style_id_for_role

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))


def _direct_value(paragraph, tag):
    ppr = paragraph._p.pPr
    if ppr is None:
        return None
    el = ppr.find(qn(tag))
    return el.get(qn("w:val")) if el is not None else None


def _style_outline(style):
    ppr = style.element.find(qn("w:pPr"))
    if ppr is None:
        return None
    el = ppr.find(qn("w:outlineLvl"))
    return el.get(qn("w:val")) if el is not None else None


def _style_by_id(doc, style_id):
    return next(s for s in doc.styles if s.style_id == style_id)


def _style_numpr(style):
    ppr = style.element.find(qn("w:pPr"))
    return ppr.find(qn("w:numPr")) if ppr is not None else None


def _style_relation(style, tag):
    element = style.element.find(qn(tag))
    return element.get(qn("w:val")) if element is not None else None


def _has_direct_child(paragraph, tag):
    ppr = paragraph._p.pPr
    return ppr is not None and ppr.find(qn(tag)) is not None


def _set_direct_numbering(paragraph, num_id_value, level_value):
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level_value))
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), str(num_id_value))
    num_pr.append(ilvl)
    num_pr.append(num_id)
    paragraph._p.get_or_add_pPr().append(num_pr)


with open(os.path.join(ROOT, "assets", "spec_std.json"), encoding="utf-8") as f:
    spec = json.load(f)
with open(os.path.join(ROOT, "assets", "rolemap_std.json"), encoding="utf-8") as f:
    rolemap = {int(k): v for k, v in json.load(f).items()}
validate_spec(spec)

with tempfile.TemporaryDirectory() as td:
    out = os.path.join(td, "styled.docx")
    messy_path = os.path.join(ROOT, "assets", "messy.docx")
    original_body_style_id = resolve_target_body_style(
        Document(messy_path), rolemap).style_id
    changes = apply_format(messy_path, spec, rolemap, out)
    doc = Document(out)

    expected_styles = {
        0: "文档标题",
        1: "格式正文",
        2: "标题 1",
        8: "落款",
        9: "日期",
    }
    for idx, name in expected_styles.items():
        assert doc.paragraphs[idx].style.name == name, (
            idx, doc.paragraphs[idx].style.name, name)
        assert _direct_value(doc.paragraphs[idx], "w:outlineLvl") is None

    assert _style_outline(_style_by_id(doc, "FormatAgentTitle")) == "0"
    assert _style_outline(_style_by_id(doc, "FormatAgentHeading1")) == "0"
    assert _style_outline(_style_by_id(doc, "FormatAgentBody")) is None

    # 全部 FormatAgent 管理样式都必须“基于：无样式”，且后续样式统一
    # 指向排版前解析出的目标正文样式。
    for role in spec["roles"]:
        managed = _style_by_id(doc, style_id_for_role(role))
        assert _style_relation(managed, "w:basedOn") is None
        assert _style_relation(managed, "w:next") == original_body_style_id

    # 字体和字号应位于样式定义，不再逐 run 直刷。
    title_rpr = _style_by_id(doc, "FormatAgentTitle").element.find(qn("w:rPr"))
    assert title_rpr.find(qn("w:rFonts")).get(qn("w:eastAsia")) == "方正小标宋简体"
    assert title_rpr.find(qn("w:sz")).get(qn("w:val")) == "44"
    run_rpr = doc.paragraphs[0].runs[0]._element.find(qn("w:rPr"))
    assert run_rpr is None or run_rpr.find(qn("w:rFonts")) is None
    assert run_rpr is None or run_rpr.find(qn("w:sz")) is None

    assert all(c.get("style_name") for c in changes)

    # 真正覆盖“源模板取样 → 在另一份目标文档创建同名样式并应用”的流程。
    with open(os.path.join(ROOT, "assets", "party_rolemap.json"), encoding="utf-8") as f:
        template_rolemap = {int(k): v for k, v in json.load(f).items()}
    template_spec = extract_rules_from_template(
        os.path.join(ROOT, "assets", "party_meeting.docx"), template_rolemap)
    assert template_spec["roles"]["heading_1"]["numbering"]["num_format"] == "chineseCounting"
    assert template_spec["roles"]["heading_1"]["numbering"]["level_text"] in {
        "%1、", "%1 、"
    }
    target = os.path.join(td, "target.docx")
    target_out = os.path.join(td, "target_styled.docx")
    target_doc = Document()
    target_doc.add_paragraph("目标文档标题")
    heading = target_doc.add_paragraph("一、目标文档一级标题")
    # 模拟问题文件中遗留的 linked character style：应用块级标题样式时必须清掉。
    target_doc.styles.add_style("UserStyle_1", WD_STYLE_TYPE.PARAGRAPH)
    char_style = target_doc.styles.add_style("LegacyCharacter", WD_STYLE_TYPE.CHARACTER)
    link = OxmlElement("w:link")
    link.set(qn("w:val"), "UserStyle_1")
    char_style.element.append(link)
    rpr = heading.runs[0]._element.get_or_add_rPr()
    rstyle = OxmlElement("w:rStyle")
    rstyle.set(qn("w:val"), char_style.style_id)
    rpr.append(rstyle)
    mark_rpr = OxmlElement("w:rPr")
    mark_rstyle = OxmlElement("w:rStyle")
    mark_rstyle.set(qn("w:val"), char_style.style_id)
    mark_rpr.append(mark_rstyle)
    heading._p.get_or_add_pPr().append(mark_rpr)
    target_doc.add_paragraph("这是一段需要套用模板正文样式的文字。")
    target_doc.save(target)
    apply_format(target, template_spec,
                 {0: "title", 1: "heading_1", 2: "body"}, target_out)
    target_result = Document(target_out)
    assert [p.style.name for p in target_result.paragraphs] == [
        "文档标题", "标题 1", "格式正文"]
    heading_style = _style_by_id(target_result, style_id_for_role("heading_1"))
    assert heading_style.font.size.pt == 16
    assert _style_by_id(target_result, style_id_for_role("body")).font.size.pt == 16
    assert target_result.paragraphs[1].text == "目标文档一级标题"
    assert _direct_value(target_result.paragraphs[1], "w:pStyle") == "FormatAgentHeading1"
    assert not _has_direct_child(target_result.paragraphs[1], "w:numPr")
    assert _style_numpr(heading_style) is not None
    out_rpr = target_result.paragraphs[1].runs[0]._element.find(qn("w:rPr"))
    assert out_rpr is None or out_rpr.find(qn("w:rStyle")) is None
    out_mark_rpr = target_result.paragraphs[1]._p.pPr.find(qn("w:rPr"))
    assert out_mark_rpr is None or out_mark_rpr.find(qn("w:rStyle")) is None

    numbering = target_result.part.numbering_part.element
    marker = None
    for abstract in numbering.findall(qn("w:abstractNum")):
        name = abstract.find(qn("w:name"))
        if name is not None and name.get(qn("w:val")) == "FormatAgent:headings":
            marker = abstract
            break
    assert marker is not None
    lvl0 = next(
        lvl for lvl in marker.findall(qn("w:lvl"))
        if lvl.get(qn("w:ilvl")) == "0")
    assert lvl0.find(qn("w:numFmt")).get(qn("w:val")) == "chineseCounting"
    assert lvl0.find(qn("w:pStyle")).get(qn("w:val")) == heading_style.style_id
    style_num_id = _style_numpr(heading_style).find(qn("w:numId")).get(qn("w:val"))
    assert int(style_num_id) > 0

    # numPr 必须按 CT_PPr 顺序位于 spacing/ind/jc/outlineLvl 之前。
    ppr_children = [child.tag for child in heading_style.element.find(qn("w:pPr"))]
    num_pos = ppr_children.index(qn("w:numPr"))
    for later_tag in ("w:spacing", "w:ind", "w:jc", "w:outlineLvl"):
        if qn(later_tag) in ppr_children:
            assert num_pos < ppr_children.index(qn(later_tag))
    if qn("w:spacing") in ppr_children and qn("w:ind") in ppr_children:
        assert ppr_children.index(qn("w:spacing")) < ppr_children.index(qn("w:ind"))
    if qn("w:ind") in ppr_children and qn("w:jc") in ppr_children:
        assert ppr_children.index(qn("w:ind")) < ppr_children.index(qn("w:jc"))
    if qn("w:jc") in ppr_children and qn("w:outlineLvl") in ppr_children:
        assert ppr_children.index(qn("w:jc")) < ppr_children.index(qn("w:outlineLvl"))

    # 对已经排版的文档再次执行，不得重复增长 FormatAgent 编号定义。
    second_out = os.path.join(td, "target_styled_twice.docx")
    before_abstracts = len(numbering.findall(qn("w:abstractNum")))
    before_nums = len(numbering.findall(qn("w:num")))
    apply_format(target_out, template_spec,
                 {0: "title", 1: "heading_1", 2: "body"}, second_out)
    twice = Document(second_out)
    assert len(twice.part.numbering_part.element.findall(qn("w:abstractNum"))) == before_abstracts
    assert len(twice.part.numbering_part.element.findall(qn("w:num"))) == before_nums

    # 用户已有同名样式时复用它，但绝不能原地改 styleId 破坏既有引用。
    same_name_src = os.path.join(td, "same_name.docx")
    same_name_out = os.path.join(td, "same_name_out.docx")
    same_name_doc = Document()
    existing = same_name_doc.styles.add_style("标题 1", WD_STYLE_TYPE.PARAGRAPH)
    existing_id = existing.style_id
    same_name_doc.add_paragraph("已有同名样式", style=existing)
    same_name_doc.add_paragraph("正文")
    same_name_doc.save(same_name_src)
    apply_format(same_name_src, spec, {0: "heading_1", 1: "body"}, same_name_out)
    reused_doc = Document(same_name_out)
    reused = reused_doc.paragraphs[0].style
    assert reused.name == "标题 1"
    assert reused.style_id == existing_id
    assert _style_relation(reused, "w:basedOn") is None
    assert _style_relation(reused, "w:next") == "Normal"

    # 模板缺少的角色必须回退到目标文档 body 段原有 pStyle 众数；即使模板
    # 明确定义了 other，也不能把 date 借用为 other。回退只改 pStyle，保留
    # 段落直接编号与 run 内加粗。
    fallback_src = os.path.join(td, "fallback_src.docx")
    fallback_out = os.path.join(td, "fallback_out.docx")
    fallback_doc = Document()
    target_body = fallback_doc.styles.add_style(
        "目标正文", WD_STYLE_TYPE.PARAGRAPH)
    target_body.paragraph_format.first_line_indent = Pt(24)
    minority_body = fallback_doc.styles.add_style(
        "少数正文", WD_STYLE_TYPE.PARAGRAPH)
    legacy = fallback_doc.styles.add_style(
        "旧日期样式", WD_STYLE_TYPE.PARAGRAPH)
    fallback_doc.add_paragraph("正文甲", style=target_body)
    fallback_doc.add_paragraph("正文乙", style=target_body)
    manual_body = fallback_doc.add_paragraph("1. 手工正文编号", style=minority_body)
    _set_direct_numbering(manual_body, 0, -1)
    missing_date = fallback_doc.add_paragraph("1. 手工日期文本", style=legacy)
    missing_date.runs[0].bold = True
    _set_direct_numbering(missing_date, 0, -1)
    real_numbered = fallback_doc.add_paragraph("真实自动编号项", style=legacy)
    _set_direct_numbering(real_numbered, 7, 0)
    fallback_doc.add_paragraph("显式其他", style=legacy)
    fallback_doc.save(fallback_src)

    fallback_spec = {
        "roles": {
            "body": {"font_eastasia": "宋体", "size_pt": 12,
                     "alignment": "justify", "first_line_indent_chars": 2},
            "other": {"font_eastasia": "黑体", "size_pt": 14,
                      "alignment": "left"},
        }
    }
    fallback_roles = {
        0: "body", 1: "body", 2: "body", 3: "date", 4: "signature",
        5: "other",
    }
    fallback_changes = apply_format(
        fallback_src, fallback_spec, fallback_roles, fallback_out)
    fallback_result = Document(fallback_out)
    resolved_body_id = target_body.style_id
    # 显式正文：手工“1.”文字保留，取消编号 numPr 删除，段落不再遮蔽
    # FormatAgentBody 的两字符首行缩进。
    assert fallback_result.paragraphs[2].text == "1. 手工正文编号"
    assert fallback_result.paragraphs[2]._p.pPr.find(qn("w:numPr")) is None
    formatted_body = _style_by_id(fallback_result, style_id_for_role("body"))
    body_ind = formatted_body.element.find(qn("w:pPr")).find(qn("w:ind"))
    assert body_ind.get(qn("w:firstLineChars")) == "200"
    assert fallback_result.paragraphs[2]._p.pPr.find(qn("w:ind")) is None

    # 缺失角色回退：与正文保持一致（套用 FormatAgentBody 的规则），
    # 文字保留、取消编号清除；bold 未被 body 规则管控，run 强调保留。
    assert fallback_result.paragraphs[3].style.style_id == style_id_for_role("body")
    assert fallback_result.paragraphs[3].text == "1. 手工日期文本"
    assert fallback_result.paragraphs[3]._p.pPr.find(qn("w:numPr")) is None
    fallback_body_ind = fallback_result.paragraphs[3].style.element.find(qn("w:pPr")).find(qn("w:ind"))
    assert fallback_body_ind.get(qn("w:firstLineChars")) == "200"

    # 真正的 numId>0 自动编号在缺失角色回退中同样完整保留。
    assert fallback_result.paragraphs[4].style.style_id == style_id_for_role("body")
    assert fallback_result.paragraphs[5].style.style_id == style_id_for_role("other")
    fallback_numpr = fallback_result.paragraphs[3]._p.pPr.find(qn("w:numPr"))
    assert fallback_numpr is None
    real_numpr = fallback_result.paragraphs[4]._p.pPr.find(qn("w:numPr"))
    assert real_numpr.find(qn("w:ilvl")).get(qn("w:val")) == "0"
    assert real_numpr.find(qn("w:numId")).get(qn("w:val")) == "7"
    fallback_rpr = fallback_result.paragraphs[3].runs[0]._element.find(qn("w:rPr"))
    assert fallback_rpr.find(qn("w:b")) is not None
    date_change = next(c for c in fallback_changes if c["idx"] == 3)
    assert date_change["fallback_to_target_body"] is True
    assert date_change["style_name"] == "格式正文"
    assert "invalid_numbering_removed" in date_change["changed_fields"]
    body_change = next(c for c in fallback_changes if c["idx"] == 2)
    assert "invalid_numbering_removed" in body_change["changed_fields"]
    real_change = next(c for c in fallback_changes if c["idx"] == 4)
    assert "invalid_numbering_removed" not in real_change["changed_fields"]
    other_change = next(c for c in fallback_changes if c["idx"] == 5)
    assert other_change["fallback_to_target_body"] is False

    for role in fallback_spec["roles"]:
        managed = _style_by_id(fallback_result, style_id_for_role(role))
        assert _style_relation(managed, "w:basedOn") is None
        assert _style_relation(managed, "w:next") == resolved_body_id

    # body 段没有显式 pStyle 时，每段都按默认样式参与众数统计；少量显式
    # “缩进正文”特例不能压过大量隐式 Normal 正文。
    default_src = os.path.join(td, "default_fallback_src.docx")
    default_out = os.path.join(td, "default_fallback_out.docx")
    default_doc = Document()
    default_doc.add_paragraph("没有显式 pStyle 的正文甲")
    default_doc.add_paragraph("没有显式 pStyle 的正文乙")
    indent_exception = default_doc.styles.add_style(
        "少量缩进正文", WD_STYLE_TYPE.PARAGRAPH)
    default_doc.add_paragraph("少数显式样式正文", style=indent_exception)
    default_doc.add_paragraph("模板未定义的日期")
    default_doc.save(default_src)
    apply_format(
        default_src, fallback_spec,
        {0: "body", 1: "body", 2: "body", 3: "date"}, default_out)
    default_result = Document(default_out)
    default_style_id = next(
        s.style_id for s in default_result.styles
        if s.type == WD_STYLE_TYPE.PARAGRAPH
        and s.element.get(qn("w:default")) in {"1", "true", "on"})
    # 模板未定义的角色回退到 FormatAgentBody（与正文保持一致）。
    assert default_result.paragraphs[3].style.style_id == style_id_for_role("body")
    for role in fallback_spec["roles"]:
        managed = _style_by_id(default_result, style_id_for_role(role))
        assert _style_relation(managed, "w:next") == default_style_id

print("命名样式流水线测试通过：样式关系、缺失角色回退、有效/取消编号与字符样式清理均正确")
