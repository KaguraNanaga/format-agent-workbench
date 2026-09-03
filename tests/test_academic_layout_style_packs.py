# -*- coding: utf-8 -*-
"""学术域、脚注、复杂页面、表格几何与英文 Style Pack 回归测试。"""

import os
import sys
import zipfile
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from core.agent import Agent
from core.apply import apply_format
from core.preflight import preflight_docx, scan_docx
from core.rules_from_template import extract_rules_from_template
from core.safe_output import replace_with_retry
from core.schema import validate_spec
from core.style_packs import get_style_pack, list_style_packs
from core.text_integrity import check_text_integrity


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
PR_NS = "http://schemas.openxmlformats.org/package/2006/relationships"


def _rule(size=11, alignment="left", **overrides):
    rule = {
        "font_eastasia": "Times New Roman",
        "font_ascii": "Times New Roman",
        "size_pt": size,
        "alignment": alignment,
    }
    rule.update(overrides)
    return rule


def _spec(**overrides):
    spec = {
        "profile": "english_academic",
        "cleanup": {"mode": "preserve_emphasis"},
        "roles": {"body": _rule()},
    }
    spec.update(overrides)
    return spec


def _append_field(paragraph, instruction, cached="1"):
    for kind in ("begin",):
        run = paragraph.add_run()
        field = OxmlElement("w:fldChar")
        field.set(qn("w:fldCharType"), kind)
        run._r.append(field)
    run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    run._r.append(instr)
    run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.append(separate)
    paragraph.add_run(cached)
    run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


def _append_split_field(paragraph, instruction_parts, cached="1"):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)
    for part in instruction_parts:
        run = paragraph.add_run()
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = part
        run._r.append(instr)
    run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.append(separate)
    paragraph.add_run(cached)
    run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


def _bookmark_run(paragraph, name, bookmark_id="1"):
    run = paragraph.runs[0]
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), bookmark_id)
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), bookmark_id)
    run._r.addprevious(start)
    run._r.addnext(end)


def _add_footnotes_part(path):
    with zipfile.ZipFile(path) as archive:
        payloads = {name: archive.read(name) for name in archive.namelist()}

    content_types = ET.fromstring(payloads["[Content_Types].xml"])
    ET.SubElement(content_types, f"{{{CT_NS}}}Override", {
        "PartName": "/word/footnotes.xml",
        "ContentType": (
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.footnotes+xml"
        ),
    })
    payloads["[Content_Types].xml"] = ET.tostring(
        content_types, encoding="utf-8", xml_declaration=True)

    rels = ET.fromstring(payloads["word/_rels/document.xml.rels"])
    ET.SubElement(rels, f"{{{PR_NS}}}Relationship", {
        "Id": "rIdFormatAgentFootnotes",
        "Type": (
            "http://schemas.openxmlformats.org/officeDocument/2006/"
            "relationships/footnotes"
        ),
        "Target": "footnotes.xml",
    })
    payloads["word/_rels/document.xml.rels"] = ET.tostring(
        rels, encoding="utf-8", xml_declaration=True)
    payloads["word/footnotes.xml"] = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:footnotes xmlns:w="{W_NS}">
  <w:footnote w:id="-1" w:type="separator"><w:p><w:r><w:t>separator</w:t></w:r></w:p></w:footnote>
  <w:footnote w:id="1"><w:p><w:r><w:t>A preserved note.</w:t></w:r></w:p></w:footnote>
</w:footnotes>""".encode("utf-8")

    rewritten = str(path) + ".rewritten"
    with zipfile.ZipFile(rewritten, "w", zipfile.ZIP_DEFLATED) as archive:
        for name, data in payloads.items():
            archive.writestr(name, data)
    replace_with_retry(rewritten, path)


def test_field_inventory_and_broken_cross_reference_detection(tmp_path):
    source = tmp_path / "fields.docx"
    output = tmp_path / "fields_out.docx"
    doc = Document()
    target = doc.add_paragraph("Cross-reference target")
    _bookmark_run(target, "TargetBookmark")
    valid = doc.add_paragraph("See ")
    _append_field(valid, "REF TargetBookmark \\h", "Cross-reference target")
    broken = doc.add_paragraph("Missing ")
    _append_field(broken, "PAGEREF MissingBookmark \\h", "Error!")
    citation = doc.add_paragraph()
    _append_field(citation, "CITATION Smith2025 \\l 1033", "(Smith, 2025)")
    authority_mark = doc.add_paragraph()
    _append_field(authority_mark, 'TA \\l "Example v. Sample" \\c 1', "")
    authorities = doc.add_paragraph()
    _append_field(authorities, "TOA \\h \\c 1", "Table of Authorities")
    index_mark = doc.add_paragraph()
    _append_field(index_mark, 'XE "document formatting"', "")
    index = doc.add_paragraph()
    _append_field(index, "INDEX \\c 2", "Index")
    doc.save(source)

    scan = scan_docx(source)
    assert scan["field_inventory"]["REF"] == 1
    assert scan["field_inventory"]["PAGEREF"] == 1
    assert scan["field_inventory"]["CITATION"] == 1
    assert scan["field_inventory"]["TA"] == 1
    assert scan["field_inventory"]["TOA"] == 1
    assert scan["field_inventory"]["XE"] == 1
    assert scan["field_inventory"]["INDEX"] == 1
    assert scan["broken_cross_references"] == ["MissingBookmark"]
    report = preflight_docx(source)
    codes = {risk["code"] for risk in report["warnings"]}
    assert "FIELD_REFRESH_REQUIRED" in codes
    assert "BROKEN_CROSS_REFERENCES" in codes
    apply_format(
        source,
        _spec(academic={"preserve_fields": True}),
        {index: "body" for index in range(8)},
        output,
    )
    output_scan = scan_docx(output)
    assert output_scan["field_inventory"]["REF"] == 1
    assert output_scan["field_inventory"]["PAGEREF"] == 1
    assert output_scan["field_inventory"]["CITATION"] == 1
    assert output_scan["field_inventory"]["TA"] == 1
    assert output_scan["field_inventory"]["TOA"] == 1
    assert output_scan["field_inventory"]["XE"] == 1
    assert output_scan["field_inventory"]["INDEX"] == 1
    assert check_text_integrity(source, output)["ok"]


def test_split_field_instruction_is_reassembled_for_cross_reference(tmp_path):
    source = tmp_path / "split_field.docx"
    doc = Document()
    target = doc.add_paragraph("Target")
    _bookmark_run(target, "SplitBookmark")
    reference = doc.add_paragraph("See ")
    _append_split_field(
        reference, [" REF ", "SplitBookmark", " \\h "], "Target")
    doc.save(source)

    scan = scan_docx(source)
    assert scan["field_inventory"] == {"REF": 1}
    assert scan["cross_reference_targets"] == ["SplitBookmark"]
    assert scan["broken_cross_references"] == []
    assert scan["story_counts"]["main"]["fields"] == 1


def test_caption_sequences_lists_and_visible_text_integrity(tmp_path):
    source = tmp_path / "captions.docx"
    output = tmp_path / "captions_out.docx"
    doc = Document()
    doc.add_paragraph("List of Figures")
    doc.add_paragraph("Figure 1. System architecture")
    doc.add_paragraph("List of Tables")
    doc.add_paragraph("Table 1. Evaluation results")
    doc.add_paragraph("Body text remains unchanged.")
    doc.save(source)

    spec = _spec(
        academic={
            "caption_numbering": True,
            "preserve_fields": True,
            "lists": {"figures": True, "tables": True},
        },
        roles={
            "body": _rule(),
            "figure_caption": _rule(9, "center"),
            "table_caption": _rule(9, "center"),
            "list_of_figures_heading": _rule(12, "center", bold=True),
            "list_of_tables_heading": _rule(12, "center", bold=True),
        },
    )
    rolemap = {
        0: "list_of_figures_heading", 1: "figure_caption",
        2: "list_of_tables_heading", 3: "table_caption", 4: "body",
    }
    changelog = apply_format(source, spec, rolemap, output)
    result = Document(output)
    visible = [paragraph.text for paragraph in result.paragraphs if paragraph.text]
    assert "Figure 1. System architecture" in visible
    assert "Table 1. Evaluation results" in visible
    instructions = [
        (element.text or "").strip()
        for element in result.element.body.iter(qn("w:instrText"))
    ]
    assert any("SEQ Figure" in instruction for instruction in instructions)
    assert any("SEQ Table" in instruction for instruction in instructions)
    assert any('TOC \\h \\z \\c "Figure"' in instruction for instruction in instructions)
    assert any('TOC \\h \\z \\c "Table"' in instruction for instruction in instructions)
    bookmarks = {
        element.get(qn("w:name"))
        for element in result.element.body.iter(qn("w:bookmarkStart"))
    }
    assert {"FormatAgentFigure1", "FormatAgentTable1"} <= bookmarks
    assert result.settings.element.find(qn("w:updateFields")).get(qn("w:val")) == "true"
    assert check_text_integrity(source, output)["ok"]
    academic_change = next(
        item for item in changelog if item["role"] == "academic_fields")
    assert "caption_sequence_fields_2" in academic_change["changed_fields"]


def test_footnote_story_formatting_preserves_note_text(tmp_path):
    source = tmp_path / "footnote.docx"
    output = tmp_path / "footnote_out.docx"
    doc = Document()
    doc.add_paragraph("Body with a footnote reference marker.")
    doc.save(source)
    _add_footnotes_part(source)
    spec = _spec(notes={
        "footnote": {
            "font_ascii": "Arial", "font_eastasia": "Arial",
            "size_pt": 9, "alignment": "left",
            "line_spacing": {"type": "exact", "pt": 12},
            "space_after_pt": 0,
        },
    })
    report = preflight_docx(source, spec=spec)
    assert "FORMAT_FOOTNOTES" in {item["code"] for item in report["warnings"]}
    apply_format(source, spec, {0: "body"}, output)

    with zipfile.ZipFile(output) as archive:
        root = ET.fromstring(archive.read("word/footnotes.xml"))
    actual = next(
        note for note in root.findall(f"{{{W_NS}}}footnote")
        if note.get(f"{{{W_NS}}}id") == "1")
    text = "".join(node.text or "" for node in actual.iter(f"{{{W_NS}}}t"))
    assert text == "A preserved note."
    spacing = actual.find(f".//{{{W_NS}}}spacing")
    assert spacing.get(f"{{{W_NS}}}line") == "240"
    fonts = actual.find(f".//{{{W_NS}}}rFonts")
    assert fonts.get(f"{{{W_NS}}}ascii") == "Arial"
    assert actual.find(f".//{{{W_NS}}}sz").get(f"{{{W_NS}}}val") == "18"
    integrity = check_text_integrity(source, output)
    assert integrity["ok"] is True
    assert integrity["story_differences"] == []


def test_table_geometry_and_merged_cells_are_preserved(tmp_path):
    source = tmp_path / "table.docx"
    output = tmp_path / "table_out.docx"
    doc = Document()
    doc.add_paragraph("Body")
    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).merge(table.cell(0, 1)).text = "Merged heading"
    table.cell(0, 2).text = "Heading 2"
    for index, cell in enumerate(table.rows[1].cells):
        cell.text = str(index + 1)
    doc.save(source)
    spec = _spec(
        table={
            "layout": "fixed", "alignment": "center", "width_pct": 90,
            "column_widths_pct": [40, 30, 30],
            "cell_margins_mm": {"top": 1, "bottom": 1, "left": 2, "right": 2},
            "repeat_header_row": True, "allow_row_break": False,
            "vertical_alignment": "center", "borders": True,
            "font_ascii": "Arial", "font_eastasia": "Arial", "size_pt": 9,
            "header_alignment": "center", "body_alignment": "right",
        },
    )
    apply_format(source, spec, {0: "body"}, output)
    result = Document(output)
    output_table = result.tables[0]
    tbl_pr = output_table._tbl.tblPr
    assert tbl_pr.find(qn("w:tblLayout")).get(qn("w:type")) == "fixed"
    assert tbl_pr.find(qn("w:tblW")).get(qn("w:type")) == "pct"
    assert tbl_pr.find(qn("w:tblW")).get(qn("w:w")) == "4500"
    assert output_table.rows[0]._tr.trPr.find(qn("w:tblHeader")) is not None
    assert output_table.rows[1]._tr.trPr.find(qn("w:cantSplit")) is not None
    grid_span = output_table.cell(0, 0)._tc.tcPr.find(qn("w:gridSpan"))
    assert grid_span is not None and grid_span.get(qn("w:val")) == "2"
    margins = tbl_pr.find(qn("w:tblCellMar"))
    assert margins.find(qn("w:left")).get(qn("w:w")) == str(round(2 / 25.4 * 1440))
    extracted = extract_rules_from_template(output, {0: "body"})
    assert extracted["table"]["layout"] == "fixed"
    assert extracted["table"]["width_pct"] == 90
    assert extracted["table"]["repeat_header_row"] is True
    assert len(extracted["table"]["column_widths_pct"]) == 3


def test_landscape_section_override_and_odd_even_first_headers(tmp_path):
    source = tmp_path / "sections.docx"
    output = tmp_path / "sections_out.docx"
    doc = Document()
    doc.add_paragraph("First section")
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_paragraph("Second section")
    doc.save(source)
    spec = _spec(page={
        "size": "letter", "orientation": "portrait",
        "margin": {"top_mm": 20, "bottom_mm": 20,
                   "left_mm": 20, "right_mm": 20},
        "section_overrides": [
            {"section_index": 1, "size": "A4", "orientation": "landscape"},
        ],
        "different_odd_even": True,
        "different_first_page": True,
        "header": {"text": "Odd", "page_number": True,
                   "font_ascii": "Arial", "size_pt": 9, "alignment": "right"},
        "even_header": {"text": "Even", "page_number": True,
                        "font_ascii": "Arial", "size_pt": 9, "alignment": "right"},
        "first_header": {"text": "First", "page_number": True,
                         "font_ascii": "Arial", "size_pt": 9, "alignment": "right"},
    })
    apply_format(source, spec, {0: "body", 1: "body", 2: "body"}, output)
    result = Document(output)
    assert len(result.sections) == 2
    assert result.sections[1].orientation == WD_ORIENT.LANDSCAPE
    assert result.sections[1].page_width > result.sections[1].page_height
    assert result.settings.odd_and_even_pages_header_footer is True
    assert result.sections[0].different_first_page_header_footer is True
    assert result.sections[0].header.paragraphs[0].text.startswith("Odd")
    assert result.sections[0].even_page_header.paragraphs[0].text.startswith("Even")
    assert result.sections[0].first_page_header.paragraphs[0].text.startswith("First")
    for story in (
        result.sections[0].header,
        result.sections[0].even_page_header,
        result.sections[0].first_page_header,
    ):
        assert any(
            (element.text or "").strip() == "PAGE"
            for element in story._element.iter(qn("w:instrText")))
    extracted = extract_rules_from_template(
        output, {index: "body" for index, _ in enumerate(result.paragraphs)})
    assert extracted["page"]["different_odd_even"] is True
    assert extracted["page"]["different_first_page"] is True
    assert extracted["page"]["header"]["text"] == "Odd"
    assert extracted["page"]["even_header"]["text"] == "Even"
    assert extracted["page"]["first_header"]["text"] == "First"


def test_multi_section_template_migrates_each_section_contract(tmp_path):
    template = tmp_path / "multi_section_template.docx"
    target = tmp_path / "multi_section_target.docx"
    output = tmp_path / "multi_section_output.docx"

    doc = Document()
    doc.add_paragraph("Template section one body")
    first = doc.sections[0]
    first.page_width, first.page_height = Mm(210), Mm(297)
    first.left_margin, first.right_margin = Mm(21), Mm(22)
    first.header.paragraphs[0].text = "First template header"
    second = doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_paragraph("Template section two body")
    second.orientation = WD_ORIENT.LANDSCAPE
    second.page_width, second.page_height = Mm(355.6), Mm(215.9)
    second.left_margin, second.right_margin = Mm(16), Mm(17)
    columns = second._sectPr.find(qn("w:cols"))
    columns.set(qn("w:num"), "2")
    page_number = OxmlElement("w:pgNumType")
    page_number.set(qn("w:fmt"), "upperRoman")
    page_number.set(qn("w:start"), "3")
    second._sectPr.insert_element_before(page_number, "w:cols", "w:docGrid")
    second.header.is_linked_to_previous = False
    second.header.paragraphs[0].text = "Second template header"
    doc.save(template)

    template_doc = Document(template)
    spec = extract_rules_from_template(
        template,
        {index: "body" for index, _ in enumerate(template_doc.paragraphs)},
    )
    validate_spec(spec)
    overrides = {
        item["section_index"]: item
        for item in spec["page"]["section_overrides"]
    }
    assert set(overrides) == {0, 1}
    assert overrides[1]["orientation"] == "landscape"
    assert overrides[1]["columns"] == 2
    assert overrides[1]["page_numbering"] == {
        "format": "upperRoman", "start": 3,
    }
    assert overrides[1]["header"]["text"] == "Second template header"

    target_doc = Document()
    target_doc.add_paragraph("Target section one body")
    target_doc.add_section(WD_SECTION.NEW_PAGE)
    target_doc.add_paragraph("Target section two body")
    target_doc.save(target)
    reopened = Document(target)
    apply_format(
        target, spec,
        {index: "body" for index, _ in enumerate(reopened.paragraphs)}, output)
    result = Document(output)
    assert result.sections[0].header.paragraphs[0].text == "First template header"
    assert result.sections[1].header.paragraphs[0].text == "Second template header"
    assert result.sections[1].header.is_linked_to_previous is False
    assert result.sections[1].orientation == WD_ORIENT.LANDSCAPE
    assert result.sections[1]._sectPr.find(qn("w:cols")).get(qn("w:num")) == "2"
    page_number = result.sections[1]._sectPr.find(qn("w:pgNumType"))
    assert page_number.get(qn("w:fmt")) == "upperRoman"
    assert page_number.get(qn("w:start")) == "3"


def test_template_uses_dominant_run_and_per_table_overrides(tmp_path):
    template = tmp_path / "heterogeneous_tables.docx"
    target = tmp_path / "heterogeneous_target.docx"
    output = tmp_path / "heterogeneous_output.docx"
    doc = Document()
    paragraph = doc.add_paragraph()
    label = paragraph.add_run("Label: ")
    label.bold = True
    label.font.size = Mm(3)
    body = paragraph.add_run("This much longer template body run controls the role style.")
    body.font.name = "Arial"
    body.font.size = Mm(4.23)  # approximately 12 pt
    first = doc.add_table(rows=2, cols=1)
    first.alignment = WD_TABLE_ALIGNMENT.LEFT
    first.cell(0, 0).text = "First header"
    first.cell(1, 0).text = "First body"
    second = doc.add_table(rows=2, cols=1)
    second.alignment = WD_TABLE_ALIGNMENT.CENTER
    second.autofit = False
    second.cell(0, 0).text = "Second header"
    second.cell(1, 0).text = "Second body"
    doc.save(template)

    spec = extract_rules_from_template(template, {0: "body"})
    assert spec["roles"]["body"]["font_ascii"] == "Arial"
    assert round(spec["roles"]["body"]["size_pt"]) == 12
    assert spec["roles"]["body"].get("bold") is not True
    assert spec["table"]["alignment"] == "left"
    assert spec["table"]["overrides"][0]["table_index"] == 1
    assert spec["table"]["overrides"][0]["alignment"] == "center"

    target_doc = Document()
    target_doc.add_paragraph("Target body")
    for header, value in (("A", "1"), ("B", "2")):
        table = target_doc.add_table(rows=2, cols=1)
        table.cell(0, 0).text = header
        table.cell(1, 0).text = value
    target_doc.save(target)
    apply_format(target, spec, {0: "body"}, output)
    result = Document(output)
    assert result.tables[0].alignment == WD_TABLE_ALIGNMENT.LEFT
    assert result.tables[1].alignment == WD_TABLE_ALIGNMENT.CENTER


def test_complex_script_style_and_custom_page_size_are_serialized_safely(tmp_path):
    source = tmp_path / "rtl.docx"
    output = tmp_path / "rtl_out.docx"
    doc = Document()
    doc.add_paragraph("هذه فقرة عربية للاختبار")
    doc.save(source)
    spec = _spec(
        page={"width_mm": 216, "height_mm": 330, "orientation": "portrait"},
        roles={"body": _rule(
            font_cs="Traditional Arabic", language="ar-SA",
            rtl=True, bidi=True, small_caps=False,
        )},
    )
    validate_spec(spec)
    apply_format(source, spec, {0: "body"}, output)
    result = Document(output)
    assert abs(result.sections[0].page_width.mm - 216) < 0.2
    assert abs(result.sections[0].page_height.mm - 330) < 0.2
    style = result.styles["格式正文"]
    rpr = style.element.find(qn("w:rPr"))
    tags = [element.tag.rsplit("}", 1)[-1] for element in rpr]
    expected_order = [
        "rFonts", "b", "bCs", "i", "iCs", "caps", "smallCaps",
        "strike", "sz", "szCs", "highlight", "u", "rtl", "cs", "lang",
    ]
    ranks = {name: index for index, name in enumerate(expected_order)}
    actual_ranks = [ranks[tag] for tag in tags if tag in ranks]
    assert actual_ranks == sorted(actual_ranks)
    fonts = rpr.find(qn("w:rFonts"))
    assert fonts.get(qn("w:cs")) == "Traditional Arabic"
    language = rpr.find(qn("w:lang"))
    assert language.get(qn("w:val")) == "ar-SA"
    assert language.get(qn("w:bidi")) == "ar-SA"
    assert rpr.find(qn("w:rtl")) is not None
    ppr = style.element.find(qn("w:pPr"))
    assert ppr.find(qn("w:bidi")) is not None


def test_explicit_table_is_wrapped_in_landscape_section(tmp_path):
    source = tmp_path / "landscape_table.docx"
    output = tmp_path / "landscape_table_out.docx"
    doc = Document()
    doc.add_paragraph("Before table")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "2"
    doc.add_paragraph("After table")
    doc.save(source)

    spec = _spec(table={"landscape_table_indices": [0]})
    apply_format(source, spec, {0: "body", 1: "body"}, output,
                 allow_risky_structure=True)
    result = Document(output)
    assert len(result.sections) == 3
    assert result.sections[0].orientation != WD_ORIENT.LANDSCAPE
    assert result.sections[1].orientation == WD_ORIENT.LANDSCAPE
    assert result.sections[2].orientation != WD_ORIENT.LANDSCAPE
    assert result.tables[0].cell(0, 0).text == "A"
    assert result.tables[0].cell(1, 1).text == "2"
    assert check_text_integrity(source, output)["ok"]


def test_style_packs_validate_and_expose_expected_baselines():
    assert set(list_style_packs()) == {
        "apa7-student", "mla9", "ieee-journal",
        "official-cn-gbt9704", "chicago18-notes-bibliography",
        "chicago18-author-date", "turabian9-student",
        "technical-manual", "us-legal-brief",
    }
    for name in list_style_packs():
        validate_spec(get_style_pack(name))
    apa = get_style_pack("apa7-student")
    assert apa["page"]["margin"]["left_mm"] == 25.4
    assert apa["roles"]["body"]["line_spacing"]["pt"] == 2.0
    assert apa["roles"]["bibliography_entry"]["hanging_indent_chars"] == 3
    mla = get_style_pack("mla9", running_head="Rivera")
    assert mla["page"]["header"]["text"] == "Rivera"
    assert mla["page"]["header"]["page_number"] is True
    ieee = get_style_pack("ieee-journal")
    assert ieee["page"]["columns"] == 2
    assert ieee["roles"]["bibliography_entry"]["size_pt"] == 8
    chicago = get_style_pack("chicago18-notes-bibliography")
    assert chicago["notes"]["footnote"]["size_pt"] == 10
    assert chicago["roles"]["block_quote"]["line_spacing"]["pt"] == 1.0
    official = get_style_pack("official-cn-gbt9704")
    assert official["profile"] == "official_cn"
    assert official["page"]["margin"]["top_mm"] == 37
    technical = get_style_pack("technical-manual")
    assert technical["roles"]["warning_box"]["shading"] == "FEE2E2"
    legal = get_style_pack("us-legal-brief")
    assert legal["legal"]["preserve_toa"] is True


def test_agent_runs_style_pack_without_llm_and_commits_atomic_outputs(tmp_path):
    source = tmp_path / "paper.docx"
    output = tmp_path / "paper_mla.docx"
    doc = Document()
    doc.add_paragraph("A Sample Paper")
    doc.add_paragraph("This paragraph must remain unchanged.")
    doc.add_paragraph("Works Cited")
    doc.add_paragraph("Rivera, Alex. Example Source.")
    doc.save(source)

    result = Agent().run(
        source,
        output,
        style_pack="mla9",
        style_pack_options={"running_head": "Rivera"},
        rolemap={
            0: "title", 1: "body", 2: "bibliography_heading",
            3: "bibliography_entry",
        },
    )
    assert result["spec"]["style_pack"] == "mla9"
    assert result["text_integrity"]["ok"] is True
    assert output.is_file()
    assert (tmp_path / "paper_mla_tracked.docx").is_file()
    assert (tmp_path / "paper_mla_report.docx").is_file()
    rendered = Document(output)
    assert rendered.sections[0].header.paragraphs[0].text.startswith("Rivera")
