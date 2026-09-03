"""可重复的真实环境验收：生成 DOCX 基准稿并跑完整排版流水线。

这个脚本不伪造 Office 转换器。Word/WPS/LibreOffice 的探测和旧格式
往返结果由调用方写入 ``environment``；本脚本负责不依赖 LLM 的原生
DOCX 端到端验收，并输出机器可读清单供渲染阶段复核。
"""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from core.agent import Agent
from core.extract import extract_paragraphs
from core.text_integrity import check_text_integrity, protected_story_texts


SPEC = {
    "profile": "general",
    "locale": "zh-CN",
    "cleanup": {"mode": "preserve_emphasis"},
    "roles": {
        "title": {
            "font_eastasia": "黑体",
            "font_ascii": "Arial",
            "size_pt": 22,
            "alignment": "center",
            "bold": True,
            "space_after_pt": 12,
        },
        "heading_1": {
            "font_eastasia": "黑体",
            "font_ascii": "Arial",
            "size_pt": 16,
            "alignment": "left",
            "bold": True,
            "keep_with_next": True,
        },
        "body": {
            "font_eastasia": "宋体",
            "font_ascii": "Times New Roman",
            "size_pt": 12,
            "alignment": "justify",
            "first_line_indent_chars": 2,
            "line_spacing": {"type": "exact", "pt": 22},
        },
    },
    "table": {
        "font_eastasia": "宋体",
        "font_ascii": "Times New Roman",
        "size_pt": 10.5,
        "alignment": "center",
        "layout": "autofit",
        "width_pct": 100,
        "repeat_header_row": True,
        "allow_row_break": False,
        "header_alignment": "center",
        "body_alignment": "left",
        "vertical_alignment": "center",
    },
}


def _set_run_fonts(run, east_asia: str, ascii_font: str) -> None:
    run.font.name = ascii_font
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)


def _append_page_field(paragraph) -> None:
    paragraph.add_run("第 ")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run = paragraph.add_run()
    begin_run._r.append(begin)
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    instruction_run = paragraph.add_run()
    instruction_run._r.append(instruction)
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run = paragraph.add_run()
    separate_run._r.append(separate)
    paragraph.add_run("1")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run = paragraph.add_run()
    end_run._r.append(end)
    paragraph.add_run(" 页 / Page")


def build_reference(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(24)
    section.bottom_margin = Mm(24)
    section.left_margin = Mm(28)
    section.right_margin = Mm(24)

    header = section.header.paragraphs[0]
    header.text = "FORMAT AGENT / 真实环境验收"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in header.runs:
        _set_run_fonts(run, "黑体", "Arial")
        run.font.size = Pt(9)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _append_page_field(footer)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Format Agent 真实环境验收")
    _set_run_fonts(run, "黑体", "Arial")
    run.font.size = Pt(20)
    run.bold = True

    heading = doc.add_paragraph("一、转换基准 / Conversion Baseline")
    heading.style = doc.styles["Heading 1"]

    body = doc.add_paragraph()
    body.add_run("中文完整性：甲乙丙，标点“引号”（括号）及 2026-09-02。")
    emphasized = body.add_run(" English integrity: bold and italic survive formatting.")
    emphasized.bold = True
    emphasized.italic = True

    doc.add_paragraph(
        "第二段用于验证正文顺序、空白归一与跨语言字体。"
        " The second paragraph verifies order, whitespace, and mixed-script fonts."
    )

    table = doc.add_table(rows=3, cols=3)
    table.style = "Table Grid"
    values = (
        ("项目 / Item", "中文值", "English value"),
        ("编号 / ID", "A-001", "42.50"),
        ("状态 / Status", "待验收", "Pending"),
    )
    for row, row_values in zip(table.rows, values):
        for cell, value in zip(row.cells, row_values):
            cell.text = value

    doc.add_page_break()
    doc.add_paragraph("第二页固定文本 / Fixed text on page two.")
    doc.add_paragraph("页眉、页脚和分页必须保留；表格不得移动到正文末尾之外。")
    doc.save(path)


def _inventory(path: Path) -> dict:
    doc = Document(path)
    section = doc.sections[0]
    table_values = [
        [[paragraph.text for paragraph in cell.paragraphs] for cell in row.cells]
        for table in doc.tables
        for row in table.rows
    ]
    return {
        "main_paragraphs": [paragraph.text for paragraph in doc.paragraphs],
        "table_rows": table_values,
        "stories": protected_story_texts(path),
        "sections": len(doc.sections),
        "page_mm": {
            "width": round(section.page_width.mm, 2),
            "height": round(section.page_height.mm, 2),
            "top": round(section.top_margin.mm, 2),
            "bottom": round(section.bottom_margin.mm, 2),
            "left": round(section.left_margin.mm, 2),
            "right": round(section.right_margin.mm, 2),
        },
    }


def _rolemap(path: Path) -> dict[int, str]:
    result = {}
    for item in extract_paragraphs(path):
        if not item.get("editable", True) or item.get("in_table"):
            continue
        text = item.get("text", "").strip()
        if text == "Format Agent 真实环境验收":
            role = "title"
        elif text.startswith("一、转换基准"):
            role = "heading_1"
        else:
            role = "body"
        result[item["idx"]] = role
    return result


def run(output_dir: Path) -> dict:
    output_dir.mkdir(parents=True, exist_ok=True)
    reference = output_dir / "reference.docx"
    final = output_dir / "native_final.docx"
    build_reference(reference)

    events = []
    result = Agent(on_event=events.append).run(
        str(reference), str(final), spec=SPEC, rolemap=_rolemap(reference)
    )
    reference_inventory = _inventory(reference)
    final_inventory = _inventory(final)
    semantic_integrity = check_text_integrity(reference, final)
    structural_checks = {
        "section_count_preserved": (
            reference_inventory["sections"] == final_inventory["sections"]
        ),
        "page_geometry_preserved": (
            reference_inventory["page_mm"] == final_inventory["page_mm"]
        ),
        "table_order_and_text_preserved": (
            reference_inventory["table_rows"] == final_inventory["table_rows"]
        ),
        "protected_story_text_preserved": (
            reference_inventory["stories"] == final_inventory["stories"]
        ),
    }
    accepted = (
        result["text_integrity"]["ok"]
        and semantic_integrity["ok"]
        and all(structural_checks.values())
    )
    payload = {
        "accepted": accepted,
        "reference": str(reference.resolve()),
        "final": str(final.resolve()),
        "native_pipeline": {
            "status": "PASS" if accepted else "FAIL",
            "agent_text_integrity": result["text_integrity"],
            "independent_text_integrity": semantic_integrity,
            "structural_checks": structural_checks,
            "preflight": result["preflight"],
            "event_log": events,
            "outputs": {
                "main": result["out_path"],
                "tracked": result["tracked_path"],
                "report_docx": result["report_docx_path"],
                "report_md": result["report_path"],
            },
        },
        "inventory": {
            "reference": reference_inventory,
            "final": final_inventory,
        },
    }
    (output_dir / "native_acceptance.json").write_text(
        json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    return payload


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--output-dir", type=Path, required=True)
    args = parser.parse_args()
    payload = run(args.output_dir.resolve())
    print(json.dumps(payload, ensure_ascii=False, indent=2))
    return 0 if payload["accepted"] else 1


if __name__ == "__main__":
    raise SystemExit(main())
