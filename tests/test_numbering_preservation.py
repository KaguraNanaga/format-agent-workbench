"""Regression coverage for target-side automatic numbering preservation."""

import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from core.apply import apply_format
from core.extract import paragraph_numbering_metadata


def _num_pr(num_id=1, level=0):
    value = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    value.append(ilvl)
    value.append(number)
    return value


def test_valid_direct_and_style_numbering_survive_uncontrolled_formatting(tmp_path):
    source = tmp_path / "numbered-headings.docx"
    output = tmp_path / "numbered-headings-formatted.docx"
    document = Document()

    direct = document.add_paragraph("直接编号标题")
    direct._p.get_or_add_pPr()._insert_numPr(_num_pr())

    source_style = document.styles.add_style(
        "源文档编号标题", WD_STYLE_TYPE.PARAGRAPH)
    source_style.element.get_or_add_pPr()._insert_numPr(_num_pr())
    inherited = document.add_paragraph("样式编号标题", style=source_style)
    document.add_paragraph("正文")
    document.save(source)

    before = Document(source)
    assert paragraph_numbering_metadata(before.paragraphs[0])["numbering_status"] == "automatic"
    assert paragraph_numbering_metadata(before.paragraphs[1])["numbering_status"] == "automatic"

    spec = {
        "roles": {
            "heading_1": {
                "font_eastasia": "黑体",
                "font_ascii": "Times New Roman",
                "size_pt": 16,
                "bold": True,
                "alignment": "left",
            },
            "body": {
                "font_eastasia": "仿宋_GB2312",
                "font_ascii": "Times New Roman",
                "size_pt": 16,
                "alignment": "justify",
            },
        }
    }
    changes = apply_format(
        source, spec, {0: "heading_1", 1: "heading_1", 2: "body"}, output)
    result = Document(output)

    for paragraph in result.paragraphs[:2]:
        metadata = paragraph_numbering_metadata(paragraph)
        assert metadata["numbering_status"] == "automatic"
        assert metadata["num_id"] == 1
        assert paragraph._p.pPr.find(qn("w:numPr")) is not None
        assert paragraph.style.font.size.pt == 16

    for index in (0, 1):
        change = next(item for item in changes if item["idx"] == index)
        assert "automatic_numbering_preserved" in change["changed_fields"]
