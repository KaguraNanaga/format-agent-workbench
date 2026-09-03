"""Regression coverage: tracked documents flatten the applied style into
explicit direct formatting so WPS can accept revisions without losing fonts.

背景：修订稿早期让“当前格式”完全依赖命名样式继承，WPS 接受全部修订后
字体解析失败并回退到默认字体。修复后修订稿在保留 pStyle 的同时把新格式
写成直接格式（与 Word 原生录制的格式修订一致）；干净稿不受影响。
"""

import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from docx.oxml.ns import qn

from core.apply import apply_format

SPEC = {
    "roles": {
        "body": {
            "font_eastasia": "仿宋_GB2312",
            "font_ascii": "Times New Roman",
            "size_pt": 16,
            "alignment": "justify",
            "first_line_indent_chars": 2,
            "line_spacing": {"type": "exact", "pt": 28},
        },
    }
}


def _current_rpr_children(run_element):
    """run 当前 rPr 中、修订记录之外的子元素标签集合。"""
    rpr = run_element.find(qn("w:rPr"))
    if rpr is None:
        return rpr, set()
    return rpr, {c.tag for c in rpr if c.tag != qn("w:rPrChange")}


def test_tracked_doc_flattens_style_into_direct_formatting(tmp_path):
    source = tmp_path / "src.docx"
    tracked = tmp_path / "tracked.docx"
    clean = tmp_path / "clean.docx"

    document = Document()
    messy = document.add_paragraph()
    run = messy.add_run("带旧直接格式的正文。")
    run.font.name = "微软雅黑"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "微软雅黑")
    plain = document.add_paragraph()
    plain.add_run("完全没有字符属性的正文。")
    document.save(source)

    apply_format(str(source), SPEC, {0: "body", 1: "body"}, str(tracked),
                 track=True)
    apply_format(str(source), SPEC, {0: "body", 1: "body"}, str(clean),
                 track=False)

    result = Document(tracked)
    for paragraph in result.paragraphs:
        ppr = paragraph._p.pPr
        # 段落级：样式里的行距/缩进/对齐摊平为当前直接格式
        assert ppr.find(qn("w:spacing")) is not None
        assert ppr.find(qn("w:ind")) is not None
        assert ppr.find(qn("w:jc")) is not None
        # 段落标记同样补齐字体
        mark_fonts = ppr.find(qn("w:rPr")).find(qn("w:rFonts"))
        assert mark_fonts is not None
        assert mark_fonts.get(qn("w:eastAsia")) == "仿宋_GB2312"
        for run in paragraph.runs:
            rpr, tags = _current_rpr_children(run._element)
            assert rpr is not None
            fonts = rpr.find(qn("w:rFonts"))
            assert fonts is not None
            assert fonts.get(qn("w:eastAsia")) == "仿宋_GB2312"
            assert fonts.get(qn("w:ascii")) == "Times New Roman"
            assert qn("w:sz") in tags
            # 有修订记录时必须放在 rPr 末尾
            change = rpr.find(qn("w:rPrChange"))
            if change is not None:
                assert rpr[-1].tag == qn("w:rPrChange")

    # 旧格式仍保存在修订记录里（拒绝修订可回退）
    first_change = result.paragraphs[0].runs[0]._element.find(
        qn("w:rPr")).find(qn("w:rPrChange"))
    old_fonts = first_change.find(qn("w:rPr")).find(qn("w:rFonts"))
    assert old_fonts.get(qn("w:eastAsia")) == "微软雅黑"

    # 干净稿保持纯样式继承，不写 FormatAgent 直接字体
    clean_doc = Document(clean)
    for paragraph in clean_doc.paragraphs:
        for run in paragraph.runs:
            fonts = run._element.find(qn("w:rPr"))
            if fonts is not None and fonts.find(qn("w:rFonts")) is not None:
                assert fonts.find(qn("w:rFonts")).get(
                    qn("w:eastAsia")) != "仿宋_GB2312"
