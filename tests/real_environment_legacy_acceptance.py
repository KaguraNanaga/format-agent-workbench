"""在真实 Word/WPS/LibreOffice 环境运行旧格式完整流水线。"""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))
sys.path.insert(0, str(Path(__file__).resolve().parent))

from core.agent import Agent
from core.text_integrity import paragraph_texts, protected_story_texts
from real_environment_acceptance import SPEC


class DeterministicBodyLabeler:
    """验收只测排版/转换，不引入网络 LLM 的不确定性。"""

    def chat_json(self, prompt):
        payload_text = prompt.split("段落清单：\n", 1)[1]
        payload, _ = json.JSONDecoder().raw_decode(payload_text)
        result = []
        for item in payload:
            text = str(item.get("text") or "").strip()
            if text == "Format Agent 真实环境验收":
                role = "title"
            elif text.startswith("一、转换基准"):
                role = "heading_1"
            else:
                role = "body"
            result.append({"idx": item["idx"], "role": role})
        return result


def _table_rows(path: Path) -> list:
    doc = Document(path)
    return [
        [[paragraph.text for paragraph in cell.paragraphs] for cell in row.cells]
        for table in doc.tables
        for row in table.rows
    ]


def _page_geometry(path: Path) -> dict:
    section = Document(path).sections[0]
    return {
        "width_mm": round(section.page_width.mm, 2),
        "height_mm": round(section.page_height.mm, 2),
        "top_mm": round(section.top_margin.mm, 2),
        "bottom_mm": round(section.bottom_margin.mm, 2),
        "left_mm": round(section.left_margin.mm, 2),
        "right_mm": round(section.right_margin.mm, 2),
    }


def run(source_dir: Path, reference: Path, output_dir: Path) -> dict:
    output_dir.mkdir(parents=True, exist_ok=True)
    expected_text = paragraph_texts(reference)
    expected_tables = _table_rows(reference)
    expected_stories = protected_story_texts(reference)
    expected_geometry = _page_geometry(reference)
    results = {}

    for extension in ("doc", "wps", "odt", "rtf"):
        print(f"RUN {extension}", flush=True)
        source = source_dir / f"reference.{extension}"
        final = output_dir / f"{extension}_final.docx"
        events = []
        try:
            result = Agent(
                llm=DeterministicBodyLabeler(), on_event=events.append
            ).run(str(source), str(final), spec=SPEC)
            actual_text = paragraph_texts(final)
            actual_tables = _table_rows(final)
            actual_stories = protected_story_texts(final)
            body_exact = actual_text == expected_text
            table_exact = actual_tables == expected_tables
            # 手工构造的 ODT 基线没有页眉页脚；其余格式应往返保留。
            story_exact = (
                not actual_stories["headers"] and not actual_stories["footers"]
                if extension == "odt"
                else actual_stories == expected_stories
            )
            geometry_exact = _page_geometry(final) == expected_geometry
            roundtrip_exact = body_exact and table_exact and story_exact
            pipeline_ok = bool(result["text_integrity"]["ok"])
            status = "PASS" if pipeline_ok and roundtrip_exact else (
                "WARN" if pipeline_ok else "FAIL"
            )
            results[extension] = {
                "status": status,
                "source": str(source.resolve()),
                "source_size": source.stat().st_size,
                "converter": result["input_conversion"]["target"]["converter"],
                "pipeline_text_integrity": result["text_integrity"],
                "roundtrip": {
                    "body_text_exact": body_exact,
                    "table_order_and_text_exact": table_exact,
                    "protected_story_exact": story_exact,
                    "page_geometry_exact": geometry_exact,
                    "expected_text": expected_text if not body_exact else None,
                    "actual_text": actual_text if not body_exact else None,
                    "expected_stories": expected_stories if not story_exact else None,
                    "actual_stories": actual_stories if not story_exact else None,
                    "actual_page_geometry": (
                        _page_geometry(final) if not geometry_exact else None
                    ),
                },
                "preflight": {
                    "ok": result["preflight"]["ok"],
                    "section_count": result["preflight"]["section_count"],
                    "story_parts": len(result["preflight"]["story_parts"]),
                    "warnings": result["preflight"]["warnings"],
                    "blockers": result["preflight"]["blockers"],
                },
                "final": str(final.resolve()),
                "tracked": str(Path(result["tracked_path"]).resolve()),
                "report_docx": str(Path(result["report_docx_path"]).resolve()),
                "errors_in_event_log": [
                    event for event in events if event.get("status") == "err"
                ],
            }
        except Exception as exc:  # noqa: BLE001 - 验收必须收集各格式失败
            results[extension] = {
                "status": "FAIL",
                "source": str(source.resolve()),
                "error_type": type(exc).__name__,
                "error": str(exc),
                "errors_in_event_log": [
                    event for event in events if event.get("status") == "err"
                ],
            }
        print(f"RESULT {extension} {results[extension]['status']}", flush=True)

    payload = {
        "results": results,
        "summary": {
            key: sum(1 for value in results.values() if value["status"] == key)
            for key in ("PASS", "WARN", "FAIL")
        },
    }
    (output_dir / "legacy_acceptance.json").write_text(
        json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    return payload


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source-dir", type=Path, required=True)
    parser.add_argument("--reference", type=Path, required=True)
    parser.add_argument("--output-dir", type=Path, required=True)
    args = parser.parse_args()
    payload = run(
        args.source_dir.resolve(), args.reference.resolve(), args.output_dir.resolve()
    )
    print(json.dumps(payload, ensure_ascii=False, indent=2))
    return 1 if payload["summary"]["FAIL"] else 0


if __name__ == "__main__":
    raise SystemExit(main())
