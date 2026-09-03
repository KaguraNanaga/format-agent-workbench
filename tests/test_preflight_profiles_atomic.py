# -*- coding: utf-8 -*-
"""Story 预检、多节保留、英文 Profile 与失败闭环回归测试。"""

import os
import json
import sys
import zipfile

import pytest
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml import OxmlElement
from docx.shared import Inches, Mm, Pt

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from core.agent import Agent
from core.apply import apply_format
from core.extract import extract_paragraphs
from core.label_roles import detect_document_profile, regex_role
from core.preflight import (
    PreflightBlockedError, merge_preflight_reports, preflight_docx, scan_docx,
)
from core.rules_from_template import extract_rules_from_template
from core.safe_output import IntegrityViolationError, UnsafeOutputPathError
from core.schema import validate_spec
from core.text_integrity import check_text_integrity


def _body_rule():
    return {
        "font_eastasia": "Times New Roman",
        "font_ascii": "Times New Roman",
        "size_pt": 12,
        "alignment": "justify",
    }


def _spec(profile="general"):
    return {
        "profile": profile,
        "cleanup": {"mode": "preserve_emphasis"},
        "page": {
            "margin": {
                "top_mm": 25,
                "bottom_mm": 25,
                "left_mm": 30,
                "right_mm": 20,
            },
            "line_grid": {"line_pt": 20},
        },
        "roles": {"body": _body_rule()},
    }


def _write_story_fixture(path):
    doc = Document()
    paragraph = doc.add_paragraph("Visible body text for the preflight fixture.")
    paragraph._p.append(OxmlElement("w:sdt"))
    paragraph._p.append(OxmlElement("w:txbxContent"))
    inserted = OxmlElement("w:ins")
    inserted.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id", "1")
    paragraph._p.append(inserted)
    doc.add_table(rows=1, cols=1).cell(0, 0).text = "table"
    doc.sections[0].header.paragraphs[0].text = "Header story"
    doc.sections[0].footer.paragraphs[0].text = "Footer story"
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_paragraph("Second section")
    doc.save(path)
    with zipfile.ZipFile(path, "a") as archive:
        archive.writestr(
            "word/footnotes.xml",
            """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:footnote w:id="1"><w:p><w:r><w:t>Footnote story</w:t></w:r></w:p></w:footnote>
</w:footnotes>""",
        )


def test_full_story_scan_and_preflight_blockers(tmp_path):
    source = tmp_path / "stories.docx"
    _write_story_fixture(source)

    scan = scan_docx(source)
    assert scan["section_count"] == 2
    assert {part["kind"] for part in scan["story_parts"]} >= {
        "main", "header", "footer", "footnotes",
    }
    assert scan["story_counts"]["main"]["tables"] == 1
    assert scan["story_counts"]["main"]["textboxes"] == 1
    assert scan["story_counts"]["main"]["content_controls"] == 1
    assert scan["story_counts"]["main"]["tracked_changes"] == 1
    assert scan["story_counts"]["footnotes"]["entries"] == 1

    report = preflight_docx(source)
    blocker_codes = {item["code"] for item in report["blockers"]}
    warning_codes = {item["code"] for item in report["warnings"]}
    assert "SOURCE_TRACKED_CHANGES" in blocker_codes
    assert "PRESERVE_ONLY_FOOTNOTES" in warning_codes
    assert next(
        item for item in report["warnings"]
        if item["code"] == "PRESERVE_ONLY_FOOTNOTES"
    )["count"] == 1
    assert "PRESERVE_ONLY_TEXTBOXES" in warning_codes
    assert "PRESERVE_ONLY_CONTENT_CONTROLS" in warning_codes
    assert report["ok"] is False


def test_ordered_nested_table_and_protected_story_extraction(tmp_path):
    source = tmp_path / "ordered_stories.docx"
    copied = tmp_path / "ordered_stories_copy.docx"
    doc = Document()
    doc.add_paragraph("Before table")
    outer = doc.add_table(rows=1, cols=1)
    cell = outer.cell(0, 0)
    cell.paragraphs[0].text = "Outer cell before"
    nested = cell.add_table(rows=1, cols=1)
    nested.cell(0, 0).text = "Nested cell"
    cell.add_paragraph("Outer cell after")
    doc.add_paragraph("After table")
    doc.sections[0].header.paragraphs[0].text = "Header story"
    doc.sections[0].footer.paragraphs[0].text = "Footer story"

    control = OxmlElement("w:sdt")
    content = OxmlElement("w:sdtContent")
    control_paragraph = OxmlElement("w:p")
    control_run = OxmlElement("w:r")
    control_text = OxmlElement("w:t")
    control_text.text = "Controlled story text"
    control_run.append(control_text)
    control_paragraph.append(control_run)
    content.append(control_paragraph)
    control.append(content)
    doc.element.body.insert(len(doc.element.body) - 1, control)
    doc.save(source)

    records = extract_paragraphs(source)
    main_texts = [
        record["text"] for record in records
        if record.get("story") == "main" and record["text"]
    ]
    assert main_texts == [
        "Before table", "Outer cell before", "Nested cell",
        "Outer cell after", "After table",
    ]
    assert [record["idx"] for record in records] == list(range(len(records)))
    assert next(
        record for record in records if record["text"] == "Nested cell"
    )["table_depth"] == 2
    protected = {
        (record["story"], record["text"])
        for record in records if not record.get("editable", True)
    }
    assert ("header", "Header story") in protected
    assert ("footer", "Footer story") in protected
    assert ("content_control", "Controlled story text") in protected

    # A load/save cycle must preserve every audited Story and nested cell text.
    Document(source).save(copied)
    assert check_text_integrity(source, copied)["ok"] is True


def test_target_and_template_preflight_risks_are_distinguished(tmp_path):
    target = tmp_path / "target.docx"
    template = tmp_path / "template.docx"
    plain = Document()
    plain.add_paragraph("Plain target content")
    plain.save(target)
    _write_story_fixture(template)

    merged = merge_preflight_reports(
        preflight_docx(target), preflight_docx(template))
    assert merged["template"]["path"] == os.path.abspath(template)
    assert merged["ok"] is False
    assert {risk["source"] for risk in merged["risks"]} == {"template"}
    assert "SOURCE_TRACKED_CHANGES" in {
        risk["code"] for risk in merged["blockers"]
        if risk["source"] == "template"
    }


def test_cli_preflight_only_scans_template_and_writes_report(tmp_path, monkeypatch):
    target = tmp_path / "target.docx"
    template = tmp_path / "template.docx"
    output = tmp_path / "planned.docx"
    plain = Document()
    plain.add_paragraph("Plain target content")
    plain.save(target)
    _write_story_fixture(template)

    import main as cli

    monkeypatch.setattr(sys, "argv", [
        "main.py", "--preflight-only",
        "--target", str(target),
        "--template", str(template),
        "--out", str(output),
    ])
    assert cli.main() == 2
    report_path = tmp_path / "planned_preflight.json"
    report = json.loads(report_path.read_text(encoding="utf-8"))
    assert report["template"]["path"] == os.path.abspath(template)
    assert any(
        risk["source"] == "template"
        and risk["code"] == "SOURCE_TRACKED_CHANGES"
        for risk in report["blockers"]
    )


def test_multi_section_page_geometry_is_preserved(tmp_path):
    source = tmp_path / "multi_section.docx"
    output = tmp_path / "multi_section_out.docx"
    doc = Document()
    doc.add_paragraph("First section paragraph")
    second = doc.add_section(WD_SECTION.NEW_PAGE)
    second.orientation = WD_ORIENT.LANDSCAPE
    second.page_width = Inches(11)
    second.page_height = Inches(8.5)
    doc.sections[0].left_margin = Mm(22)
    second.left_margin = Mm(16)
    doc.sections[0].header.paragraphs[0].text = "First-section header"
    second.header.is_linked_to_previous = False
    second.header.paragraphs[0].text = "Second-section header"
    doc.add_paragraph("Second section paragraph")
    doc.save(source)

    before = Document(source)
    before_geometry = [
        (section.start_type, section.orientation,
         section.page_width, section.page_height)
        for section in before.sections
    ]
    before_margins = [
        (section.top_margin, section.bottom_margin,
         section.left_margin, section.right_margin)
        for section in before.sections
    ]
    rolemap = {idx: "body" for idx, _ in enumerate(before.paragraphs)}
    spec = _spec()
    spec["page"]["header"] = {
        "text": "Template header",
        "font_ascii": "Arial",
        "size_pt": 10,
        "alignment": "center",
    }
    apply_format(source, spec, rolemap, output)

    result = Document(output)
    after_geometry = [
        (section.start_type, section.orientation,
         section.page_width, section.page_height)
        for section in result.sections
    ]
    assert after_geometry == before_geometry
    assert len(result.sections) == 2
    assert [
        (section.top_margin, section.bottom_margin,
         section.left_margin, section.right_margin)
        for section in result.sections
    ] == before_margins
    assert result.sections[0].header.paragraphs[0].text == "Template header"
    assert result.sections[1].header.paragraphs[0].text == "Second-section header"
    assert result.sections[1].header.is_linked_to_previous is False
    assert all(
        section._sectPr.find(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}docGrid"
        ) is not None
        for section in result.sections
    )


def test_multi_section_structural_rebuild_requires_override(tmp_path):
    source = tmp_path / "multi_section.docx"
    doc = Document()
    doc.add_paragraph("Section one")
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_paragraph("Section two")
    doc.save(source)

    risky = _spec("thesis")
    risky["structure"] = {"enabled": True, "mode": "thesis"}
    report = preflight_docx(source, spec=risky)
    assert "MULTI_SECTION_STRUCTURAL_REBUILD" in {
        item["code"] for item in report["blockers"]
    }
    approved = preflight_docx(
        source, spec=risky, allow_risky_structure=True)
    assert "MULTI_SECTION_STRUCTURAL_REBUILD" in {
        item["code"] for item in approved["warnings"]
    }


@pytest.mark.parametrize(
    ("paragraphs", "expected"),
    [
        ([
            "Quarterly Operations Review",
            "This document summarizes delivery, staffing, and open decisions for the next quarter.",
        ], "english_general"),
        ([
            "A Reproducible Evaluation of Document Formatting Agents",
            "Abstract",
            "This paper evaluates deterministic formatting pipelines across representative documents.",
            "Keywords: document automation; reproducibility",
            "Introduction",
            "Methods",
            "References",
        ], "english_academic"),
        ([
            "SERVICES AGREEMENT",
            "WHEREAS the parties wish to record the terms of their agreement;",
            "ARTICLE I DEFINITIONS",
            "SECTION 1.1 Defined Terms",
            "IN WITNESS WHEREOF, the parties have executed this Agreement.",
        ], "english_legal"),
    ],
)
def test_english_profile_detection(paragraphs, expected):
    records = [{"text": text} for text in paragraphs]
    assert detect_document_profile(records) == expected
    validate_spec(_spec(expected))


def test_english_academic_and_legal_roles():
    assert regex_role("Abstract", profile="english_academic") == "abstract_heading"
    assert regex_role(
        "Abstract: This study evaluates a deterministic pipeline.",
        profile="english_academic",
    ) == "abstract_body"
    assert regex_role(
        "Keywords: formatting; document automation",
        profile="english_academic",
    ) == "keywords"
    assert regex_role("References", profile="english_academic") == "bibliography_heading"
    assert regex_role("Chapter 2 Related Work", profile="english_academic") == "chapter_heading"
    assert regex_role("2.1 Experimental Setup", profile="english_academic") == "heading_2"
    assert regex_role("Figure 2. System architecture", profile="english_academic") == "figure_caption"
    assert regex_role("Table 3. Evaluation results", profile="english_academic") == "table_caption"
    assert regex_role("Equation 4. Governing equation", profile="english_academic") == "equation"
    assert regex_role("Author Note", profile="english_academic") == "author_note"
    assert regex_role("Corresponding author: a@example.org", profile="english_academic") == "correspondence"
    assert regex_role("A. Scope", profile="english_academic") == "heading_2"
    assert regex_role("(a) Definitions", profile="english_legal") == "heading_3"
    assert regex_role("ARTICLE IV REMEDIES", profile="english_legal") == "heading_1"
    assert regex_role("Section 4.2 Termination", profile="english_legal") == "heading_2"
    assert regex_role("Table of Authorities", profile="english_legal") == "table_of_authorities_heading"
    assert regex_role('"Effective Date" means 1 January 2026.', profile="english_legal") == "legal_definition"
    assert regex_role("IN WITNESS WHEREOF", profile="english_legal") == "signature_block"
    assert regex_role("Exhibit A", profile="english_legal") == "attachment_label"


def test_english_academic_template_builds_semantic_roles(tmp_path):
    template = tmp_path / "academic_template.docx"
    doc = Document()
    doc.add_paragraph("A Reproducible Study of Formatting Agents")
    doc.add_paragraph("Abstract")
    doc.add_paragraph(
        "This paper presents a deterministic method for formatting long documents."
    )
    doc.add_paragraph("Keywords: formatting; reproducibility")
    introduction = doc.add_paragraph("Introduction")
    introduction.runs[0].font.size = Pt(14)
    introduction.runs[0].bold = True
    doc.add_paragraph("References")
    doc.add_paragraph("Doe, J. (2025). A representative reference entry.")
    doc.save(template)

    rolemap = {
        0: "title",
        1: "heading_1",
        2: "body",
        3: "body",
        4: "heading_1",
        5: "heading_1",
        6: "body",
    }
    spec = extract_rules_from_template(template, rolemap)
    validate_spec(spec)
    assert spec["profile"] == "english_academic"
    assert spec["cleanup"]["mode"] == "preserve_emphasis"
    assert spec["roles"]["chapter_heading"]["size_pt"] == 14
    assert spec["roles"]["chapter_heading"]["bold"] is True
    assert {
        "abstract_heading", "abstract_body", "keywords",
        "chapter_heading", "bibliography_heading", "bibliography_entry",
        "equation", "appendix_heading",
    } <= set(spec["roles"])
    assert spec["roles"]["chapter_heading"].get("page_break_before") is not True
    assert spec["roles"]["bibliography_heading"].get("page_break_before") is not True


def test_integrity_failure_does_not_publish_any_candidate(tmp_path, monkeypatch):
    source = tmp_path / "source.docx"
    output = tmp_path / "result.docx"
    report = tmp_path / "result_report.md"
    report_docx = tmp_path / "result_report.docx"
    tracked = tmp_path / "result_tracked.docx"

    source_doc = Document()
    source_doc.add_paragraph("Source content must not change.")
    source_doc.save(source)
    for path, marker in ((output, "OLD MAIN"), (report_docx, "OLD REPORT"), (tracked, "OLD TRACKED")):
        existing = Document()
        existing.add_paragraph(marker)
        existing.save(path)
    report.write_text("OLD MARKDOWN REPORT", encoding="utf-8")
    before = {
        path: path.read_bytes()
        for path in (output, report, report_docx, tracked)
    }

    import core.text_integrity as text_integrity

    monkeypatch.setattr(
        text_integrity,
        "check_text_integrity",
        lambda *args, **kwargs: {
            "ok": False,
            "added": ["unexpected"],
            "removed": ["Source content must not change."],
        },
    )
    with pytest.raises(IntegrityViolationError):
        Agent().run(
            source,
            output,
            spec=_spec(),
            rolemap={0: "body"},
            report_path=report,
        )

    assert all(path.read_bytes() == content for path, content in before.items())
    assert not list(tmp_path.glob(".*.tmp*"))


def test_preflight_blocker_keeps_existing_output_unchanged(tmp_path):
    source = tmp_path / "tracked_source.docx"
    _write_story_fixture(source)
    output = tmp_path / "result.docx"
    existing = Document()
    existing.add_paragraph("OLD FINAL")
    existing.save(output)
    before = output.read_bytes()

    with pytest.raises(PreflightBlockedError):
        Agent().run(
            source,
            output,
            spec=_spec(),
            rolemap={idx: "body" for idx in range(3)},
        )

    assert output.read_bytes() == before
    assert not (tmp_path / "result_report.md").exists()
    assert not (tmp_path / "result_report.docx").exists()
    assert not (tmp_path / "result_tracked.docx").exists()


def test_success_commits_complete_output_set(tmp_path):
    source = tmp_path / "source.docx"
    output = tmp_path / "result.docx"
    doc = Document()
    doc.add_paragraph("A complete output set is committed only after validation.")
    doc.save(source)

    result = Agent().run(
        source,
        output,
        spec=_spec("english_general"),
        rolemap={0: "body"},
    )

    assert output.exists()
    assert (tmp_path / "result_report.md").exists()
    assert (tmp_path / "result_report.docx").exists()
    assert (tmp_path / "result_tracked.docx").exists()
    assert result["preflight"]["ok"] is True
    assert Document(output).paragraphs[0].text.startswith("A complete output set")
    assert not list(tmp_path.glob(".*.tmp*"))


def test_source_path_cannot_be_used_as_output(tmp_path):
    source = tmp_path / "same.docx"
    doc = Document()
    doc.add_paragraph("Do not overwrite this source.")
    doc.save(source)
    with pytest.raises(UnsafeOutputPathError):
        Agent().run(
            source,
            source,
            spec=_spec(),
            rolemap={0: "body"},
        )
