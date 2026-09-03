# -*- coding: utf-8 -*-
"""编号正文识别回归测试。

覆盖：真自动编号、手工 1./2./1.2、numId=0/ilvl<0 取消编号、
连续正文列表、带强标题元数据的短小数字标题，以及模板 body 代表段
不得把偶然的列表编号抽成全局 body.numbering。
"""

import os
import sys
import tempfile

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from core.extract import extract_paragraphs, manual_number_prefix
from core.label_roles import label_roles, regex_role
from core.rules_from_template import extract_rules_from_template


def _child(parent, tag, value=None):
    element = OxmlElement(tag)
    if value is not None:
        element.set(qn("w:val"), str(value))
    parent.append(element)
    return element


def _add_numbering(doc, num_format="decimal", level_text="%1.", left=720, hanging=360):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(el.get(qn("w:abstractNumId")))
        for el in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(el.get(qn("w:numId")))
        for el in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max([0] + num_ids) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    _child(abstract, "w:multiLevelType", "singleLevel")
    lvl = _child(abstract, "w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    _child(lvl, "w:start", 1)
    _child(lvl, "w:numFmt", num_format)
    _child(lvl, "w:lvlText", level_text)
    ppr = _child(lvl, "w:pPr")
    ind = _child(ppr, "w:ind")
    ind.set(qn("w:start"), str(left))  # 验证 start/left 两种 OOXML 写法均能读
    ind.set(qn("w:hanging"), str(hanging))

    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(abstract)
    else:
        numbering.insert(numbering.index(first_num), abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    _child(num, "w:abstractNumId", abstract_id)
    numbering.append(num)
    return num_id


def _set_num_pr(paragraph, num_id, level=0):
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    _child(num_pr, "w:ilvl", level)
    _child(num_pr, "w:numId", num_id)
    ppr._insert_numPr(num_pr)


def _set_direct_indent(paragraph, first_line=640, first_line_chars=200):
    ind = paragraph._p.get_or_add_pPr().get_or_add_ind()
    ind.set(qn("w:firstLine"), str(first_line))
    ind.set(qn("w:firstLineChars"), str(first_line_chars))


assert manual_number_prefix("1. 同意该议案")["parts"] == [1]
assert manual_number_prefix("2、同意该议案")["parts"] == [2]
assert manual_number_prefix("1.2 项目范围")["parts"] == [1, 2]
assert manual_number_prefix("1.2.1 实施范围")["parts"] == [1, 2, 1]
assert manual_number_prefix("投资1.2万元") is None
assert manual_number_prefix("1.2万元投资") is None
assert regex_role("1. 经济效益") is None  # 数字前缀本身不足以强制标题


with tempfile.TemporaryDirectory() as td:
    source = os.path.join(td, "numbered-source.docx")
    doc = Document()
    manual_1 = doc.add_paragraph(
        "1. 同意向科创基金委派常任委员，并由相关专家完成后续表决工作。")
    manual_2 = doc.add_paragraph(
        "2. 同意常任委员和专家委员对投资事项做出同意表决。")
    for paragraph in (manual_1, manual_2):
        _set_num_pr(paragraph, 0, -1)
        _set_direct_indent(paragraph)

    auto_num_id = _add_numbering(doc, left=720, hanging=360)
    auto_1 = doc.add_paragraph("自动编号正文第一项。")
    auto_2 = doc.add_paragraph("自动编号正文第二项。")
    _set_num_pr(auto_1, auto_num_id)
    _set_num_pr(auto_2, auto_num_id)

    heading_style = doc.styles.add_style("标题 3", WD_STYLE_TYPE.PARAGRAPH)
    outline = heading_style.element.get_or_add_pPr().get_or_add_outlineLvl()
    outline.set(qn("w:val"), "2")
    doc.add_paragraph("1.2 项目范围", style=heading_style)
    doc.add_paragraph("1. 短正文", style="Normal")
    doc.save(source)

    records = extract_paragraphs(source)
    for idx in (0, 1):
        assert records[idx]["numbering_status"] == "cancelled"
        assert records[idx]["list_kind"] == "manual"
        assert records[idx]["list_sequence"] is True
        assert records[idx]["indent_first_line_twips"] == 640
        assert records[idx]["indent_first_line_chars"] == 200
    for idx in (2, 3):
        assert records[idx]["numbering_status"] == "automatic"
        assert records[idx]["list_kind"] == "automatic"
        assert records[idx]["list_sequence"] is True
        assert records[idx]["num_format"] == "decimal"
        assert records[idx]["level_text"] == "%1."
        assert records[idx]["numbering_left_twips"] == 720
        assert records[idx]["numbering_hanging_twips"] == 360

    assert regex_role(records[0]["text"], records[0]) == "body"
    assert regex_role(records[2]["text"], records[2]) == "body"
    assert regex_role(records[4]["text"], records[4]) == "heading_3"
    assert regex_role(records[5]["text"], records[5]) == "body"

    class MustNotRunLLM:
        def chat_json(self, _prompt):
            raise AssertionError("明确的编号正文/标题应由结构元数据确定，不应调用 LLM")

    rolemap = label_roles(records, MustNotRunLLM())
    assert [rolemap[idx] for idx in range(6)] == [
        "body", "body", "body", "body", "heading_3", "body"]

    # 不确定的数字短段交给 LLM，但 prompt 必须携带编号/样式/大纲元数据。
    ambiguous = dict(records[5])
    ambiguous.update({
        "idx": 99, "style_name": "Mystery", "outline_level": None,
        "numbering_status": "none", "list_sequence": False,
    })

    class CaptureLLM:
        prompt = None

        def chat_json(self, prompt):
            self.prompt = prompt
            return {"roles": [{"idx": 99, "role": "body"}]}

    llm = CaptureLLM()
    assert label_roles([ambiguous], llm) == {99: "body"}
    for field in ("style_name", "outline_level", "numbering_status", "list_kind",
                  "num_id", "indent_first_line_twips"):
        assert f'"{field}"' in llm.prompt

    # 模板的第一个 body 是自动列表项：必须改选普通正文作代表，
    # 且 body 规则不得带 numbering；标题编号及其缩进仍正常抽取。
    template = os.path.join(td, "template.docx")
    tpl = Document()
    body_num_id = _add_numbering(tpl, left=900, hanging=450)
    body_list = tpl.add_paragraph("模板中的编号正文项。")
    body_list.runs[0].font.size = Pt(9)
    _set_num_pr(body_list, body_num_id)
    body_style = tpl.styles.add_style("模板正文", WD_STYLE_TYPE.PARAGRAPH)
    body_ppr = body_style.element.get_or_add_pPr()
    body_ind = _child(body_ppr, "w:ind")
    body_ind.set(qn("w:firstLineChars"), "200")
    body_ind.set(qn("w:firstLine"), "640")
    body_spacing = _child(body_ppr, "w:spacing")
    body_spacing.set(qn("w:line"), "600")
    body_spacing.set(qn("w:lineRule"), "exact")
    _child(body_ppr, "w:jc", "both")
    plain_body = tpl.add_paragraph("模板中的普通正文代表段。", style=body_style)
    plain_body.runs[0].font.size = Pt(16)

    heading_num_id = _add_numbering(
        tpl, num_format="chineseCounting", level_text="%1、", left=720, hanging=360)
    heading = tpl.add_paragraph("标题文本")
    heading.runs[0].font.size = Pt(18)
    _set_num_pr(heading, heading_num_id)
    tpl.save(template)

    spec = extract_rules_from_template(template, {0: "body", 1: "body", 2: "heading_1"})
    assert spec["roles"]["body"]["size_pt"] == 16
    assert spec["roles"]["body"]["alignment"] == "justify"
    assert spec["roles"]["body"]["first_line_indent_chars"] == 2
    assert spec["roles"]["body"]["line_spacing"] == {"type": "exact", "pt": 30.0}
    assert "numbering" not in spec["roles"]["body"]
    heading_numbering = spec["roles"]["heading_1"]["numbering"]
    assert heading_numbering["num_format"] == "chineseCounting"
    assert heading_numbering["level_text"] == "%1、"
    assert heading_numbering["left_twips"] == 720
    assert heading_numbering["hanging_twips"] == 360

print("编号正文识别测试通过：自动/手工/取消编号、序列、缩进与 body 模板防污染均正常")
