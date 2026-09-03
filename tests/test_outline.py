# -*- coding: utf-8 -*-
"""确定性测试：二级标题 regex 预识别 + 大纲级别写入 + spec_std 校验。
不调用 LLM，直接跑。运行: python tests/test_outline.py
"""
import json
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from docx.oxml.ns import qn

from core.apply import apply_format
from core.extract import extract_paragraphs
from core.label_roles import regex_role
from core.schema import validate_spec

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

# ---- 1. regex_role 单测 ----
cases = {
    "一、背景": "heading_1",
    "二、主要内容": "heading_1",
    "（一）投资背景": "heading_2",
    "(二)投资方案": "heading_2",
    "（三）回报预测": "heading_2",
    "1. 经济效益": None,  # 数字前缀本身不足以强制标题
    "1、直接经济效益": None,
    "图1 系统架构图": "figure_caption",
    "图 2-1 技术路线图": "figure_caption",
    "表1 主要财务数据": "table_caption",
    "表 3-2 对比分析表": "table_caption",
    "为深入贯彻落实上级决策部署，现就有关工作通知如下。": None,  # 正文
    "一、项目背景的基本情况如下。": None,  # 以句号结尾，按正文处理
    "": None,
}
for text, want in cases.items():
    got = regex_role(text)
    assert got == want, "regex_role(%r) = %r, want %r" % (text, got, want)
print("1. regex_role 单测通过（%d 例）" % len(cases))

assert regex_role("1. 经济效益", {
    "manual_number": "1", "list_kind": "manual", "char_count": 7,
    "ends_with_sentence_punct": False, "outline_level": 2, "style_name": "标题 3",
}) == "heading_3"
assert regex_role("1. 同意该议案", {
    "manual_number": "1", "list_kind": "manual", "char_count": 8,
    "ends_with_sentence_punct": False, "outline_level": None, "style_name": "Normal",
}) == "body"

# ---- 2. spec_std 通过校验（含新增 heading_2/heading_3）----
with open(os.path.join(ROOT, "assets", "spec_std.json"), encoding="utf-8") as f:
    spec = json.load(f)
validate_spec(spec)
with open(os.path.join(ROOT, "assets", "spec_thesis_std.json"), encoding="utf-8") as f:
    thesis_spec = json.load(f)
validate_spec(thesis_spec)
print("2. spec_std.json / spec_thesis_std.json 校验通过")

# ---- 3. 端到端：构造含二级标题的文档，apply 后验证大纲级别与字体 ----
from docx import Document as NewDoc

src = os.path.join(ROOT, "out", "_outline_src.docx")
out = os.path.join(ROOT, "out", "outline_test.docx")
os.makedirs(os.path.dirname(src), exist_ok=True)
d = NewDoc()
d.add_paragraph("关于向投资基金合伙企业委派委员的议案")   # 0 title
d.add_paragraph("一、背景")                                 # 1 heading_1
d.add_paragraph("为贯彻落实发展规划，推动产业整体发展，特设立本基金。")  # 2 body
d.add_paragraph("（一）投资背景")                           # 3 heading_2
d.add_paragraph("在基金层面，作为政策性基金，科创基金完成了多笔投资。")  # 4 body
d.add_paragraph("（二）投资方案")                           # 5 heading_2
d.add_paragraph("星驰智行已于2024年6月与全体股东终止了全部特殊权利。")   # 6 body
d.add_paragraph("某某局办公室")                              # 7 signature
d.add_paragraph("2026年8月28日")                             # 8 date
d.save(src)

rolemap = {0: "title", 1: "heading_1", 2: "body", 3: "heading_2",
           4: "body", 5: "heading_2", 6: "body", 7: "signature", 8: "date"}
changelog = apply_format(src, spec, rolemap, out)

doc = Document(out)
def _outline_val(p):
    ppr = p._p.find(qn("w:pPr"))
    if ppr is not None:
        ol = ppr.find(qn("w:outlineLvl"))
        if ol is not None:
            return ol.get(qn("w:val"))
    # 新流水线把大纲级别定义在命名样式中，而不是逐段直刷。
    style_ppr = p.style.element.find(qn("w:pPr"))
    if style_ppr is None:
        return None
    ol = style_ppr.find(qn("w:outlineLvl"))
    return ol.get(qn("w:val")) if ol is not None else None

n_h2 = 0
for idx, p in enumerate(doc.paragraphs):
    role = rolemap.get(idx)
    if role == "title":
        assert _outline_val(p) == "0", "title 大纲级别应为 0"
    elif role == "heading_1":
        assert _outline_val(p) == "0", "heading_1 大纲级别应为 0"
    elif role == "heading_2":
        n_h2 += 1
        assert _outline_val(p) == "1", "heading_2 大纲级别应为 1"
        assert p.style.name == "标题 2", "heading_2 必须绑定真实中文命名样式"
        rpr = p.style.element.find(qn("w:rPr"))
        font = rpr.find(qn("w:rFonts")).get(qn("w:eastAsia"))
        assert font == "楷体_GB2312", "heading_2 应为楷体_GB2312，实际 %s" % font
    elif role == "body":
        assert _outline_val(p) is None, "body 不应有大纲级别"
print("3. 端到端 apply 通过：识别二级标题 %d 个，大纲级别 XML 全部正确" % n_h2)

# ---- 4. 渲染验证文件完整性（手工脚本模式才调用 Office/LibreOffice）----
# pytest 收集模块时不能拉起 Word COM；渲染属于较慢的环境集成测试。
if __name__ == "__main__":
    from core.render import render_docx_to_png
    pages = render_docx_to_png(out, os.path.join(ROOT, "out", "outline_render"))
    print("4. 渲染通过：%d 页 -> out/outline_render/" % len(pages))
    print("ALL OK")
