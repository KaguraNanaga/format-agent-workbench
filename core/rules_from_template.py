# 模板 docx → FormatSpec（格式源第二种，确定性读取，不依赖 VLM）。
# 思路：给定模板的 RoleMap，为每个角色找代表段落，用 effective_props 读生效
# 字体/字号/加粗，用 python-docx 读对齐/行距/缩进，页面级读 section 页边距和行网格。
# 未知/读不到的字段不编——留给 LLM 规范抽取或人肉 JSON 补。

from copy import deepcopy
import json
import re

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn

from core.effective_props import effective_props
from core.extract import manual_number_prefix, paragraph_numbering_metadata
from core.profiles import detect_profile_from_texts
from core.schema import validate_spec

_ALIGN_MAP = {0: "left", 1: "center", 2: "right", 3: "justify"}
_ALIGN_XML_MAP = {
    "left": "left",
    "start": "left",
    "center": "center",
    "right": "right",
    "end": "right",
    "both": "justify",
    "distribute": "justify",
    "thaiDistribute": "justify",
}


def _effective_ppr_elements(paragraph):
    """按“段落直接格式 → 段落样式链”顺序产出 pPr。"""
    direct = paragraph._p.pPr
    if direct is not None:
        yield direct
    style = paragraph.style
    seen = set()
    while style is not None and style.style_id not in seen:
        seen.add(style.style_id)
        ppr = style.element.find(qn("w:pPr"))
        if ppr is not None:
            yield ppr
        style = style.base_style


def _effective_ppr_child(paragraph, tag):
    for ppr in _effective_ppr_elements(paragraph):
        child = ppr.find(qn(tag))
        if child is not None:
            return child
    return None


def _effective_ppr_attr(paragraph, child_tag, *attrs):
    """逐层、逐属性读取有效值，允许直接格式只覆盖同一元素的部分属性。"""
    for ppr in _effective_ppr_elements(paragraph):
        child = ppr.find(qn(child_tag))
        if child is None:
            continue
        for attr in attrs:
            value = child.get(qn(attr))
            if value is not None:
                return value
    return None


def _numbering_element_for_paragraph(p):
    """返回段落直接或样式链继承的 numPr。"""
    ppr = p._p.pPr
    if ppr is not None:
        num_pr = ppr.find(qn("w:numPr"))
        if num_pr is not None:
            return num_pr
    style = p.style
    seen = set()
    while style is not None and style.style_id not in seen:
        seen.add(style.style_id)
        ppr = style.element.find(qn("w:pPr"))
        if ppr is not None:
            num_pr = ppr.find(qn("w:numPr"))
            if num_pr is not None:
                return num_pr
        style = style.base_style
    return None


def _find_by_attr(parent, tag, attr, value):
    for element in parent.findall(qn(tag)):
        if element.get(qn(attr)) == str(value):
            return element
    return None


def _val(parent, tag, default=None):
    element = parent.find(qn(tag)) if parent is not None else None
    return element.get(qn("w:val")) if element is not None else default


def _int_or_none(value):
    try:
        return int(value)
    except (TypeError, ValueError):
        return None


def _numbering_rule(doc, paragraph, role):
    """解析段落的有效 OOXML 编号级别，转成可跨文档复制的规则。"""
    # 正文的自动列表是局部段落结构，不是“所有正文”的样式规则。
    # 若把偶然命中的第一个 body 列表项抽成 body.numbering，会使全文正文
    # 都被编号。当前只允许语义标题角色产生全局样式编号。
    if not role.startswith("heading_"):
        return None
    num_pr = _numbering_element_for_paragraph(paragraph)
    if num_pr is None:
        return None
    num_id = _int_or_none(_val(num_pr, "w:numId"))
    level = _int_or_none(_val(num_pr, "w:ilvl", "0"))
    if num_id is None or num_id <= 0 or level is None:
        return None

    numbering = doc.part.numbering_part.element
    num = _find_by_attr(numbering, "w:num", "w:numId", num_id)
    if num is None:
        return None
    abstract_id = _int_or_none(_val(num, "w:abstractNumId"))
    abstract = _find_by_attr(
        numbering, "w:abstractNum", "w:abstractNumId", abstract_id)
    if abstract is None:
        return None

    lvl = None
    override = _find_by_attr(num, "w:lvlOverride", "w:ilvl", level)
    if override is not None:
        lvl = override.find(qn("w:lvl"))
    if lvl is None:
        lvl = _find_by_attr(abstract, "w:lvl", "w:ilvl", level)
    if lvl is None:
        return None

    num_format = _val(lvl, "w:numFmt")
    level_text = _val(lvl, "w:lvlText")
    if not num_format or level_text is None:
        return None
    override_start = _int_or_none(_val(override, "w:startOverride")) if override is not None else None
    rule = {
        "group": "headings" if role.startswith("heading_") else role,
        "level": level,
        "num_format": num_format,
        "level_text": level_text,
        "start": override_start or _int_or_none(_val(lvl, "w:start", "1")) or 1,
        "suffix": _val(lvl, "w:suff", "tab"),
        "alignment": _val(lvl, "w:lvlJc", "left"),
    }
    if lvl.find(qn("w:isLgl")) is not None:
        rule["is_legal"] = True
    level_restart = _int_or_none(_val(lvl, "w:lvlRestart"))
    if level_restart is not None:
        rule["level_restart"] = level_restart

    ppr = lvl.find(qn("w:pPr"))
    ind = ppr.find(qn("w:ind")) if ppr is not None else None
    if ind is not None:
        left = ind.get(qn("w:left")) or ind.get(qn("w:start"))
        for value, key in (
            (left, "left_twips"),
            (ind.get(qn("w:hanging")), "hanging_twips"),
            (ind.get(qn("w:firstLine")), "first_line_twips"),
        ):
            parsed = _int_or_none(value)
            if parsed is not None:
                rule[key] = parsed
    tabs = ppr.find(qn("w:tabs")) if ppr is not None else None
    if tabs is not None:
        for tab in tabs.findall(qn("w:tab")):
            if tab.get(qn("w:val")) in {"num", "left"}:
                value = _int_or_none(tab.get(qn("w:pos")))
                if value is not None:
                    rule["tab_pos_twips"] = value
                    break
    rpr = lvl.find(qn("w:rPr"))
    if rpr is not None:
        fonts = rpr.find(qn("w:rFonts"))
        if fonts is not None:
            eastasia = fonts.get(qn("w:eastAsia"))
            ascii_font = fonts.get(qn("w:ascii")) or fonts.get(qn("w:hAnsi"))
            if eastasia:
                rule["font_eastasia"] = eastasia
            if ascii_font:
                rule["font_ascii"] = ascii_font
        size = _int_or_none(_val(rpr, "w:sz"))
        if size is not None:
            rule["size_pt"] = size / 2
        bold = rpr.find(qn("w:b"))
        if bold is not None:
            rule["bold"] = bold.get(qn("w:val"), "1") not in {"0", "false", "off"}
    return rule


def _is_numbered_body_candidate(paragraph):
    """用于选正文代表段：优先避开真自动列表和手工 1./1.2 列表。"""
    metadata = paragraph_numbering_metadata(paragraph)
    return bool(
        metadata.get("numbering_status") == "automatic"
        or manual_number_prefix(paragraph.text) is not None
    )


def _representative_paragraphs(paragraphs, rolemap):
    """按角色选主流格式代表段，避免首段或短标签 Run 污染全局规则。"""
    from collections import Counter

    candidates = {}
    for idx, role in sorted(rolemap.items()):
        if 0 <= idx < len(paragraphs):
            candidates.setdefault(role, []).append(paragraphs[idx])
    representatives = {}
    for role, role_paragraphs in candidates.items():
        if role == "body":
            role_paragraphs = [
                paragraph for paragraph in role_paragraphs
                if paragraph.text.strip() and not _is_numbered_body_candidate(paragraph)
            ] or role_paragraphs
        signatures = []
        for paragraph in role_paragraphs:
            props = effective_props(paragraph)
            signatures.append((
                paragraph.style.style_id if paragraph.style is not None else None,
                props.get("eastasia"), props.get("ascii"), props.get("cs"),
                props.get("size_pt"), props.get("bold"), props.get("italic"),
                _para_alignment(paragraph),
            ))
        dominant = Counter(signatures).most_common(1)[0][0]
        matching = [
            paragraph for paragraph, signature in zip(role_paragraphs, signatures)
            if signature == dominant
        ]
        representatives[role] = max(
            matching, key=lambda paragraph: len(paragraph.text.strip()))
    return representatives


def _para_alignment(p):
    jc = _effective_ppr_child(p, "w:jc")
    if jc is not None:
        value = jc.get(qn("w:val"))
        if value in _ALIGN_XML_MAP:
            return _ALIGN_XML_MAP[value]
    a = p.alignment
    return _ALIGN_MAP.get(int(a)) if a is not None else None


def _para_line_spacing(p):
    spacing = _effective_ppr_child(p, "w:spacing")
    if spacing is not None and spacing.get(qn("w:line")) is not None:
        line = _int_or_none(spacing.get(qn("w:line")))
        line_rule = spacing.get(qn("w:lineRule"), "auto")
        if line is not None and line_rule == "exact":
            return {"type": "exact", "pt": round(line / 20, 1)}
        if line is not None and line_rule == "auto":
            return {"type": "multiple", "pt": round(line / 240, 2)}
    pf = p.paragraph_format
    rule = pf.line_spacing_rule
    if rule == WD_LINE_SPACING.EXACTLY and pf.line_spacing is not None:
        return {"type": "exact", "pt": round(pf.line_spacing.pt, 1)}
    if rule == WD_LINE_SPACING.MULTIPLE and pf.line_spacing is not None:
        return {"type": "multiple", "pt": round(float(pf.line_spacing), 2)}
    return None


def _para_indent_chars(p, size_pt):
    """首行缩进字符数：优先读 XML firstLineChars，否则用磅值/字号反推。"""
    # hanging 与 firstLine 互斥；较近一层明确设置悬挂时，不能继续继承
    # 样式中的首行缩进。
    for ppr in _effective_ppr_elements(p):
        ind = ppr.find(qn("w:ind"))
        if ind is None:
            continue
        if ind.get(qn("w:hanging")) is not None or ind.get(qn("w:hangingChars")) is not None:
            return None
        flc = ind.get(qn("w:firstLineChars"))
        if flc is not None:
            return round(int(flc) / 100, 1)
        first_line = ind.get(qn("w:firstLine"))
        if first_line is not None and size_pt:
            return round((int(first_line) / 20) / size_pt, 1)
    fl = p.paragraph_format.first_line_indent
    if fl is not None and size_pt:
        return round(fl.pt / size_pt, 1)
    return None


def _para_spacing_pt(p, attr):
    value = _effective_ppr_attr(p, "w:spacing", attr)
    parsed = _int_or_none(value)
    return round(parsed / 20, 1) if parsed is not None else None


def _ppr_toggle(p, tag):
    element = _effective_ppr_child(p, tag)
    if element is None:
        return None
    return element.get(qn("w:val"), "1") not in {"0", "false", "off"}


def _paragraph_rule(doc, p, role):
    props = effective_props(p)
    size_pt = props.get("size_pt") or 10.5
    rule = {
        "font_eastasia": props.get("eastasia") or "宋体",
        "size_pt": size_pt,
        "bold": bool(props.get("bold")),
    }
    if props.get("ascii"):
        rule["font_ascii"] = props["ascii"]
    if props.get("cs"):
        rule["font_cs"] = props["cs"]
    if props.get("language"):
        rule["language"] = props["language"]
    for field in ("italic", "underline", "strike", "caps", "small_caps", "rtl"):
        if props.get(field) is not None:
            rule[field] = bool(props[field])
    bidi = _ppr_toggle(p, "w:bidi")
    if bidi is not None:
        rule["bidi"] = bidi
    for field in ("color", "highlight"):
        if props.get(field) is not None:
            rule[field] = props[field]
    alignment = _para_alignment(p)
    if alignment:
        rule["alignment"] = alignment
    line_spacing = _para_line_spacing(p)
    if line_spacing:
        rule["line_spacing"] = line_spacing
    first_indent = _para_indent_chars(p, size_pt)
    if first_indent:
        rule["first_line_indent_chars"] = first_indent
    for field, attr in (("space_before_pt", "w:before"), ("space_after_pt", "w:after")):
        value = _para_spacing_pt(p, attr)
        if value is not None:
            rule[field] = value
    for field, tag in (
        ("keep_with_next", "w:keepNext"),
        ("keep_together", "w:keepLines"),
        ("page_break_before", "w:pageBreakBefore"),
        ("widow_control", "w:widowControl"),
    ):
        value = _ppr_toggle(p, tag)
        if value is not None:
            rule[field] = value
    numbering = _numbering_rule(doc, p, role)
    if numbering is not None:
        rule["numbering"] = numbering
    rule.setdefault("alignment", "justify" if role in {"body", "abstract_body"} else "left")
    return rule


def _detect_profile(doc):
    return detect_profile_from_texts(
        p.text.strip() for p in doc.paragraphs if p.text.strip())


def _cleanup_mode_for_template(doc, profile):
    if profile in {
        "english_academic", "english_legal", "english_technical",
        "english_legal_brief",
    }:
        return "preserve_emphasis"
    if profile == "thesis":
        text = " ".join(paragraph.text for paragraph in doc.paragraphs)
        latin = len(re.findall(r"[A-Za-z]", text))
        cjk = len(re.findall(r"[\u3400-\u9fff]", text))
        if latin >= 50 and latin >= max(1, round(cjk * 0.1)):
            return "preserve_emphasis"
        return "strict"
    return "controlled"


def _section_page_numbering(section):
    element = section._sectPr.find(qn("w:pgNumType"))
    if element is None:
        return {}
    result = {}
    if element.get(qn("w:fmt")):
        result["format"] = element.get(qn("w:fmt"))
    if element.get(qn("w:start")):
        result["start"] = _int_or_none(element.get(qn("w:start")))
    return result


def _thesis_structure(doc):
    """从论文模板提取可安全迁移的结构契约，不复制示例作者或声明内容。"""
    institution = next((
        paragraph.text.strip()
        for paragraph in doc.paragraphs
        if paragraph.text.strip()
        and str(paragraph.style.name if paragraph.style is not None else "") == "主题"
    ), "")
    if not institution:
        for section in doc.sections:
            text = next((
                paragraph.text.strip() for paragraph in section.header.paragraphs
                if paragraph.text.strip()
            ), "")
            if text:
                institution = text.split("\t", 1)[0].strip()
                break
    institution = institution or "学位论文"

    front_numbering = {"format": "upperRoman", "start": 1}
    body_numbering = {"format": "decimal", "start": 1}
    for section in doc.sections:
        page_numbering = _section_page_numbering(section)
        if page_numbering.get("format") in {"upperRoman", "lowerRoman"}:
            front_numbering.update(page_numbering)
        if page_numbering.get("start") == 1 and page_numbering.get("format") not in {
            "upperRoman", "lowerRoman", "upperLetter", "lowerLetter",
        }:
            body_numbering.update(page_numbering)

    return {
        "enabled": True,
        "mode": "thesis",
        "cover": {
            "enabled": True,
            "logo": True,
            "institution": institution,
            "type_label": institution,
            "metadata": {
                "姓名": "（待填写）",
                "学号": "（待填写）",
                "导师": "（待填写）",
                "院系": "（待填写）",
                "学科/专业": "（待填写）",
                "申请学位": "（待填写）",
            },
            "date_text": "日期：____年__月__日",
        },
        "front_matter": {
            "abstract": True,
            "toc": True,
            # 法律声明需要作者明确确认和签名，不能从示例模板盲拷贝。
            "declarations": False,
        },
        "page_numbering": {
            "front_format": front_numbering.get("format", "upperRoman"),
            "front_start": front_numbering.get("start") or 1,
            "body_format": body_numbering.get("format", "decimal"),
            "body_start": body_numbering.get("start") or 1,
        },
        "running_header": {
            "left_text": institution,
            "chapter_style_name": "章标题",
            "header_distance_mm": round(doc.sections[0].header_distance.mm, 1),
            "footer_distance_mm": round(doc.sections[0].footer_distance.mm, 1),
        },
    }


def _first_paragraph(doc, predicate):
    return next((p for p in doc.paragraphs if predicate(p.text.strip().replace(" ", ""))), None)


def _add_thesis_roles(doc, roles, english=False):
    """把通用模板样式扩展成论文语义角色，并加入分页/段落连续性约束。"""
    body = deepcopy(roles["body"])
    heading_1 = deepcopy(roles.get("heading_1") or body)
    heading_2 = roles.get("heading_2")

    for role in ("heading_1", "heading_2", "heading_3"):
        if role in roles:
            roles[role]["keep_with_next"] = True
            roles[role]["keep_together"] = True
    if "heading_1" in roles and not english:
        roles["heading_1"]["page_break_before"] = True

    chapter_paragraph = None
    if english:
        academic_sections = {
            "introduction", "literaturereview", "relatedwork", "method",
            "methods", "methodology", "results", "discussion",
            "conclusion", "conclusions", "acknowledgments",
            "acknowledgements",
        }
        chapter_paragraph = _first_paragraph(
            doc,
            lambda text: (
                text.lower() in academic_sections
                or text.lower().startswith("chapter")
            ),
        )
    chapter = (
        _paragraph_rule(doc, chapter_paragraph, "chapter_heading")
        if chapter_paragraph is not None else deepcopy(heading_1)
    )
    chapter.update({
        "keep_with_next": True, "keep_together": True, "outline_level": 0,
    })
    if not english:
        chapter["page_break_before"] = True
    roles["chapter_heading"] = chapter

    abstract_heading_p = _first_paragraph(
        doc, (lambda text: text.lower() == "abstract") if english
        else (lambda text: text == "摘要"))
    abstract_heading = (
        _paragraph_rule(doc, abstract_heading_p, "abstract_heading")
        if abstract_heading_p is not None else deepcopy(heading_1)
    )
    abstract_heading.update({"keep_with_next": True, "keep_together": True})
    abstract_heading.pop("page_break_before", None)
    roles["abstract_heading"] = abstract_heading

    abstract_body = deepcopy(body)
    abstract_body.pop("numbering", None)
    abstract_body["first_line_indent_chars"] = 0
    abstract_body["label_prefix"] = {
        "text": ["Abstract:", "Abstract："] if english else ["摘要：", "摘要:"],
        "bold": True,
    }
    roles["abstract_body"] = abstract_body

    keywords = deepcopy(body)
    keywords.pop("numbering", None)
    keywords["first_line_indent_chars"] = 0
    keywords["label_prefix"] = {
        "text": ["Keywords:", "Keyword:", "Keywords：", "Keyword："]
        if english else ["关键词：", "关键词:", "关键字：", "关键字:"],
        "bold": True,
    }
    roles["keywords"] = keywords

    bibliography_heading_p = _first_paragraph(
        doc,
        (lambda text: text.lower() in {"references", "bibliography", "workscited"})
        if english else (lambda text: text == "参考文献"))
    bibliography_heading = (
        _paragraph_rule(doc, bibliography_heading_p, "bibliography_heading")
        if bibliography_heading_p is not None else deepcopy(heading_1)
    )
    bibliography_heading.update({
        "keep_with_next": True, "keep_together": True, "outline_level": 0,
    })
    if not english:
        bibliography_heading["page_break_before"] = True
    roles["bibliography_heading"] = bibliography_heading

    bibliography_entry = deepcopy(roles.get("attachment") or body)
    bibliography_entry.pop("numbering", None)
    bibliography_entry.pop("first_line_indent_chars", None)
    bibliography_entry.update({
        "alignment": "left", "left_indent_chars": 2,
        "hanging_indent_chars": 2, "widow_control": True,
    })
    roles["bibliography_entry"] = bibliography_entry

    equation = deepcopy(body)
    equation.pop("numbering", None)
    equation.pop("first_line_indent_chars", None)
    equation.update({"alignment": "center", "keep_together": True})
    roles["equation"] = equation

    appendix = deepcopy(chapter)
    roles["appendix_heading"] = appendix

    for role in ("list_of_figures_heading", "list_of_tables_heading"):
        list_heading = deepcopy(heading_1)
        list_heading.update({
            "keep_with_next": True, "keep_together": True,
            "outline_level": 0,
        })
        if not english:
            list_heading["page_break_before"] = True
        roles[role] = list_heading

    if heading_2 is not None:
        heading_2.setdefault("keep_with_next", True)


def _section_page_rule(section):
    """提取单节纸张、边距、栏数和页码格式。"""
    rule = {}
    if section.page_width is not None and section.page_height is not None:
        width = round(section.page_width.mm, 1)
        height = round(section.page_height.mm, 1)
        short, long = min(width, height), max(width, height)
        for name, (wanted_short, wanted_long) in {
            "A3": (297.0, 420.0),
            "A4": (210.0, 297.0),
            "A5": (148.0, 210.0),
            "letter": (215.9, 279.4),
            "legal": (215.9, 355.6),
        }.items():
            if abs(short - wanted_short) <= 1 and abs(long - wanted_long) <= 1:
                rule["size"] = name
                break
        else:
            rule["width_mm"] = short
            rule["height_mm"] = long
        rule["orientation"] = "landscape" if width > height else "portrait"
    margins = {
        key: round(getattr(section, attr).mm, 1)
        for key, attr in (
            ("top_mm", "top_margin"), ("bottom_mm", "bottom_margin"),
            ("left_mm", "left_margin"), ("right_mm", "right_margin"),
        )
        if getattr(section, attr) is not None
        and 5 <= getattr(section, attr).mm <= 50
    }
    if margins:
        rule["margin"] = margins
    columns = section._sectPr.find(qn("w:cols"))
    try:
        number = int(columns.get(qn("w:num"), "1")) if columns is not None else 1
    except ValueError:
        number = 1
    rule["columns"] = max(1, min(number, 4))
    page_numbering = _section_page_numbering(section)
    if page_numbering:
        rule["page_numbering"] = page_numbering
    return rule


def _page_section(doc):
    """单节输出全局规则；多节按索引保存各节契约，避免主流值覆盖特例。"""
    section_rules = [_section_page_rule(section) for section in doc.sections]
    if len(section_rules) == 1:
        page = section_rules[0]
        if page.get("columns") == 1:
            page.pop("columns", None)
    else:
        page = {"section_overrides": [
            {"section_index": index, **rule}
            for index, rule in enumerate(section_rules)
        ]}
    line_pitches = []
    for section in doc.sections:
        doc_grid = section._sectPr.find(qn("w:docGrid"))
        value = doc_grid.get(qn("w:linePitch")) if doc_grid is not None else None
        line_pitches.append(value)
    if line_pitches and len(set(line_pitches)) == 1 and line_pitches[0]:
        page["line_grid"] = {
            "line_pt": round(int(line_pitches[0]) / 20, 1)}
    return page




def _has_toc(doc):
    """模板里是否有目录域（TOC field）。"""
    body = doc.element.body
    for el in body.iter(qn("w:instrText")):
        if "TOC" in (el.text or ""):
            return True
    return False


def _header_footer_story_rule(hf, *, represent_empty=False):
    """读取一个页眉/页脚 Story；动态域不固化缓存文字。"""
    try:
        text_ps = [p for p in hf.paragraphs if p.text.strip()]
    except Exception:
        return None
    instructions = [
        (el.text or "").strip()
        for el in hf._element.iter(qn("w:instrText"))
        if (el.text or "").strip()
    ]
    has_page_field = any(
        instruction.upper().split(" ", 1)[0] == "PAGE"
        for instruction in instructions)
    has_dynamic_text = any(
        instruction.upper().split(" ", 1)[0] != "PAGE"
        for instruction in instructions)
    if not text_ps and not has_page_field and not represent_empty:
        return None
    rule = {}
    if text_ps and not has_dynamic_text:
        paragraph = text_ps[0]
        rule["text"] = paragraph.text.strip()
        props = effective_props(paragraph)
        if props.get("eastasia"):
            rule["font_eastasia"] = props["eastasia"]
        if props.get("ascii"):
            rule["font_ascii"] = props["ascii"]
        if props.get("cs"):
            rule["font_cs"] = props["cs"]
        if props.get("language"):
            rule["language"] = props["language"]
        if props.get("rtl") is not None:
            rule["rtl"] = bool(props["rtl"])
        if props.get("size_pt"):
            rule["size_pt"] = props["size_pt"]
        if props.get("bold") is not None:
            rule["bold"] = bool(props["bold"])
        alignment = _para_alignment(paragraph)
        if alignment:
            rule["alignment"] = alignment
    elif has_dynamic_text:
        rule["preserve_text"] = True
    elif represent_empty:
        rule["text"] = ""
    if has_page_field:
        rule["page_number"] = True
    return rule


def _header_footer_rules(doc):
    """读取所有节的默认/偶数页/首页 Story，并保留独立分节差异。"""
    rules = {}
    variants = (
        ("header", "header"), ("footer", "footer"),
        ("even_header", "even_page_header"),
        ("even_footer", "even_page_footer"),
        ("first_header", "first_page_header"),
        ("first_footer", "first_page_footer"),
    )
    if doc.settings.odd_and_even_pages_header_footer:
        rules["different_odd_even"] = True
    if doc.sections[0].different_first_page_header_footer:
        rules["different_first_page"] = True
    first = doc.sections[0]
    for which, accessor in variants:
        rule = _header_footer_story_rule(getattr(first, accessor))
        if rule is not None:
            rules[which] = rule

    overrides = []
    for section_index, section in enumerate(doc.sections[1:], start=1):
        override = {"section_index": section_index}
        if section.different_first_page_header_footer:
            override["different_first_page"] = True
        for which, accessor in variants:
            story = getattr(section, accessor)
            if story.is_linked_to_previous:
                continue
            rule = _header_footer_story_rule(story, represent_empty=True)
            if rule is not None:
                override[which] = rule
        if len(override) > 1:
            overrides.append(override)
    if overrides:
        rules["section_overrides"] = overrides
    return rules


def _merge_page_rules(base, additions):
    result = deepcopy(base)
    base_overrides = {
        item["section_index"]: deepcopy(item)
        for item in result.pop("section_overrides", [])
    }
    addition_overrides = additions.get("section_overrides") or []
    for key, value in additions.items():
        if key != "section_overrides":
            result[key] = deepcopy(value)
    for item in addition_overrides:
        index = item["section_index"]
        base_overrides.setdefault(index, {"section_index": index}).update(
            deepcopy(item))
    if base_overrides:
        result["section_overrides"] = [
            base_overrides[index] for index in sorted(base_overrides)
        ]
    return result


def _table_rule_for_table(t):
    """读取一张表的字体、几何、表头和跨页规则。"""
    rule = {}
    table_alignment = t.alignment
    if table_alignment is not None:
        rule["alignment"] = {0: "left", 1: "center", 2: "right"}.get(
            int(table_alignment), "left")
    tbl_pr = t._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is not None and layout.get(qn("w:type")) in {"fixed", "autofit"}:
        rule["layout"] = layout.get(qn("w:type"))
        rule["autofit"] = rule["layout"] == "autofit"
    width = tbl_pr.find(qn("w:tblW"))
    if width is not None:
        try:
            raw_width = int(width.get(qn("w:w")))
        except (TypeError, ValueError):
            raw_width = None
        if raw_width is not None and width.get(qn("w:type")) == "pct":
            rule["width_pct"] = round(raw_width / 50, 2)
        elif raw_width is not None and width.get(qn("w:type")) == "dxa" and raw_width > 0:
            rule["preferred_width_mm"] = round(raw_width / 1440 * 25.4, 2)
    grid_widths = []
    for grid_column in t._tbl.tblGrid.findall(qn("w:gridCol")):
        try:
            grid_widths.append(int(grid_column.get(qn("w:w"))))
        except (TypeError, ValueError):
            grid_widths = []
            break
    if grid_widths and sum(grid_widths) > 0:
        rule["column_widths_pct"] = [
            round(value / sum(grid_widths) * 100, 2) for value in grid_widths]
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is not None:
        values = {}
        for edge in ("top", "bottom", "left", "right"):
            element = margins.find(qn(f"w:{edge}"))
            try:
                values[edge] = round(int(element.get(qn("w:w"))) / 1440 * 25.4, 2)
            except (AttributeError, TypeError, ValueError):
                pass
        if values:
            rule["cell_margins_mm"] = values
    rows = t.rows
    if not rows:
        return None
    first_tr_pr = rows[0]._tr.trPr
    rule["repeat_header_row"] = bool(
        first_tr_pr is not None
        and first_tr_pr.find(qn("w:tblHeader")) is not None)
    sample_tr_pr = rows[-1]._tr.trPr
    rule["allow_row_break"] = not bool(
        sample_tr_pr is not None
        and sample_tr_pr.find(qn("w:cantSplit")) is not None)
    vertical = rows[0].cells[0].vertical_alignment
    if vertical is not None:
        rule["vertical_alignment"] = {
            0: "top", 1: "center", 3: "bottom",
        }.get(int(vertical), "top")
    header_props = None
    for p in rows[0].cells[0].paragraphs:
        header_props = effective_props(p)
        a = _para_alignment(p)
        if a:
            rule["header_alignment"] = a
        break
    if header_props and header_props.get("bold") is not None:
        rule["header_bold"] = bool(header_props["bold"])
    # 正文行字体字号
    body_row = rows[1] if len(rows) > 1 else rows[0]
    for p in body_row.cells[0].paragraphs:
        props = effective_props(p)
        if props.get("eastasia"):
            rule["font_eastasia"] = props["eastasia"]
        if props.get("ascii"):
            rule["font_ascii"] = props["ascii"]
        if props.get("cs"):
            rule["font_cs"] = props["cs"]
        if props.get("language"):
            rule["language"] = props["language"]
        if props.get("rtl") is not None:
            rule["rtl"] = bool(props["rtl"])
        if props.get("size_pt"):
            rule["size_pt"] = props["size_pt"]
        a = _para_alignment(p)
        if a:
            rule["body_alignment"] = a
        break
    if header_props:
        rule.setdefault("font_eastasia", header_props.get("eastasia") or "宋体")
        rule.setdefault("size_pt", header_props.get("size_pt") or 10.5)

    # 边框：直接 tblBorders 或表格样式（Table Grid 等）继承都算有边框
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is not None:
        any_border = any(
            (el.get(qn("w:val")) or "single") not in ("none", "nil")
            for el in borders)
        rule["borders"] = any_border
    else:
        style_name = ""
        try:
            style_name = (t.style.name or "").lower() if t.style else ""
        except Exception:
            style_name = ""
        if "grid" in style_name or "网格" in style_name:
            rule["borders"] = True
    return rule or None


def _table_rule(doc):
    """提取模板中的主流表格规则，并为异构表格生成按索引覆盖。"""
    if not doc.tables:
        return None
    rules = [_table_rule_for_table(table) or {} for table in doc.tables]
    signatures = [
        json.dumps(rule, ensure_ascii=False, sort_keys=True)
        for rule in rules
    ]
    from collections import Counter
    dominant_signature = Counter(signatures).most_common(1)[0][0]
    dominant_index = signatures.index(dominant_signature)
    result = deepcopy(rules[dominant_index])
    overrides = []
    for index, (rule, signature) in enumerate(zip(rules, signatures)):
        if signature == dominant_signature:
            continue
        override = {"table_index": index}
        override.update(deepcopy(rule))
        overrides.append(override)
    if overrides:
        result["overrides"] = overrides
    return result or None


def extract_rules_from_template(template_path, rolemap):
    """模板 docx + RoleMap → FormatSpec（经 schema 校验）。
    rolemap: {idx: role}。每个角色取第一个代表段落读格式。
    要求 rolemap 里至少有 body 角色，否则抛 ValueError。
    """
    doc = Document(template_path)
    paras = doc.paragraphs
    roles = {}
    representatives = _representative_paragraphs(paras, rolemap)
    for role, p in representatives.items():
        roles[role] = _paragraph_rule(doc, p, role)

    if "body" not in roles:
        raise ValueError("模板中没有标注 body 角色的段落，无法确定正文格式")
    profile = _detect_profile(doc)
    if profile in {"thesis", "english_academic"}:
        _add_thesis_roles(doc, roles, english=profile == "english_academic")

    page_rule = _page_section(doc)
    if profile == "thesis" and len(doc.sections) > 1:
        # 论文结构模式会按封面/前置页/正文重新建节；以末节正文几何作为基线，
        # 不把模板旧节索引强行映射到目标稿。
        line_grid = page_rule.get("line_grid")
        page_rule = _section_page_rule(doc.sections[-1])
        page_rule.pop("page_numbering", None)
        if page_rule.get("columns") == 1:
            page_rule.pop("columns", None)
        if line_grid:
            page_rule["line_grid"] = line_grid
    spec = {
        "profile": profile,
        "locale": "en-US" if profile.startswith("english_") else "zh-CN",
        "cleanup": {
            "mode": _cleanup_mode_for_template(doc, profile)
        },
        "page": page_rule,
        "roles": roles,
    }
    if profile == "thesis":
        spec["structure"] = _thesis_structure(doc)
    # 页眉页脚 + 表格规则（模板有就读，没有就不编）
    header_footer_rules = _header_footer_rules(doc)
    if profile == "thesis":
        header_footer_rules.pop("section_overrides", None)
    spec["page"] = _merge_page_rules(
        spec["page"], header_footer_rules)
    table_rule = _table_rule(doc)
    if table_rule:
        spec["table"] = table_rule
    if _has_toc(doc):
        spec["toc"] = {"enabled": True, "levels": [1, 2]}
    # 行网格一致性：模板里的 docGrid 常是 Word 默认值（如 15.6pt），与正文实际
    # 固定行距不一致时，网格会干扰排版。正文有明确固定行距时，以正文行距为准。
    body_ls = (roles.get("body") or {}).get("line_spacing") or {}
    if body_ls.get("type") == "exact" and body_ls.get("pt"):
        grid = spec["page"].setdefault("line_grid", {})
        if grid.get("line_pt") != body_ls["pt"]:
            grid["line_pt"] = body_ls["pt"]
    validate_spec(spec)
    return spec


if __name__ == "__main__":
    import json
    import sys
    with open(sys.argv[2], encoding="utf-8") as f:
        rolemap = {int(k): v for k, v in json.load(f).items()}
    spec = extract_rules_from_template(sys.argv[1], rolemap)
    print(json.dumps(spec, ensure_ascii=False, indent=2))
