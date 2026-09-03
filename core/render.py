# 渲染管线冒烟脚本 —— 在本机先跑通, 别到现场才试。
# 你的机器: Word COM 可用 (Office16), LibreOffice 不可用 -> 走 Word COM。
# 输出: docx -> PDF -> PNG, 供视觉验证/视觉读模板用。
#
# 环境要求: pywin32 (pip install pywin32)。python-docx 1.2.0 已装。
# 若现场没有 Word, 退回 LibreOffice headless, 需提前在赛前机器装好。
# 本脚本在 "python" 下纯 python-docx + win32com, 只依赖 docx + comtypes/pywin32。

import os
import json
import re
import shutil
import subprocess
import sys
import tempfile
import uuid
import zipfile
from pathlib import Path

from lxml import etree

from core.runtime import worker_command


class RenderQualityError(RuntimeError):
    """渲染成功生成文件，但页面内容明显不可用。"""


class CJKGlyphLossError(RenderQualityError):
    """PDF/PNG 中的中文字形因字体替换失败而丢失。"""


_CJK_RE = re.compile(r"[\u3400-\u4dbf\u4e00-\u9fff\uf900-\ufaff]")
_FONT_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_DRAWING_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
_OFFICE_REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_PACKAGE_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
_CONTENT_TYPES_NS = "http://schemas.openxmlformats.org/package/2006/content-types"


def _installed_font_families():
    """从 fontconfig 读取可用字体族；Windows 无 fc-list 时返回空集。"""
    command = shutil.which("fc-list")
    if not command:
        return set()
    try:
        result = subprocess.run(
            [command, "-f", "%{family}\n"], capture_output=True, text=True,
            timeout=15, check=False)
    except (OSError, subprocess.SubprocessError):
        return set()
    families = set()
    for line in result.stdout.splitlines():
        for family in line.split(","):
            value = family.strip()
            if value:
                families.add(value.casefold())
    return families


def _cjk_fallback_font():
    explicit = os.environ.get("LIBREOFFICE_CJK_FALLBACK_FONT", "").strip()
    if explicit:
        return explicit
    installed = _installed_font_families()
    candidates = (
        "Microsoft YaHei", "Noto Sans SC", "Noto Sans CJK SC", "Source Han Sans SC",
        "WenQuanYi Zen Hei", "Hiragino Sans GB", "PingFang SC",
        "Arial Unicode MS",
    )
    if installed:
        for candidate in candidates:
            if candidate.casefold() in installed:
                return candidate
    # Windows 常见可用；若仍不存在，二次渲染质量门会报出来。
    return "Microsoft YaHei"


def _font_file_for_family(family):
    command = shutil.which("fc-match")
    if not command:
        return None
    try:
        result = subprocess.run(
            [command, "-f", "%{file}\n", family], capture_output=True, text=True,
            timeout=15, check=False)
    except (OSError, subprocess.SubprocessError):
        return None
    path = result.stdout.splitlines()[0].strip() if result.stdout.splitlines() else ""
    return path if os.path.isfile(path) else None


def _looks_like_cjk_font_name(value):
    value = str(value or "")
    lowered = value.casefold()
    return bool(
        _CJK_RE.search(value)
        or "gb2312" in lowered
        or any(token in lowered for token in (
            "simsun", "simhei", "fangsong", "kaiti", "yahei",
            "songti", "heiti", "pingfang", "hiragino sans gb",
            "noto sans cjk", "source han sans",
        ))
    )


def _make_cjk_fallback_copy(source_path, output_path, fallback_font):
    """只为 LibreOffice 渲染制作字体替代副本，绝不改写交付 DOCX。"""
    changed = 0
    with zipfile.ZipFile(source_path, "r") as source:
        infos = source.infolist()
        payloads = {info.filename: source.read(info.filename) for info in infos}

    for filename, raw in list(payloads.items()):
        if not (filename.startswith("word/") and filename.endswith(".xml")):
            continue
        try:
            root = etree.fromstring(raw)
        except etree.XMLSyntaxError:
            continue

        for fonts in root.iter(f"{{{_FONT_NS}}}rFonts"):
            eastasia = f"{{{_FONT_NS}}}eastAsia"
            eastasia_theme = f"{{{_FONT_NS}}}eastAsiaTheme"
            if fonts.get(eastasia) != fallback_font:
                fonts.set(eastasia, fallback_font)
                changed += 1
            fonts.attrib.pop(eastasia_theme, None)
            for name in ("ascii", "hAnsi", "cs"):
                attribute = f"{{{_FONT_NS}}}{name}"
                if _looks_like_cjk_font_name(fonts.get(attribute)):
                    fonts.set(attribute, fallback_font)

        # Word 主题字体也可以覆盖段落样式的 East Asia 字体。
        for element in root.iter(f"{{{_DRAWING_NS}}}ea"):
            element.set("typeface", fallback_font)
        for element in root.iter(f"{{{_DRAWING_NS}}}font"):
            if element.get("script") in {"Hans", "Hant"}:
                element.set("typeface", fallback_font)
        payloads[filename] = etree.tostring(
            root, xml_declaration=True, encoding="UTF-8", standalone=True)

    # LibreOffice on macOS 有时看不到系统/用户字体。将开源兜底字体
    # 嵌入这份临时 DOCX，使渲染结果不再依赖字体搜索路径。
    font_file = _font_file_for_family(fallback_font)
    font_table_name = "word/fontTable.xml"
    if font_file and font_table_name in payloads:
        font_key = uuid.uuid4()
        relationship_id = "rIdFormatAgentFallback"
        font_part_name = "word/fonts/formatagent-fallback.odttf"

        font_data = bytearray(Path(font_file).read_bytes())
        obfuscation_key = font_key.bytes[::-1]
        for index in range(min(32, len(font_data))):
            font_data[index] ^= obfuscation_key[index % 16]
        payloads[font_part_name] = bytes(font_data)

        font_root = etree.fromstring(payloads[font_table_name])
        font_element = None
        for candidate in font_root.iter(f"{{{_FONT_NS}}}font"):
            if candidate.get(f"{{{_FONT_NS}}}name") == fallback_font:
                font_element = candidate
                break
        if font_element is None:
            font_element = etree.SubElement(font_root, f"{{{_FONT_NS}}}font")
            font_element.set(f"{{{_FONT_NS}}}name", fallback_font)
        for embed in list(font_element):
            if embed.tag == f"{{{_FONT_NS}}}embedRegular":
                font_element.remove(embed)
        embed = etree.SubElement(font_element, f"{{{_FONT_NS}}}embedRegular")
        embed.set(f"{{{_OFFICE_REL_NS}}}id", relationship_id)
        embed.set(f"{{{_FONT_NS}}}fontKey", "{" + str(font_key).upper() + "}")
        payloads[font_table_name] = etree.tostring(
            font_root, xml_declaration=True, encoding="UTF-8", standalone=True)

        rels_name = "word/_rels/fontTable.xml.rels"
        if rels_name in payloads:
            rels_root = etree.fromstring(payloads[rels_name])
        else:
            rels_root = etree.Element(f"{{{_PACKAGE_REL_NS}}}Relationships")
        for relationship in list(rels_root):
            if relationship.get("Id") == relationship_id:
                rels_root.remove(relationship)
        relationship = etree.SubElement(
            rels_root, f"{{{_PACKAGE_REL_NS}}}Relationship")
        relationship.set("Id", relationship_id)
        relationship.set(
            "Type", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/font")
        relationship.set("Target", "fonts/formatagent-fallback.odttf")
        payloads[rels_name] = etree.tostring(
            rels_root, xml_declaration=True, encoding="UTF-8", standalone=True)

        content_types_name = "[Content_Types].xml"
        content_root = etree.fromstring(payloads[content_types_name])
        part_name = "/" + font_part_name
        existing_override = next((
            element for element in content_root
            if element.tag == f"{{{_CONTENT_TYPES_NS}}}Override"
            and element.get("PartName") == part_name
        ), None)
        if existing_override is None:
            override = etree.SubElement(
                content_root, f"{{{_CONTENT_TYPES_NS}}}Override")
            override.set("PartName", part_name)
            override.set(
                "ContentType",
                "application/vnd.openxmlformats-officedocument.obfuscatedFont")
        payloads[content_types_name] = etree.tostring(
            content_root, xml_declaration=True, encoding="UTF-8", standalone=True)
        changed += 1

    original_names = {info.filename for info in infos}
    with zipfile.ZipFile(output_path, "w") as target:
        for info in infos:
            target.writestr(info, payloads[info.filename])
        for filename, data in payloads.items():
            if filename not in original_names:
                target.writestr(filename, data)
    return changed


def _docx_cjk_count(docx_path):
    from docx import Document

    doc = Document(docx_path)
    chunks = [paragraph.text for paragraph in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                chunks.extend(paragraph.text for paragraph in cell.paragraphs)
    for section in doc.sections:
        for part in (section.header, section.footer):
            chunks.extend(paragraph.text for paragraph in part.paragraphs)
    return len(_CJK_RE.findall("".join(chunks)))


def _rendered_ink_stats(paths):
    from PIL import Image

    ink_pixels = 0
    total_pixels = 0
    for path in paths:
        with Image.open(path) as source:
            image = source.convert("L")
            image.thumbnail((700, 700))
            histogram = image.histogram()
            ink_pixels += sum(histogram[:245])
            total_pixels += image.width * image.height
    return ink_pixels, total_pixels


def validate_rendered_pages(docx_path, pages):
    """阻止空页或中文字形丢失的 PNG 继续送给视觉模型。"""
    if not pages:
        raise RenderQualityError("渲染器没有生成任何 PNG 页面")
    for path in pages:
        if not os.path.isfile(path) or os.path.getsize(path) <= 0:
            raise RenderQualityError(f"渲染页面不存在或为空: {path}")
    cjk_count = _docx_cjk_count(docx_path)
    ink_pixels, total_pixels = _rendered_ink_stats(pages)
    ink_ratio = ink_pixels / total_pixels if total_pixels else 0.0
    ink_per_cjk = ink_pixels / cjk_count if cjk_count else None
    # 同时要求“页面几乎全白”和“每个中文字对应的墨迹极少”，
    # 避免把故意留白的短文档误报为缺字。
    if (
        cjk_count >= 20
        and ink_ratio < 0.003
        and ink_per_cjk is not None
        and ink_per_cjk < 12
    ):
        raise CJKGlyphLossError(
            "LibreOffice 已生成页面，但中文字形几乎全部丢失"
            f"（中文 {cjk_count} 字，墨迹比 {ink_ratio:.4f}）。"
            "通常是方正小标宋/仿宋_GB2312/楷体_GB2312 等字体未安装。")
    return {
        "page_count": len(pages),
        "cjk_count": cjk_count,
        "ink_ratio": ink_ratio,
        "ink_per_cjk": ink_per_cjk,
    }


def docx_to_pdf_word_com(docx_path, pdf_path):
    """在有超时边界的独立进程中用 Microsoft Word 导出 PDF。"""
    return _docx_to_pdf_com_candidate(
        docx_path, pdf_path, "Word.Application", "Microsoft Word")


def _docx_to_pdf_com_candidate(docx_path, pdf_path, prog_id, display_name):
    from core.input_conversion import (
        _com_timeout_seconds,
        _office_pids,
        _office_process_names,
        _terminate_new_office_processes,
    )

    source = Path(docx_path).expanduser().resolve()
    destination = Path(pdf_path).expanduser().resolve()
    process_names = _office_process_names(prog_id)
    before = _office_pids(process_names)
    command = worker_command(
        "com-pdf", source, destination, prog_id, display_name)
    creationflags = getattr(subprocess, "CREATE_NO_WINDOW", 0)
    timeout = _com_timeout_seconds()
    try:
        result = subprocess.run(
            command, capture_output=True, text=True, encoding="utf-8", timeout=timeout,
            check=False, creationflags=creationflags,
        )
    except subprocess.TimeoutExpired as exc:
        terminated = _terminate_new_office_processes(process_names, before)
        cleanup = f"；已终止新建 Office 进程 {terminated}" if terminated else ""
        raise RuntimeError(
            f"{display_name} PDF 渲染在 {timeout:g} 秒后超时{cleanup}") from exc
    output = (result.stdout or "").strip().splitlines()
    try:
        payload = json.loads(output[-1]) if output else {}
    except json.JSONDecodeError:
        payload = {}
    if result.returncode or not payload.get("ok"):
        detail = payload.get("error") or (result.stderr or result.stdout or "无诊断输出").strip()
        raise RuntimeError(f"{display_name} PDF 渲染失败：{detail}")
    if not destination.is_file() or destination.read_bytes()[:5] != b"%PDF-":
        raise RuntimeError(f"{display_name} 没有生成有效 PDF：{destination}")
    return str(destination)


def docx_to_pdf_office_com(docx_path, pdf_path):
    """Microsoft Word 无响应或不可用时继续回退到 WPS。"""
    errors = []
    for prog_id, display_name in (
        ("Word.Application", "Microsoft Word"),
        ("Kwps.Application", "WPS"),
        ("wps.Application", "WPS"),
    ):
        try:
            return _docx_to_pdf_com_candidate(
                docx_path, pdf_path, prog_id, display_name)
        except RuntimeError as exc:
            errors.append(str(exc))
    raise RuntimeError("；".join(errors))


def _soffice_candidates():
    """按优先级列出跨平台候选；不依赖 Streamlit 进程是否继承终端 PATH。"""
    candidates = []
    for key in ("LIBREOFFICE_PATH", "SOFFICE_PATH"):
        value = os.environ.get(key)
        if value:
            value_path = Path(value).expanduser()
            if value_path.is_dir():
                value_path = value_path / "Contents" / "MacOS" / "soffice"
            candidates.append(str(value_path))
    for command in ("soffice", "libreoffice"):
        found = shutil.which(command)
        if found:
            candidates.append(found)

    home = Path.home()
    candidates.extend([
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        "/Applications/LibreOfficeDev.app/Contents/MacOS/soffice",
        str(home / "Applications/LibreOffice.app/Contents/MacOS/soffice"),
        str(home / "Applications/LibreOfficeDev.app/Contents/MacOS/soffice"),
        "/opt/homebrew/bin/soffice",
        "/usr/local/bin/soffice",
        "/usr/bin/soffice",
        "/usr/bin/libreoffice",
        # Codex 桌面环境的文档运行时兜底；正式部署仍推荐系统安装或显式变量。
        str(home / ".cache/codex-runtimes/codex-primary-runtime/dependencies/bin/override/soffice"),
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ])
    # 去重但保持优先级，便于错误信息按实际搜索顺序展示。
    return list(dict.fromkeys(candidates))


def _find_soffice():
    """返回可执行且能响应 ``--version`` 的 LibreOffice 路径。"""
    for candidate in _soffice_candidates():
        if not os.path.isfile(candidate) or not os.access(candidate, os.X_OK):
            continue
        try:
            probe = subprocess.run(
                [candidate, "--version"], capture_output=True, text=True,
                timeout=10, check=False)
        except (OSError, subprocess.SubprocessError):
            continue
        if probe.returncode == 0:
            return candidate
    return None


def libreoffice_status():
    """供网页诊断展示，避免只给一句“没安装”。"""
    path = _find_soffice()
    version = None
    if path:
        try:
            probe = subprocess.run(
                [path, "--version"], capture_output=True, text=True,
                timeout=10, check=False)
            if probe.returncode == 0:
                version = (probe.stdout or probe.stderr).strip() or None
        except (OSError, subprocess.SubprocessError):
            version = None
    return {
        "available": path is not None,
        "path": path,
        "version": version,
        "searched": _soffice_candidates(),
    }


def renderer_status():
    """渲染器总诊断（供网页用）：Windows 优先 Word COM，其次 LibreOffice；
    其他平台走 LibreOffice。与 docx_to_pdf 的实际选择保持一致——
    之前网页只看 LibreOffice，Windows 上有 Word 也会误判"渲染不可用"。
    """
    if sys.platform == "win32":
        try:
            import win32com.client  # noqa: F401
            return {"available": True, "name": "Word COM",
                    "version": "Microsoft Word (COM)", "path": None}
        except ImportError:
            pass
    lo = libreoffice_status()
    return {"available": lo["available"], "name": "LibreOffice",
            "version": lo["version"], "path": lo["path"],
            "searched": lo.get("searched")}


def docx_to_pdf_libreoffice(docx_path, pdf_path):
    """用 LibreOffice headless 导出 PDF，供 macOS/Linux 使用。"""
    soffice = _find_soffice()
    if not soffice:
        searched = "\n  - ".join(_soffice_candidates())
        raise RuntimeError(
            "未找到可用的 LibreOffice。可设置 LIBREOFFICE_PATH 指向 soffice，"
            "或在 macOS 安装 LibreOffice。\n已检查：\n  - " + searched)
    docx_path = os.path.abspath(docx_path)
    pdf_path = os.path.abspath(pdf_path)
    with tempfile.TemporaryDirectory() as work:
        converted_dir = os.path.join(work, "converted")
        profile_dir = os.path.join(work, "profile")
        os.makedirs(converted_dir, exist_ok=True)
        os.makedirs(profile_dir, exist_ok=True)
        command = [
            soffice,
            f"-env:UserInstallation={Path(profile_dir).resolve().as_uri()}",
            "--headless", "--norestore", "--convert-to", "pdf",
            "--outdir", converted_dir, docx_path,
        ]
        render_env = dict(os.environ)
        # Homebrew LibreOffice 26.x 在 macOS 上有时启用一份没有
        # cachedir/font 搜索路径的内置 fontconfig，结果是中文全部
        # 落到 LinuxLibertine 并只剩数字。显式指向 Homebrew 的完整配置。
        if sys.platform == "darwin" and not render_env.get("FONTCONFIG_FILE"):
            for config in (
                "/opt/homebrew/etc/fonts/fonts.conf",
                "/usr/local/etc/fonts/fonts.conf",
            ):
                if os.path.isfile(config):
                    render_env["FONTCONFIG_FILE"] = config
                    break
        result = subprocess.run(
            command, env=render_env, capture_output=True, text=True, timeout=120,
            check=False)
        generated = os.path.join(
            converted_dir, os.path.splitext(os.path.basename(docx_path))[0] + ".pdf")
        if result.returncode != 0 or not os.path.isfile(generated):
            details = (result.stderr or result.stdout or "无详细错误")[-800:]
            raise RuntimeError(f"LibreOffice 导出 PDF 失败：{details}")
        shutil.copy2(generated, pdf_path)
    return pdf_path


def docx_to_pdf(docx_path, pdf_path):
    """按平台选择 DOCX → PDF 渲染器。"""
    if sys.platform == "win32":
        errors = []
        try:
            return docx_to_pdf_office_com(docx_path, pdf_path)
        except RuntimeError as exc:
            errors.append(str(exc))
        try:
            return docx_to_pdf_libreoffice(docx_path, pdf_path)
        except RuntimeError as exc:
            errors.append(str(exc))
        raise RuntimeError("DOCX 渲染失败：" + "；".join(errors))
    return docx_to_pdf_libreoffice(docx_path, pdf_path)


def pdf_to_png(pdf_path, png_dir, dpi=150):
    """PDF 转 PNG (逐页)。用 fitz (PyMuPDF) 或 pymupdf。备选: pdfium / pdf2image。
    优先 fitz, 没装则降级为 pymupdf。
    """
    os.makedirs(png_dir, exist_ok=True)
    try:
        import pymupdf as fitz
    except ImportError:
        try:
            import fitz  # 兼容较旧的 PyMuPDF
        except ImportError:
            fitz = None
    pages = []
    if fitz is not None:
        pdf = fitz.open(pdf_path)
        try:
            for i, page in enumerate(pdf):
                pix = page.get_pixmap(dpi=dpi)
                png_path = os.path.join(png_dir, f"page_{i:02d}.png")
                pix.save(png_path)
                pages.append(png_path)
        finally:
            pdf.close()
        return pages

    try:
        import pypdfium2 as pdfium
    except ImportError:
        pdfium = None
    if pdfium is not None:
        pdf = pdfium.PdfDocument(pdf_path)
        try:
            for i in range(len(pdf)):
                page = pdf[i]
                bitmap = page.render(scale=dpi / 72.0)
                png_path = os.path.join(png_dir, f"page_{i:02d}.png")
                try:
                    bitmap.to_pil().save(png_path)
                finally:
                    bitmap.close()
                    page.close()
                pages.append(png_path)
        finally:
            pdf.close()
        return pages

    try:
        from pdf2image import convert_from_path
    except ImportError as exc:
        raise RuntimeError(
            "需要安装 PyMuPDF、pypdfium2 或 pdf2image 之一") from exc
    for i, image in enumerate(convert_from_path(pdf_path, dpi=dpi)):
        png_path = os.path.join(png_dir, f"page_{i:02d}.png")
        image.save(png_path, "PNG")
        pages.append(png_path)
    return pages


def _render_once(docx_path, png_dir, dpi):
    os.makedirs(png_dir, exist_ok=True)
    with tempfile.TemporaryDirectory() as tmpdir:
        pdf_path = os.path.join(tmpdir, "out.pdf")
        docx_to_pdf(docx_path, pdf_path)
        return pdf_to_png(pdf_path, png_dir, dpi)


def render_docx_to_png(docx_path, png_dir, dpi=150, on_event=None):
    """一键: docx -> pdf -> png 列表，并在送 VLM 前检查缺字。

    LibreOffice 因 Windows 中文字体缺失而渲染近乎空白时，仅对临时
    渲染副本做 CJK 字体替代后重试。输出 DOCX 始终不受影响。
    """
    on_event = on_event or (lambda _message: None)
    pages = _render_once(docx_path, png_dir, dpi)
    try:
        validate_rendered_pages(docx_path, pages)
        return pages
    except CJKGlyphLossError as original_error:
        if sys.platform == "win32":
            raise

        fallback_font = _cjk_fallback_font()
        on_event(
            "LibreOffice 渲染检测到中文字体缺失；正在用 "
            f"{fallback_font} 制作仅供视觉检查的临时渲染副本")
        for path in pages:
            if os.path.isfile(path):
                os.unlink(path)
        with tempfile.TemporaryDirectory() as fallback_dir:
            fallback_docx = os.path.join(fallback_dir, "render-font-fallback.docx")
            changed = _make_cjk_fallback_copy(
                docx_path, fallback_docx, fallback_font=fallback_font)
            if changed <= 0:
                raise original_error
            pages = _render_once(fallback_docx, png_dir, dpi)
            try:
                validate_rendered_pages(fallback_docx, pages)
            except CJKGlyphLossError as fallback_error:
                raise RenderQualityError(
                    f"{original_error} 已尝试临时替换为 {fallback_font}，"
                    f"但中文仍未正常渲染: {fallback_error}") from fallback_error
        on_event(
            f"临时字体替代渲染成功（替换 {changed} 处）。"
            "该替代只影响视觉检查图，不会改写输出 DOCX")
        return pages


if __name__ == "__main__":
    # 冒烟: 用你手边任意一个 docx 试跑, 无误再进主流程。
    if len(sys.argv) < 2:
        print("用法: python hackathon-render.py <某个docx路径>")
        sys.exit(1)
    src = sys.argv[1]
    png_dir = "recon_render_test"
    pages = render_docx_to_png(src, png_dir)
    print(f"渲染成功: {len(pages)} 页 -> {png_dir}/page_*.png")
