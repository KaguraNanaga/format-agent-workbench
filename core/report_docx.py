# docx 版修改对照报告：比 markdown 更适合交付/归档。
# build_report_docx(changelog, spec, path) -> path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def _fmt_fields(fields):
    return "、".join(fields) if fields else "（无字段改动）"


def build_report_docx(changelog, spec, path):
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("排版修改对照报告")
    run.font.size = Pt(18)
    run.font.bold = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(f"共 {len(changelog)} 个段落参与排版 · 由 Format Agent 生成")
    meta_run.font.size = Pt(10.5)

    # ---- 页面设置 ----
    page = spec.get("page") or {}
    if page:
        doc.add_heading("一、页面设置", level=1)
        margin = page.get("margin") or {}
        if margin:
            doc.add_paragraph(
                f"页边距（毫米）：上 {margin.get('top_mm', '-')} / 下 {margin.get('bottom_mm', '-')}"
                f" / 左 {margin.get('left_mm', '-')} / 右 {margin.get('right_mm', '-')}",
                style="List Bullet")
        lg = page.get("line_grid") or {}
        if lg.get("line_pt"):
            doc.add_paragraph(f"行网格：每行 {lg['line_pt']} 磅", style="List Bullet")

    # ---- 角色规则 ----
    roles = spec.get("roles") or {}
    if roles:
        doc.add_heading("二、套用的格式规则（FormatSpec）", level=1)
        for role, rule in roles.items():
            parts = []
            if rule.get("font_eastasia"):
                parts.append(str(rule["font_eastasia"]))
            if rule.get("size_pt"):
                parts.append(f"{rule['size_pt']} 磅")
            if rule.get("alignment"):
                parts.append(f"对齐 {rule['alignment']}")
            ls = rule.get("line_spacing") or {}
            if ls.get("pt"):
                parts.append(f"行距 {ls['pt']} 磅")
            if rule.get("first_line_indent_chars"):
                parts.append(f"首行缩进 {rule['first_line_indent_chars']} 字符")
            doc.add_paragraph(f"{role}：{'，'.join(parts)}", style="List Bullet")

    # ---- 段落明细 ----
    doc.add_heading("三、段落修改明细", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["段落", "角色", "Word 样式", "改动字段", "内容摘要"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        for r in cell.paragraphs[0].runs:
            r.font.bold = True
    for c in changelog:
        row = table.add_row().cells
        row[0].text = str(c["idx"])
        row[1].text = str(c["role"])
        row[2].text = str(c.get("style_name", "-"))
        row[3].text = _fmt_fields(c["changed_fields"])
        row[4].text = str(c.get("text", ""))
        if c.get("fallback_to_target_body"):
            row[3].text += "（模板未规定，与正文一致）"

    doc.save(path)
    return path
