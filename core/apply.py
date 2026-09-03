# 执行器接线（PLAN.md 第 7 节）：
# apply_format(docx_path, spec, rolemap, out_path) -> changelog
# 对每个段落按 RoleMap 取角色、从 FormatSpec 取规则，调用 core/executor.py 的
# 确定性函数改 XML；页边距/行网格走 section 级别。LLM 不碰 docx。

from copy import deepcopy

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.shared import Mm

from core.executor import (
    set_doc_grid,
    set_run_fonts,
)
from core.style_set import (
    apply_named_style,
    clear_invalid_numbering_override,
    effective_automatic_numbering_override,
    ensure_role_styles,
    resolve_target_body_style,
)
from core.track_changes import (
    flatten_style_formatting,
    mark_paragraph_revision,
    snapshot_paragraph,
)

_HF_ALIGN = {"left": 0, "center": 1, "right": 2, "justify": 3}
_TABLE_ALIGN = {
    "left": WD_TABLE_ALIGNMENT.LEFT,
    "center": WD_TABLE_ALIGNMENT.CENTER,
    "right": WD_TABLE_ALIGNMENT.RIGHT,
}
_CELL_VERTICAL_ALIGN = {
    "top": WD_CELL_VERTICAL_ALIGNMENT.TOP,
    "center": WD_CELL_VERTICAL_ALIGNMENT.CENTER,
    "bottom": WD_CELL_VERTICAL_ALIGNMENT.BOTTOM,
}
_PAGE_SIZES_MM = {
    "A3": (297.0, 420.0),
    "A4": (210.0, 297.0),
    "A5": (148.0, 210.0),
    "letter": (215.9, 279.4),
    "legal": (215.9, 355.6),
}

_ROLE_FALLBACKS = {
    "abstract_heading": ("heading_1", "body"),
    "abstract_body": ("body",),
    "keywords": ("body",),
    "chapter_heading": ("heading_1",),
    "bibliography_heading": ("heading_1",),
    "bibliography_entry": ("attachment", "body"),
    "equation": ("body",),
    "appendix_heading": ("heading_1",),
    "list_of_figures_heading": ("heading_1",),
    "list_of_tables_heading": ("heading_1",),
    "block_quote": ("body",),
    "code_block": ("body",),
    "byline": ("subtitle", "body"),
    "affiliation": ("body",),
    "author_note": ("body",),
    "correspondence": ("body",),
    "salutation": ("body",),
    "complimentary_close": ("signature", "body"),
    "cc": ("attachment", "body"),
    "enclosure": ("attachment", "body"),
    "legal_definition": ("body",),
    "signature_block": ("signature", "body"),
    "table_of_authorities_heading": ("heading_1",),
    "heading_4": ("heading_3", "body"),
    "recipient": ("body",),
    "closing": ("body",),
    "document_number": ("subtitle", "body"),
    "copy_to": ("attachment", "body"),
    "warning_box": ("body",),
    "caution_box": ("body",),
    "note_box": ("body",),
    "tip_box": ("body",),
    "procedure_step": ("heading_3", "body"),
    "command": ("code_block", "body"),
    "court_caption": ("title", "body"),
    "case_number": ("subtitle", "body"),
    "brief_title": ("title",),
    "table_of_contents_heading": ("heading_1",),
    "authority_entry": ("bibliography_entry", "body"),
    "counsel_block": ("signature_block", "body"),
    "certificate_heading": ("heading_1",),
    "certificate_body": ("body",),
}


def _resolved_role(role, roles):
    if role in roles:
        return role
    return next((candidate for candidate in _ROLE_FALLBACKS.get(role, ()) if candidate in roles), None)


def _apply_section_geometry(section, rule):
    size = rule.get("size")
    custom_width = rule.get("width_mm")
    custom_height = rule.get("height_mm")
    orientation = rule.get("orientation")
    if not size and custom_width is None and not orientation:
        return []
    if size:
        width_mm, height_mm = _PAGE_SIZES_MM[size]
    elif custom_width is not None and custom_height is not None:
        width_mm, height_mm = sorted((float(custom_width), float(custom_height)))
    else:
        width_mm = section.page_width.mm
        height_mm = section.page_height.mm
        width_mm, height_mm = min(width_mm, height_mm), max(width_mm, height_mm)
    landscape = (
        orientation == "landscape"
        if orientation else section.orientation == WD_ORIENT.LANDSCAPE
    )
    section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
    if landscape:
        width_mm, height_mm = max(width_mm, height_mm), min(width_mm, height_mm)
    else:
        width_mm, height_mm = min(width_mm, height_mm), max(width_mm, height_mm)
    section.page_width = Mm(width_mm)
    section.page_height = Mm(height_mm)
    return [value for value in (
                                "page_size" if size or custom_width is not None else None,
                                "page_orientation" if orientation else None)
            if value]


def _apply_margin(section, margin):
    for key, attr in (
        ("top_mm", "top_margin"), ("bottom_mm", "bottom_margin"),
        ("left_mm", "left_margin"), ("right_mm", "right_margin"),
    ):
        if margin.get(key) is not None:
            setattr(section, attr, Mm(margin[key]))


def _apply_header_footer_distances(section, rule):
    changed = []
    if rule.get("header_distance_mm") is not None:
        section.header_distance = Mm(rule["header_distance_mm"])
        changed.append("header_distance")
    if rule.get("footer_distance_mm") is not None:
        section.footer_distance = Mm(rule["footer_distance_mm"])
        changed.append("footer_distance")
    return changed


def _set_section_columns(section, number):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    if not isinstance(number, int) or number < 1:
        return False
    columns = section._sectPr.find(qn("w:cols"))
    if columns is None:
        columns = OxmlElement("w:cols")
        section._sectPr.append(columns)
    columns.set(qn("w:num"), str(number))
    return True


def _set_section_page_numbering(section, rule):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    if not isinstance(rule, dict) or not rule:
        return False
    element = section._sectPr.find(qn("w:pgNumType"))
    if element is None:
        element = OxmlElement("w:pgNumType")
        section._sectPr.append(element)
    if rule.get("format") is not None:
        element.set(qn("w:fmt"), str(rule["format"]))
    if rule.get("start") is not None:
        element.set(qn("w:start"), str(int(rule["start"])))
    return True


def _apply_header_footer(doc, page):
    """应用默认/偶数页/首页页眉页脚，并保留多节链接和独立 Story。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    changed = []
    variants = (
        ("header", "header"), ("footer", "footer"),
        ("even_header", "even_page_header"),
        ("even_footer", "even_page_footer"),
        ("first_header", "first_page_header"),
        ("first_footer", "first_page_footer"),
    )
    overrides = {
        value.get("section_index"): value
        for value in page.get("section_overrides") or []
        if isinstance(value, dict) and isinstance(value.get("section_index"), int)
    }
    if (
        page.get("different_odd_even")
        or any(page.get(key) for key, _ in variants[2:4])
        or any(
            override.get("different_odd_even")
            or any(override.get(key) for key, _ in variants[2:4])
            for override in overrides.values()
        )
    ):
        doc.settings.odd_and_even_pages_header_footer = True

    processed_parts = set()
    text_applied = {key: False for key, _ in variants}
    for section_index, section in enumerate(doc.sections):
        override = overrides.get(section_index) or {}
        if (
            override.get("different_first_page")
            or page.get("different_first_page")
            or any((override.get(key) if key in override else page.get(key))
                   for key, _ in variants[4:])
        ):
            section.different_first_page_header_footer = True
        for rule_key, accessor in variants:
            explicit_override = rule_key in override
            rule = override.get(rule_key) if explicit_override else page.get(rule_key)
            if not isinstance(rule, dict) or not rule:
                continue
            container = getattr(section, accessor)
            if explicit_override and section_index > 0 and container.is_linked_to_previous:
                container.is_linked_to_previous = False
            # 链接到前一节的页眉页脚共享同一个 part：只处理一次，不破坏链接关系。
            part_key = str(container.part.partname)
            if part_key in processed_parts:
                continue
            processed_parts.add(part_key)
            # 模板文字只替换首节对应 Story，保持旧版单节语义；已有多节文档中
            # 其余独立页眉/页脚往往承载章名或分册信息，不能被同一文本抹平。
            active_paragraph = None
            if (
                rule.get("text") is not None
                and not rule.get("preserve_text")
                and (explicit_override or not text_applied[rule_key])
            ):
                p = (
                    container.paragraphs[0]
                    if container.paragraphs else container.add_paragraph()
                )
                for run in p.runs:
                    run.text = ""
                if p.runs:
                    p.runs[0].text = rule["text"]
                else:
                    p.add_run(rule["text"])
                active_paragraph = p
                changed.append(f"{rule_key}_text")
                if not explicit_override:
                    text_applied[rule_key] = True
            has_page_field = any(
                (element.text or "").strip() == "PAGE"
                for element in container._element.iter(qn("w:instrText"))
            )
            if rule.get("page_number") and not has_page_field:
                p = active_paragraph or (
                    container.paragraphs[0]
                    if container.paragraphs else container.add_paragraph())
                prefix = rule.get("page_number_prefix")
                suffix = rule.get("page_number_suffix")
                if p.text and prefix is None:
                    p.add_run(" ")
                if prefix:
                    p.add_run(prefix)
                run = p.add_run()
                fld_begin = OxmlElement("w:fldChar")
                fld_begin.set(qn("w:fldCharType"), "begin")
                instr = OxmlElement("w:instrText")
                instr.set(qn("xml:space"), "preserve")
                instr.text = "PAGE"
                fld_end = OxmlElement("w:fldChar")
                fld_end.set(qn("w:fldCharType"), "end")
                run._element.append(fld_begin)
                run._element.append(instr)
                run._element.append(fld_end)
                if suffix:
                    p.add_run(suffix)
                changed.append(f"{rule_key}_page_number")
            font_kwargs = {}
            if rule.get("font_eastasia"):
                font_kwargs["eastasia"] = rule["font_eastasia"]
            if rule.get("font_ascii"):
                font_kwargs["ascii_font"] = rule["font_ascii"]
            if rule.get("font_cs"):
                font_kwargs["complex_font"] = rule["font_cs"]
            if rule.get("size_pt") is not None:
                font_kwargs["size_pt"] = rule["size_pt"]
            if rule.get("bold") is not None:
                font_kwargs["bold"] = rule["bold"]
            if rule.get("language"):
                font_kwargs["language"] = rule["language"]
            if rule.get("rtl") is not None:
                font_kwargs["rtl"] = rule["rtl"]
            for p in container.paragraphs:
                if font_kwargs:
                    for run in p.runs:
                        set_run_fonts(run, **font_kwargs)
                if rule.get("alignment") in _HF_ALIGN:
                    p.alignment = _HF_ALIGN[rule["alignment"]]
            if font_kwargs:
                changed.append(f"{rule_key}_font")
    return changed


def _apply_table_rule(doc, table_rule, table_indices=None):
    """应用表格字体和几何；不改变行列数、合并关系或单元格内容。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    if not isinstance(table_rule, dict) or not table_rule:
        return []
    overrides = table_rule.get("overrides") if table_indices is None else None
    if isinstance(overrides, list) and overrides:
        base_rule = {
            key: value for key, value in table_rule.items()
            if key not in {"overrides", "landscape_table_indices"}
        }
        all_changes = set(_apply_table_rule(doc, base_rule)) if base_rule else set()
        for override in overrides:
            index = override.get("table_index")
            if not isinstance(index, int) or not 0 <= index < len(doc.tables):
                all_changes.add(f"table_override_{index}_not_applied")
                continue
            merged = dict(base_rule)
            merged.update({
                key: value for key, value in override.items()
                if key != "table_index"
            })
            all_changes.update(_apply_table_rule(doc, merged, [index]))
            all_changes.add(f"table_{index}_override")
        return sorted(all_changes)
    changed_fields = set()
    selected_tables = (
        [doc.tables[index] for index in table_indices]
        if table_indices is not None else doc.tables
    )
    for table in selected_tables:
        tbl_pr = table._tbl.tblPr
        if table_rule.get("alignment") in _TABLE_ALIGN:
            table.alignment = _TABLE_ALIGN[table_rule["alignment"]]
            changed_fields.add("table_alignment")
        layout_value = table_rule.get("layout")
        if layout_value is None and table_rule.get("autofit") is not None:
            layout_value = "autofit" if table_rule["autofit"] else "fixed"
        if layout_value:
            table.autofit = layout_value == "autofit"
            layout = tbl_pr.find(qn("w:tblLayout"))
            if layout is None:
                layout = OxmlElement("w:tblLayout")
                tbl_pr.append(layout)
            layout.set(qn("w:type"), layout_value)
            changed_fields.add("table_layout")
        width_pct = table_rule.get("width_pct")
        preferred_width_mm = table_rule.get("preferred_width_mm")
        if width_pct is not None or preferred_width_mm is not None:
            width = tbl_pr.find(qn("w:tblW"))
            if width is None:
                width = OxmlElement("w:tblW")
                tbl_pr.append(width)
            if width_pct is not None:
                width.set(qn("w:type"), "pct")
                width.set(qn("w:w"), str(round(width_pct * 50)))
            else:
                width.set(qn("w:type"), "dxa")
                width.set(
                    qn("w:w"), str(round(preferred_width_mm / 25.4 * 1440)))
            changed_fields.add("table_width")
        column_widths = table_rule.get("column_widths_pct")
        if column_widths:
            grid_columns = table._tbl.tblGrid.findall(qn("w:gridCol"))
            if len(grid_columns) == len(column_widths):
                total_twips = (
                    round(preferred_width_mm / 25.4 * 1440)
                    if preferred_width_mm is not None else 9000
                )
                denominator = sum(column_widths)
                for grid_column, percentage in zip(grid_columns, column_widths):
                    grid_column.set(
                        qn("w:w"),
                        str(round(total_twips * percentage / denominator)),
                    )
                changed_fields.add("table_column_widths")
        cell_margins = table_rule.get("cell_margins_mm") or {}
        if cell_margins:
            margins = tbl_pr.find(qn("w:tblCellMar"))
            if margins is None:
                margins = OxmlElement("w:tblCellMar")
                tbl_pr.append(margins)
            for edge in ("top", "left", "bottom", "right"):
                if cell_margins.get(edge) is None:
                    continue
                element = margins.find(qn(f"w:{edge}"))
                if element is None:
                    element = OxmlElement(f"w:{edge}")
                    margins.append(element)
                element.set(qn("w:w"), str(round(cell_margins[edge] / 25.4 * 1440)))
                element.set(qn("w:type"), "dxa")
            changed_fields.add("table_cell_margins")
        if table_rule.get("borders") is not None:
            borders = tbl_pr.find(qn("w:tblBorders"))
            if borders is None:
                borders = OxmlElement("w:tblBorders")
                tbl_pr.append(borders)
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                el = borders.find(qn(f"w:{edge}"))
                if el is None:
                    el = OxmlElement(f"w:{edge}")
                    borders.append(el)
                el.set(qn("w:val"), "single" if table_rule["borders"] else "nil")
                el.set(qn("w:sz"), "4")
            changed_fields.add("table_borders")
        for r_i, row in enumerate(table.rows):
            is_header = r_i == 0
            tr_pr = row._tr.get_or_add_trPr()
            if is_header and table_rule.get("repeat_header_row") is not None:
                header = tr_pr.find(qn("w:tblHeader"))
                if table_rule["repeat_header_row"]:
                    if header is None:
                        header = OxmlElement("w:tblHeader")
                        tr_pr.append(header)
                    header.set(qn("w:val"), "true")
                elif header is not None:
                    tr_pr.remove(header)
                changed_fields.add("table_repeat_header")
            if table_rule.get("allow_row_break") is not None:
                cant_split = tr_pr.find(qn("w:cantSplit"))
                if table_rule["allow_row_break"]:
                    if cant_split is not None:
                        tr_pr.remove(cant_split)
                elif cant_split is None:
                    tr_pr.append(OxmlElement("w:cantSplit"))
                changed_fields.add("table_row_break")
            for cell in row.cells:
                if table_rule.get("vertical_alignment") in _CELL_VERTICAL_ALIGN:
                    cell.vertical_alignment = _CELL_VERTICAL_ALIGN[
                        table_rule["vertical_alignment"]]
                    changed_fields.add("table_vertical_alignment")
                for p in cell.paragraphs:
                    font_kwargs = {}
                    if table_rule.get("font_eastasia"):
                        font_kwargs["eastasia"] = table_rule["font_eastasia"]
                    if table_rule.get("font_ascii"):
                        font_kwargs["ascii_font"] = table_rule["font_ascii"]
                    if table_rule.get("font_cs"):
                        font_kwargs["complex_font"] = table_rule["font_cs"]
                    if table_rule.get("size_pt") is not None:
                        font_kwargs["size_pt"] = table_rule["size_pt"]
                    if table_rule.get("language"):
                        font_kwargs["language"] = table_rule["language"]
                    if table_rule.get("rtl") is not None:
                        font_kwargs["rtl"] = table_rule["rtl"]
                    if is_header and table_rule.get("header_bold"):
                        font_kwargs["bold"] = True
                    if font_kwargs:
                        for run in p.runs:
                            set_run_fonts(run, **font_kwargs)
                        changed_fields.add("table_font")
                    align_key = "header_alignment" if is_header else "body_alignment"
                    if table_rule.get(align_key) in _HF_ALIGN:
                        p.alignment = _HF_ALIGN[table_rule[align_key]]
                        changed_fields.add("table_text_alignment")
    return sorted(changed_fields)


def _set_sect_landscape(sect_pr):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    section_type = sect_pr.find(qn("w:type"))
    if section_type is None:
        section_type = OxmlElement("w:type")
        sect_pr.insert(0, section_type)
    section_type.set(qn("w:val"), "nextPage")
    page_size = sect_pr.find(qn("w:pgSz"))
    if page_size is None:
        page_size = OxmlElement("w:pgSz")
        page_margin = sect_pr.find(qn("w:pgMar"))
        if page_margin is not None:
            sect_pr.insert(list(sect_pr).index(page_margin), page_size)
        else:
            sect_pr.append(page_size)
    try:
        width = int(page_size.get(qn("w:w")))
        height = int(page_size.get(qn("w:h")))
    except (TypeError, ValueError):
        width, height = 11906, 16838  # A4 twips
    page_size.set(qn("w:w"), str(max(width, height)))
    page_size.set(qn("w:h"), str(min(width, height)))
    page_size.set(qn("w:orient"), "landscape")
    columns = sect_pr.find(qn("w:cols"))
    if columns is not None:
        columns.set(qn("w:num"), "1")


def _section_break_paragraph(sect_pr):
    from docx.oxml import OxmlElement

    paragraph = OxmlElement("w:p")
    properties = OxmlElement("w:pPr")
    properties.append(sect_pr)
    paragraph.append(properties)
    return paragraph


def _wrap_tables_in_landscape_sections(doc, table_indices):
    """把点名的顶层表格包进横向节；不移动表格、不改变内容与合并关系。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    if not table_indices:
        return []
    tables = list(doc.tables)
    bad = [index for index in table_indices if index >= len(tables)]
    if bad:
        raise ValueError(
            f"table.landscape_table_indices 超出表格数量 {len(tables)}：{bad}")
    body = doc.element.body
    base_final = deepcopy(body.sectPr)
    changed = []
    for index in sorted(table_indices):
        table_element = tables[index]._tbl
        has_content_before = table_element.getprevious() is not None
        next_element = table_element.getnext()
        has_content_after = (
            next_element is not None and next_element.tag != body.sectPr.tag)
        if has_content_before:
            portrait_break = deepcopy(base_final)
            section_type = portrait_break.find(qn("w:type"))
            if section_type is None:
                section_type = OxmlElement("w:type")
                portrait_break.insert(0, section_type)
            section_type.set(qn("w:val"), "nextPage")
            table_element.addprevious(_section_break_paragraph(portrait_break))
        if has_content_after:
            landscape_break = deepcopy(base_final)
            _set_sect_landscape(landscape_break)
            table_element.addnext(_section_break_paragraph(landscape_break))
        else:
            _set_sect_landscape(body.sectPr)
        changed.append(f"table_{index}_landscape_section")
    return changed



def _apply_columns(doc, num_cols, before_element=None):
    """多栏排版：before_element（w:p 元素）之后的内容进入多栏（连续分节符切分）。
    before_element 为 None 时整篇多栏。"""
    import copy
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    if not isinstance(num_cols, int) or num_cols < 2:
        return False
    final_sect = doc.sections[0]._sectPr
    if before_element is not None:
        # 在切分点前的段落上插入单栏的段落级 sectPr（该段成为第一节的结尾）
        first_sect = copy.deepcopy(final_sect)
        cols = first_sect.find(qn("w:cols"))
        if cols is None:
            cols = OxmlElement("w:cols")
            doc_grid = first_sect.find(qn("w:docGrid"))
            if doc_grid is not None:
                first_sect.insert(list(first_sect).index(doc_grid), cols)
            else:
                first_sect.append(cols)
        cols.set(qn("w:num"), "1")
        ppr = before_element.find(qn("w:pPr"))
        if ppr is None:
            ppr = OxmlElement("w:pPr")
            before_element.insert(0, ppr)
        ppr.append(first_sect)
    cols = final_sect.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        doc_grid = final_sect.find(qn("w:docGrid"))
        if doc_grid is not None:
            final_sect.insert(list(final_sect).index(doc_grid), cols)
        else:
            final_sect.append(cols)
    cols.set(qn("w:num"), str(num_cols))
    return True


def _insert_toc(doc, before_paragraph, levels=(1, 2)):
    """在 before_paragraph 前插入目录标题 + TOC 域。Word 打开后更新域即生成目录。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    def _new_para():
        p = OxmlElement("w:p")
        before_paragraph._p.addprevious(p)
        return p

    # 目录标题段
    title_p = _new_para()
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "目录"
    r.append(t)
    title_p.append(r)

    # TOC 域
    lvl = "-".join(str(x) for x in sorted(levels))
    field_p = _new_para()

    def _fld(kind):
        el = OxmlElement("w:fldChar")
        el.set(qn("w:fldCharType"), kind)
        return el

    r1 = OxmlElement("w:r"); r1.append(_fld("begin"))
    r2 = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-%d" \\h \\z \\u' % max(levels)
    r2.append(instr)
    r3 = OxmlElement("w:r"); r3.append(_fld("separate"))
    r4 = OxmlElement("w:r")
    t4 = OxmlElement("w:t")
    t4.text = "（在 Word 中右键此处选择「更新域」生成目录）"
    r4.append(t4)
    r5 = OxmlElement("w:r"); r5.append(_fld("end"))
    for r in (r1, r2, r3, r4, r5):
        field_p.append(r)
    return True


def apply_format(docx_path, spec, rolemap, out_path, track=False,
                 template_path=None, allow_risky_structure=False):
    """应用 FormatSpec × RoleMap，输出 docx，返回 changelog list[dict]。
    rolemap: {idx: role}（idx 对应 extract.py 的段落序号）。
    模板未明确指定的角色统一与正文保持一致（套用 body 规则与样式）。

    页面级：页边距/行网格 + 页眉页脚（page.header/footer，footer 支持页码域）；
    表格：spec.table 规则（首行表头加粗居中、单元格字体字号、边框）。
    track=True 时输出修订模式文档：段落/字符格式改动以 w:pPrChange /
    w:rPrChange 记录，Word 审阅视图可见。
    """
    doc = Document(docx_path)
    from core.extract import iter_main_paragraphs
    roles = spec.get("roles", {})
    cleanup_mode = (spec.get("cleanup") or {}).get("mode", "controlled")
    # 必须在创建/更新 FormatAgent 样式之前解析，避免把新样式误认成目标原样式。
    target_body_style = resolve_target_body_style(doc, rolemap)
    # 同名目标样式可能本身承载自动编号。ensure_role_styles 会重置这些样式，
    # 因此必须先逐段保存有效编号；模板未控制编号时再转成段落直接覆盖。
    preserved_numbering = {}
    for idx, (paragraph, table_depth) in enumerate(iter_main_paragraphs(doc)):
        if table_depth:
            continue
        role = rolemap.get(idx, rolemap.get(str(idx)))
        if role is None:
            continue
        resolved_role = _resolved_role(role, roles)
        rule = (
            roles.get(resolved_role, {})
            if resolved_role is not None else roles.get("body", {})
        )
        if not isinstance(rule.get("numbering"), dict):
            num_pr = effective_automatic_numbering_override(paragraph)
            if num_pr is not None:
                preserved_numbering[idx] = num_pr
    role_styles = ensure_role_styles(
        doc, spec, target_body_style=target_body_style)

    # ---- 页面级 ----
    page = spec.get("page") or {}
    extra_changes = []
    margin = page.get("margin") or {}
    # 多节文档默认保留各节纸张与横竖方向；显式 override 只影响点名的节。
    if len(doc.sections) == 1 or allow_risky_structure:
        for section in doc.sections:
            extra_changes.extend(_apply_section_geometry(section, page))
            _apply_margin(section, margin)
            extra_changes.extend(_apply_header_footer_distances(section, page))
    for override in page.get("section_overrides") or []:
        index = override["section_index"]
        if index >= len(doc.sections):
            raise ValueError(
                f"page.section_overrides[{index}] 超出目标文档节数 {len(doc.sections)}")
        section = doc.sections[index]
        extra_changes.extend(_apply_section_geometry(section, override))
        _apply_margin(section, override.get("margin") or {})
        extra_changes.extend(_apply_header_footer_distances(section, override))
        if _set_section_columns(section, override.get("columns")):
            extra_changes.append(f"section_{index}_columns")
        if _set_section_page_numbering(
                section, override.get("page_numbering")):
            extra_changes.append(f"section_{index}_page_numbering")
        extra_changes.append(f"section_{index}_override")
    line_grid = page.get("line_grid") or {}
    if line_grid.get("line_pt") is not None:
        set_doc_grid(doc, line_pt=line_grid["line_pt"])
    # ---- 页眉页脚 + 表格 ----
    extra_changes.extend(_apply_header_footer(doc, page))
    extra_changes.extend(_apply_table_rule(doc, spec.get("table")))
    # ---- 段落级 ----
    changelog = []
    rev_id = 1
    for idx, (p, table_depth) in enumerate(iter_main_paragraphs(doc)):
        if table_depth:
            continue
        role = rolemap.get(idx, rolemap.get(str(idx)))
        if role is None:
            continue  # 未被标注的段落不动
        snapshot = snapshot_paragraph(p) if track else None
        resolved_role = _resolved_role(role, roles)
        if resolved_role is not None:
            rule = roles[resolved_role]
            style = role_styles[resolved_role]
            changed = apply_named_style(
                p, style, rule, role=role, cleanup_mode=cleanup_mode,
                preserved_numbering=preserved_numbering.get(idx))
            fallback_to_target_body = resolved_role != role
        else:
            # 模板未规定的角色与正文保持一致：套用 body 的规则和命名样式，
            # 同时清掉会遮蔽样式的直接格式（如原文自带的加粗/异体字）。
            # 真实自动编号保留（body 不在清编号集合内），仅清“取消编号”残留。
            body_rule = roles.get("body", {})
            body_style = role_styles.get("body", target_body_style)
            changed = apply_named_style(
                p, body_style, body_rule, role="body", cleanup_mode=cleanup_mode,
                preserved_numbering=preserved_numbering.get(idx))
            style = body_style
            fallback_to_target_body = True
        if track:
            # 修订稿额外把新样式摊平成直接格式（WPS 兼容；干净稿不做）。
            flatten_style_formatting(p, getattr(style, "_element", None))
            rev_id = mark_paragraph_revision(p, snapshot, rev_id_start=rev_id)
        changelog.append({
            "idx": idx,
            "role": role,
            "style_name": style.name,
            "text": p.text.strip()[:30],
            "changed_fields": changed,
            "fallback_to_target_body": fallback_to_target_body,
        })

    # ---- 脚注、题注序号、图表目录与现有域 ----
    from core.academic_fields import apply_academic_features
    academic_result = apply_academic_features(
        doc, spec, rolemap, role_styles=role_styles)
    if academic_result["changed_fields"] or academic_result["diagnostics"]:
        changelog.append({
            "idx": -3,
            "role": "academic_fields",
            "style_name": "-",
            "text": "脚注/题注/图表目录/引用域",
            "changed_fields": academic_result["changed_fields"],
            "fallback_to_target_body": False,
            "captions": academic_result["captions"],
            "diagnostics": academic_result["diagnostics"],
        })

    # ---- 技术手册图文邻接 + 法律 TA/TOA 域 ----
    from core.technical_features import apply_technical_features
    technical_result = apply_technical_features(doc, spec, rolemap)
    if technical_result["changed_fields"] or technical_result["diagnostics"]:
        changelog.append({
            "idx": -4, "role": "technical_features", "style_name": "-",
            "text": "技术手册图文邻接绑定",
            "changed_fields": technical_result["changed_fields"],
            "fallback_to_target_body": False,
            "diagnostics": technical_result["diagnostics"],
        })

    from core.legal_features import apply_legal_features
    legal_result = apply_legal_features(
        doc, spec, rolemap, role_styles=role_styles)
    if legal_result["changed_fields"] or legal_result["diagnostics"]:
        changelog.append({
            "idx": -5, "role": "legal_fields", "style_name": "-",
            "text": "法律 TA/TOA 域",
            "changed_fields": legal_result["changed_fields"],
            "fallback_to_target_body": False,
            "diagnostics": legal_result["diagnostics"],
            "allowed_additions": legal_result["allowed_additions"],
        })

    # ---- 点名表格的横向节 ----
    # 使用表格前后的节属性包裹现有 w:tbl；不复制、搬动或重建表格。
    landscape_indices = (
        (spec.get("table") or {}).get("landscape_table_indices") or [])
    extra_changes.extend(
        _wrap_tables_in_landscape_sections(doc, landscape_indices))

    # ---- 目录 + 多栏（结构级，在段落样式完成后执行）----
    toc = spec.get("toc") or {}
    columns = page.get("columns")
    structure_enabled = bool((spec.get("structure") or {}).get("enabled"))
    # 论文结构模式由 thesis_structure 生成独立目录节；不要再插一次旧版占位目录。
    want_toc = (
        isinstance(toc, dict) and toc.get("enabled") and not structure_enabled)
    want_cols = isinstance(columns, int) and columns >= 2
    if want_cols and len(doc.sections) > 1 and not allow_risky_structure:
        raise ValueError(
            "多节源文档不能直接启用 page.columns；"
            "请先通过能力预检并显式允许结构重构")
    if want_toc or want_cols:
        first_heading = None
        for idx, (p, table_depth) in enumerate(iter_main_paragraphs(doc)):
            if table_depth:
                continue
            r = rolemap.get(idx, rolemap.get(str(idx)))
            if r in (
                "heading_1", "heading_2", "heading_3", "chapter_heading",
                "bibliography_heading", "appendix_heading",
            ):
                first_heading = p
                break
        if want_toc and first_heading is not None:
            # 目录要能收录我们的命名样式标题：强制补大纲级别
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn as _qn
            for role, lvl in (("heading_1", "0"), ("heading_2", "1"), ("heading_3", "2")):
                style = role_styles.get(role)
                if style is None:
                    continue
                ppr = style.element.get_or_add_pPr()
                if ppr.find(_qn("w:outlineLvl")) is None:
                    ol = OxmlElement("w:outlineLvl")
                    ol.set(_qn("w:val"), lvl)
                    ppr.append(ol)
            _insert_toc(doc, first_heading, levels=toc.get("levels") or (1, 2))
            extra_changes.append("toc_inserted")
        if want_cols:
            # 多栏从第一个标题开始（标题/副标题/摘要保持单栏，符合论文惯例）
            before_el = first_heading._p.getprevious() if first_heading is not None else None
            if _apply_columns(doc, columns, before_element=before_el):
                extra_changes.append(f"columns_{columns}")

    if extra_changes:
        changelog.append({
            "idx": -1,
            "role": "page",
            "style_name": "-",
            "text": "页眉页脚/表格/目录/多栏（页面级）",
            "changed_fields": extra_changes,
            "fallback_to_target_body": False,
        })

    # ---- 论文结构级：安全重建封面/前置页/目录/正文/参考文献 ----
    # 放在段落样式之后执行，确保 STYLEREF 与 TOC 能引用稳定的命名样式。
    if structure_enabled:
        from core.thesis_structure import assemble_thesis_structure
        structure_result = assemble_thesis_structure(
            doc, spec, rolemap, template_path=template_path,
            allow_risky_structure=allow_risky_structure)
        if structure_result:
            changelog.append({
                "idx": -2,
                "role": "structure",
                "style_name": "-",
                "text": "论文封面/前置页/目录/分节/动态页眉页码",
                "changed_fields": structure_result["changed_fields"],
                "fallback_to_target_body": False,
                "allowed_additions": structure_result["allowed_additions"],
                "stripped_prefixes": structure_result["stripped_prefixes"],
                "section_kinds": structure_result["section_kinds"],
            })

    doc.save(out_path)
    return changelog


def write_report(changelog, spec, report_path):
    """把 changelog 写成 markdown 修改对照报告。"""
    lines = ["# 排版修改对照报告", ""]
    if spec.get("style_pack"):
        lines.extend([f"- Style Pack：`{spec['style_pack']}`"])
        for note in spec.get("style_pack_notes") or []:
            lines.append(f"- 注意：{note}")
        lines.append("")
    page = spec.get("page") or {}
    if page:
        lines.append("## 页面设置")
        if page.get("size") or page.get("orientation"):
            lines.append(
                f"- 纸张/方向：{page.get('size', '保留')} / "
                f"{page.get('orientation', '保留')}")
        margin = page.get("margin") or {}
        if margin:
            lines.append(
                f"- 页边距（mm）：上 {margin.get('top_mm', '-')} / 下 {margin.get('bottom_mm', '-')}"
                f" / 左 {margin.get('left_mm', '-')} / 右 {margin.get('right_mm', '-')}")
        lg = page.get("line_grid") or {}
        if lg.get("line_pt"):
            lines.append(f"- 行网格：{lg['line_pt']} 磅/行")
        lines.append("")
    lines.append("## 段落修改明细")
    lines.append("")
    lines.append("| 段落 | 角色 | Word 样式 | 改动字段 | 内容摘要 |")
    lines.append("|---|---|---|---|---|")
    for c in changelog:
        fields = ", ".join(c["changed_fields"]) if c["changed_fields"] else "（无字段改动）"
        lines.append(
            f"| {c['idx']} | {c['role']} | {c.get('style_name', '-')} | {fields} | {c['text']} |")
    lines.append("")
    with open(report_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    return report_path
