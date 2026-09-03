"""法律 brief 的 TA/TOA 域支持；不判断引证内容是否准确。"""

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from core.extract import iter_main_paragraphs


def _role(rolemap, index):
    return rolemap.get(index, rolemap.get(str(index)))


def _field_instructions(element):
    instructions = []
    stack = []
    for node in element.iter():
        if node.tag == qn("w:fldChar"):
            kind = node.get(qn("w:fldCharType"))
            if kind == "begin":
                stack.append([])
            elif kind == "end" and stack:
                instruction = "".join(stack.pop()).strip()
                if instruction:
                    instructions.append(instruction)
        elif node.tag == qn("w:instrText"):
            value = node.text or ""
            if stack:
                stack[-1].append(value)
            elif value.strip():
                instructions.append(value.strip())
        elif node.tag == qn("w:fldSimple"):
            value = (node.get(qn("w:instr")) or "").strip()
            if value:
                instructions.append(value)
    instructions.extend(
        instruction for chunks in stack
        if (instruction := "".join(chunks).strip())
    )
    return instructions


def _add_run_property(run, name):
    rpr = run._element.get_or_add_rPr()
    element = OxmlElement(name)
    element.set(qn("w:val"), "1")
    rpr.append(element)


def _append_field(paragraph, instruction, hidden=False):
    for kind in ("begin", None, "separate", "end"):
        run = paragraph.add_run()
        if hidden:
            _add_run_property(run, "w:vanish")
        if kind is None:
            element = OxmlElement("w:instrText")
            element.set(qn("xml:space"), "preserve")
            element.text = f" {instruction.strip()} "
        else:
            element = OxmlElement("w:fldChar")
            element.set(qn("w:fldCharType"), kind)
        run._element.append(element)


def _set_update_fields_on_open(document):
    settings = document.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def _quote(value):
    return str(value).replace("\\", "\\\\").replace('"', '\\"')


def _ta_instruction(mark):
    instruction = f'TA \\l "{_quote(mark["long"])}"'
    if mark.get("short"):
        instruction += f' \\s "{_quote(mark["short"])}"'
    instruction += f' \\c {int(mark.get("category", 1))}'
    if mark.get("bold"):
        instruction += " \\b"
    return instruction


def _paragraph_after(paragraph):
    element = OxmlElement("w:p")
    paragraph._p.addnext(element)
    return Paragraph(element, paragraph._parent)


def apply_legal_features(document, spec, rolemap, role_styles=None):
    """保留既有域，并按显式 citation_marks/insert_toa 配置新增域。"""
    config = spec.get("legal") or {}
    enabled = spec.get("profile") == "english_legal_brief" or bool(config)
    if not enabled:
        return {"changed_fields": [], "diagnostics": [], "allowed_additions": []}

    changed = []
    diagnostics = []
    allowed_additions = []
    existing = _field_instructions(document.element)
    existing_ta = sum(1 for value in existing if value.upper().startswith("TA "))
    existing_toa = sum(1 for value in existing if value.upper().startswith("TOA"))
    if existing_ta or existing_toa:
        diagnostics.append({
            "code": "EXISTING_LEGAL_FIELDS_PRESERVED",
            "ta": existing_ta, "toa": existing_toa,
        })

    paragraphs = [
        (index, paragraph)
        for index, (paragraph, table_depth) in enumerate(iter_main_paragraphs(document))
        if table_depth == 0
    ]
    marked = 0
    missing = []
    mark_all = config.get("mark_all", True)
    for mark in config.get("citation_marks") or []:
        needle = mark["text"]
        instruction = _ta_instruction(mark)
        matches = 0
        for _, paragraph in paragraphs:
            occurrences = paragraph.text.count(needle)
            if not occurrences:
                continue
            prior = _field_instructions(paragraph._p)
            if any(value.strip() == instruction for value in prior):
                continue
            count = occurrences if mark_all else 1
            for _ in range(count):
                _append_field(paragraph, instruction, hidden=True)
                marked += 1
                matches += 1
            if not mark_all:
                break
        if not matches:
            missing.append(needle)
    if marked:
        changed.append(f"ta_marks_{marked}")
        _set_update_fields_on_open(document)
    if missing:
        diagnostics.append({
            "code": "CITATION_MARK_TEXT_NOT_FOUND",
            "texts": missing,
            "message": "未找到这些显式引证标记文本，未猜测替代位置。",
        })

    if config.get("insert_toa"):
        if existing_toa:
            diagnostics.append({"code": "TOA_ALREADY_PRESENT"})
        else:
            heading = next((
                paragraph for index, paragraph in paragraphs
                if _role(rolemap, index) == "table_of_authorities_heading"
            ), None)
            if heading is None and config.get("create_heading"):
                heading = document.add_paragraph("TABLE OF AUTHORITIES")
                style = (role_styles or {}).get("table_of_authorities_heading")
                if style is not None:
                    heading.style = style
                allowed_additions.append("TABLE OF AUTHORITIES")
                changed.append("toa_heading_created")
            if heading is None:
                diagnostics.append({
                    "code": "TOA_HEADING_REQUIRED",
                    "message": "未找到 Table of Authorities 标题；为避免重排正文，没有自动创建。",
                })
            else:
                field_paragraph = _paragraph_after(heading)
                body_style = (role_styles or {}).get("authority_entry") or (
                    (role_styles or {}).get("body"))
                if body_style is not None:
                    field_paragraph.style = body_style
                instruction = config.get("toa_instruction") or 'TOA \\h \\c "1"'
                _append_field(field_paragraph, instruction, hidden=False)
                _set_update_fields_on_open(document)
                changed.append("toa_field_inserted")

    return {
        "changed_fields": changed,
        "diagnostics": diagnostics,
        "allowed_additions": allowed_additions,
    }
