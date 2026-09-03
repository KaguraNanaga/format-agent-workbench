"""论文第三阶段：封面、前置页、目录、多节、动态页眉和页码。

这里重建“结构”，不盲拷贝模板正文或示例作者信息。源文档的标题、摘要、
关键词和正文只移动/拆分标签，不改写语义内容；校徽等视觉资产可从模板提取。
"""

from copy import deepcopy
from io import BytesIO
import re
import zipfile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt


_FRONT_HEADING = "摘　要"
_TOC_HEADING = "目　录"

_COVER_LABELS = {
    "姓名": "姓名", "学生姓名": "姓名", "作者": "姓名",
    "学号": "学号", "学生学号": "学号",
    "导师": "导师", "指导教师": "导师", "指导老师": "导师",
    "院系": "院系", "学院": "院系", "所在学院": "院系",
    "专业": "学科/专业", "学科": "学科/专业", "学科专业": "学科/专业",
    "申请学位": "申请学位", "学位": "申请学位",
}


def _canonical_cover_label(value):
    compact = re.sub(r"[\s：:]", "", value or "")
    return _COVER_LABELS.get(compact)


def extract_cover_metadata(docx_path):
    """从目标稿的显式标签或两列表格保守提取封面元数据。"""
    document = Document(docx_path)
    result = {}
    for paragraph in document.paragraphs:
        match = re.match(r"^\s*([^：:]{1,12})[：:]\s*(.+?)\s*$", paragraph.text)
        if not match:
            continue
        label = _canonical_cover_label(match.group(1))
        value = match.group(2).strip()
        if label and value and "待填写" not in value:
            result.setdefault(label, value)
    for table in document.tables:
        for row in table.rows:
            if len(row.cells) < 2:
                continue
            label = _canonical_cover_label(row.cells[0].text)
            value = row.cells[1].text.strip()
            if label and value and "待填写" not in value:
                result.setdefault(label, value)
    return result


def _role_for_index(rolemap, index):
    return rolemap.get(index, rolemap.get(str(index))) if isinstance(rolemap, dict) else None


def _set_run_font(run, eastasia="宋体", ascii_font="Times New Roman", size_pt=None,
                  bold=None, italic=None):
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:eastAsia"), eastasia)
    fonts.set(qn("w:ascii"), ascii_font)
    fonts.set(qn("w:hAnsi"), ascii_font)
    if size_pt is not None:
        run.font.size = Pt(float(size_pt))
    if bold is not None:
        run.bold = bool(bold)
    if italic is not None:
        run.italic = bool(italic)


def _format_paragraph(paragraph, *, size_pt, bold=False, alignment=None,
                      before=0, after=0, line_spacing=None):
    if alignment is not None:
        paragraph.alignment = alignment
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line_spacing is not None:
        pf.line_spacing = line_spacing
    for run in paragraph.runs:
        _set_run_font(run, size_pt=size_pt, bold=bold)


def _add_field(paragraph, instruction, cached_text=""):
    """加入标准复杂域 begin/instr/separate/result/end。"""
    begin_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run._element.append(begin)

    code_run = paragraph.add_run()
    code = OxmlElement("w:instrText")
    code.set(qn("xml:space"), "preserve")
    code.text = instruction
    code_run._element.append(code)

    separate_run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run._element.append(separate)

    if cached_text:
        result_run = paragraph.add_run(cached_text)
        _set_run_font(result_run, size_pt=9)

    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._element.append(end)


def _clear_story(story):
    root = story._element
    for child in list(root):
        root.remove(child)
    paragraph = OxmlElement("w:p")
    root.append(paragraph)
    from docx.text.paragraph import Paragraph
    return Paragraph(paragraph, story)


def _set_header(section, left_text, right_text=None, right_style_name=None,
                cached_right=""):
    section.header.is_linked_to_previous = False
    paragraph = _clear_story(section.header)
    # python-docx 的 Length 相减会返回 EMU 整数；1 twip = 635 EMU。
    usable_twips = int(round(
        (section.page_width - section.left_margin - section.right_margin) / 635))
    ppr = paragraph._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(usable_twips))
    tabs.append(tab)
    ppr.append(tabs)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1

    left = paragraph.add_run(left_text or "")
    _set_run_font(left, size_pt=9)
    paragraph.add_run("\t")
    if right_style_name:
        _add_field(
            paragraph, f'STYLEREF "{right_style_name}"', cached_text=cached_right)
    elif right_text:
        right = paragraph.add_run(right_text)
        _set_run_font(right, size_pt=9)


def _set_footer(section):
    section.footer.is_linked_to_previous = False
    paragraph = _clear_story(section.footer)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    _add_field(paragraph, "PAGE", cached_text="1")


def _clear_header_footer(section):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    _clear_story(section.header)
    _clear_story(section.footer)


def _set_page_numbering(section, number_format=None, start=None):
    sect_pr = section._sectPr
    existing = sect_pr.find(qn("w:pgNumType"))
    if existing is not None:
        sect_pr.remove(existing)
    if number_format is None and start is None:
        return
    element = OxmlElement("w:pgNumType")
    if number_format is not None:
        element.set(qn("w:fmt"), str(number_format))
    if start is not None:
        element.set(qn("w:start"), str(int(start)))
    successors = (
        "w:cols", "w:formProt", "w:vAlign", "w:noEndnote", "w:titlePg",
        "w:textDirection", "w:bidi", "w:rtlGutter", "w:docGrid",
    )
    for tag in successors:
        successor = sect_pr.find(qn(tag))
        if successor is not None:
            sect_pr.insert(list(sect_pr).index(successor), element)
            break
    else:
        sect_pr.append(element)


def _set_update_fields_on_open(document):
    settings = document.settings.element
    element = settings.find(qn("w:updateFields"))
    if element is None:
        element = OxmlElement("w:updateFields")
        compat = settings.find(qn("w:compat"))
        if compat is not None:
            settings.insert(list(settings).index(compat), element)
        else:
            settings.append(element)
    element.set(qn("w:val"), "true")


def _bare_section_properties(base_sect_pr):
    sect_pr = deepcopy(base_sect_pr)
    for tag in (
        "w:headerReference", "w:footerReference", "w:pgNumType", "w:titlePg",
    ):
        for element in list(sect_pr.findall(qn(tag))):
            sect_pr.remove(element)
    type_el = sect_pr.find(qn("w:type"))
    if type_el is None:
        type_el = OxmlElement("w:type")
        sect_pr.insert(0, type_el)
    type_el.set(qn("w:val"), "nextPage")
    return sect_pr


def _attach_section_break(paragraph, base_sect_pr):
    ppr = paragraph._p.get_or_add_pPr()
    old = ppr.find(qn("w:sectPr"))
    if old is not None:
        ppr.remove(old)
    ppr.append(_bare_section_properties(base_sect_pr))


def _remove_existing_section_breaks(paragraphs):
    for paragraph in paragraphs:
        ppr = paragraph._p.pPr
        sect_pr = ppr.find(qn("w:sectPr")) if ppr is not None else None
        if sect_pr is not None:
            ppr.remove(sect_pr)


def _style_for_role(document, spec, role):
    from core.style_set import style_id_for_role
    style_id = style_id_for_role(role)
    return next((style for style in document.styles if style.style_id == style_id), None)


def _extract_logo(template_path):
    """按版面中的近方形、小尺寸图片选择校徽；无法判断时不猜。"""
    if not template_path:
        return None
    try:
        template = Document(template_path)
        candidates = []
        for order, shape in enumerate(template.inline_shapes):
            blips = shape._inline.xpath(".//a:blip")
            if not blips:
                continue
            relationship_id = blips[0].get(qn("r:embed"))
            part = template.part.related_parts.get(relationship_id)
            if part is None or not getattr(part, "blob", None):
                continue
            width_mm = shape.width.mm
            height_mm = shape.height.mm
            ratio = width_mm / height_mm if height_mm else 99
            square_penalty = abs(ratio - 1)
            size_penalty = 0 if 12 <= max(width_mm, height_mm) <= 65 else 2
            candidates.append((square_penalty + size_penalty + order * 0.02, part.blob))
        if candidates:
            score, blob = min(candidates, key=lambda item: item[0])
            return blob if score < 1.25 else None
        # 浮动图片没有 inline_shape 尺寸时，仅在包内存在唯一近方形图片时采用。
        package_candidates = []
        for part in template.part.package.parts:
            image = getattr(part, "image", None)
            if image is None or not getattr(part, "blob", None):
                continue
            width = getattr(image, "px_width", 0)
            height = getattr(image, "px_height", 0)
            if width and height and abs(width / height - 1) <= 0.2:
                package_candidates.append(part.blob)
        return package_candidates[0] if len(package_candidates) == 1 else None
    except (OSError, KeyError, ValueError, zipfile.BadZipFile):
        return None


def _add_cover(document, title_text, cover, template_path, base_sect_pr):
    elements = []
    additions = []

    if cover.get("logo"):
        logo = _extract_logo(template_path)
        if logo:
            document.add_picture(BytesIO(logo), width=Mm(30))
            logo_paragraph = document.paragraphs[-1]
            logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            logo_paragraph.paragraph_format.space_before = Pt(36)
            logo_paragraph.paragraph_format.space_after = Pt(8)
            elements.append(logo_paragraph._p)

    institution = str(cover.get("type_label") or cover.get("institution") or "学位论文")
    institution_p = document.add_paragraph(institution)
    _format_paragraph(
        institution_p, size_pt=18, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        before=8, after=72)
    elements.append(institution_p._p)
    additions.append(institution)

    title_p = document.add_paragraph(title_text)
    _format_paragraph(
        title_p, size_pt=22, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        before=0, after=84, line_spacing=1.2)
    title_p.paragraph_format.keep_together = True
    elements.append(title_p._p)

    metadata = cover.get("metadata") or {}
    if metadata:
        table = document.add_table(rows=len(metadata), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        for row, (label, value) in zip(table.rows, metadata.items()):
            row.cells[0].width = Mm(34)
            row.cells[1].width = Mm(80)
            row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            row.cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            label_p = row.cells[0].paragraphs[0]
            value_p = row.cells[1].paragraphs[0]
            label_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            value_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            label_run = label_p.add_run(f"{label}：")
            value_run = value_p.add_run(value)
            _set_run_font(label_run, size_pt=12, bold=True)
            _set_run_font(value_run, size_pt=12)
            for paragraph in (label_p, value_p):
                paragraph.paragraph_format.space_before = Pt(2)
                paragraph.paragraph_format.space_after = Pt(2)
            additions.extend((f"{label}：", value))
        elements.append(table._tbl)

    date_text = str(cover.get("date_text") or "")
    if date_text:
        date_p = document.add_paragraph(date_text)
        _format_paragraph(
            date_p, size_pt=12, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
            before=28, after=0)
        elements.append(date_p._p)
        additions.append(date_text)

    end_p = document.add_paragraph()
    end_p.paragraph_format.space_before = Pt(0)
    end_p.paragraph_format.space_after = Pt(0)
    _attach_section_break(end_p, base_sect_pr)
    elements.append(end_p._p)
    return elements, additions


def _move_before(elements, anchor_element):
    for element in elements:
        anchor_element.addprevious(element)


def _strip_text_prefix(paragraph, prefixes):
    prefix = next((item for item in prefixes if paragraph.text.startswith(item)), None)
    if not prefix:
        return None
    remaining = len(prefix)
    for run in paragraph.runs:
        if remaining <= 0:
            break
        text = run.text
        if remaining >= len(text):
            run.text = ""
            remaining -= len(text)
        else:
            run.text = text[remaining:]
            remaining = 0
    return prefix if remaining == 0 else None


def _add_front_heading(document, text, spec, anchor):
    style = _style_for_role(document, spec, "abstract_heading")
    paragraph = document.add_paragraph(text, style=style)
    paragraph._p.getparent().remove(paragraph._p)
    anchor._p.addprevious(paragraph._p)
    return paragraph


def _add_toc(document, spec, first_heading, base_sect_pr, levels):
    style = _style_for_role(document, spec, "abstract_heading")
    title = document.add_paragraph(_TOC_HEADING, style=style)
    field = document.add_paragraph()
    field.paragraph_format.space_before = Pt(6)
    field.paragraph_format.space_after = Pt(0)
    max_level = max(levels or (1, 2, 3))
    _add_field(field, f'TOC \\o "1-{max_level}" \\h \\z \\u')
    _attach_section_break(field, base_sect_pr)
    _move_before((title._p, field._p), first_heading._p)
    return title, field


def _section_break_before(document, paragraph, base_sect_pr):
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(0)
    _attach_section_break(spacer, base_sect_pr)
    spacer._p.getparent().remove(spacer._p)
    paragraph._p.addprevious(spacer._p)
    return spacer


def _apply_geometry(section, page, header_rule):
    margin = page.get("margin") or {}
    if margin.get("top_mm") is not None:
        section.top_margin = Mm(margin["top_mm"])
    if margin.get("bottom_mm") is not None:
        section.bottom_margin = Mm(margin["bottom_mm"])
    if margin.get("left_mm") is not None:
        section.left_margin = Mm(margin["left_mm"])
    if margin.get("right_mm") is not None:
        section.right_margin = Mm(margin["right_mm"])
    if header_rule.get("header_distance_mm") is not None:
        section.header_distance = Mm(header_rule["header_distance_mm"])
    if header_rule.get("footer_distance_mm") is not None:
        section.footer_distance = Mm(header_rule["footer_distance_mm"])
    section.start_type = WD_SECTION.NEW_PAGE
    section.different_first_page_header_footer = False


def _configure_sections(document, kinds, spec, first_chapter_text):
    structure = spec.get("structure") or {}
    numbering = structure.get("page_numbering") or {}
    header_rule = structure.get("running_header") or {}
    page = spec.get("page") or {}
    sections = list(document.sections)
    if len(sections) != len(kinds):
        raise ValueError(
            f"论文结构节数异常：预期 {len(kinds)}，实际 {len(sections)}")

    from core.style_set import style_name_for_role
    chapter_style_name = style_name_for_role(
        "chapter_heading", (spec.get("roles") or {}).get("chapter_heading"))
    left_text = str(header_rule.get("left_text") or "学位论文")
    front_started = False
    for section, kind in zip(sections, kinds):
        _apply_geometry(section, page, header_rule)
        if kind == "cover":
            _clear_header_footer(section)
            _set_page_numbering(section)
        elif kind in {"abstract", "toc"}:
            _set_header(
                section, left_text,
                right_text="摘要" if kind == "abstract" else "目录")
            _set_footer(section)
            _set_page_numbering(
                section,
                numbering.get("front_format", "upperRoman"),
                numbering.get("front_start", 1) if not front_started else None)
            front_started = True
        elif kind == "body":
            _set_header(
                section, left_text, right_style_name=chapter_style_name,
                cached_right=first_chapter_text)
            _set_footer(section)
            _set_page_numbering(
                section, numbering.get("body_format", "decimal"),
                numbering.get("body_start", 1))
        elif kind == "references":
            _set_header(section, left_text, right_text="参考文献")
            _set_footer(section)
            _set_page_numbering(
                section, numbering.get("body_format", "decimal"), None)


def assemble_thesis_structure(document, spec, rolemap, template_path=None,
                              allow_risky_structure=False):
    """在已经完成命名样式排版的文档上组装论文结构。

    返回用于 changelog/文本一致性校验的结构化结果。
    """
    structure = spec.get("structure") or {}
    if not structure.get("enabled") or structure.get("mode") != "thesis":
        return None
    if len(document.sections) > 1 and not allow_risky_structure:
        raise ValueError(
            "源文档已有多个分节；默认拒绝为论文结构删除或重排原分节")

    from core.extract import iter_main_paragraphs
    indexed_paragraphs = [
        (index, paragraph)
        for index, (paragraph, depth) in enumerate(iter_main_paragraphs(document))
        if depth == 0
    ]
    original_paragraphs = [paragraph for _, paragraph in indexed_paragraphs]
    role_paragraphs = {}
    for index, paragraph in indexed_paragraphs:
        role = _role_for_index(rolemap, index)
        if role:
            role_paragraphs.setdefault(role, []).append(paragraph)

    title = next(iter(role_paragraphs.get("title", [])), None)
    abstract = next(iter(role_paragraphs.get("abstract_body", [])), None)
    keywords = next(iter(role_paragraphs.get("keywords", [])), None)
    first_heading = next(iter(role_paragraphs.get("chapter_heading", [])), None)
    if first_heading is None:
        first_heading = next(iter(role_paragraphs.get("heading_1", [])), None)
    references = next(iter(role_paragraphs.get("bibliography_heading", [])), None)
    if title is None or first_heading is None:
        raise ValueError("论文结构组装至少需要 title 和 chapter_heading/heading_1 角色")

    _remove_existing_section_breaks(original_paragraphs)
    base_sect_pr = deepcopy(document.sections[-1]._sectPr)
    allowed_additions = []
    stripped_prefixes = []
    kinds = []

    cover = structure.get("cover") or {}
    if cover.get("enabled", True):
        cover_elements, additions = _add_cover(
            document, title.text, cover, template_path, base_sect_pr)
        anchor = abstract or keywords or first_heading
        title._p.getparent().remove(title._p)
        _move_before(cover_elements, anchor._p)
        allowed_additions.extend(additions)
        kinds.append("cover")

    front = structure.get("front_matter") or {}
    if front.get("abstract", True) and abstract is not None:
        abstract_heading = _add_front_heading(
            document, _FRONT_HEADING, spec, abstract)
        allowed_additions.append(_FRONT_HEADING)
        stripped = _strip_text_prefix(abstract, ("摘要：", "摘要:"))
        if stripped:
            stripped_prefixes.append(stripped)
        abstract_end = keywords or abstract
        _attach_section_break(abstract_end, base_sect_pr)
        kinds.append("abstract")

    if front.get("toc", True):
        levels = (spec.get("toc") or {}).get("levels") or [1, 2, 3]
        _add_toc(document, spec, first_heading, base_sect_pr, levels)
        allowed_additions.append(_TOC_HEADING)
        kinds.append("toc")

    if references is not None:
        _section_break_before(document, references, base_sect_pr)
    kinds.append("body")
    if references is not None:
        kinds.append("references")

    _configure_sections(document, kinds, spec, first_heading.text.strip())
    _set_update_fields_on_open(document)
    return {
        "changed_fields": [
            "cover_assembled", "front_matter_sections", "dynamic_toc",
            "running_headers", "page_numbering",
        ],
        "allowed_additions": allowed_additions,
        "stripped_prefixes": stripped_prefixes,
        "section_kinds": kinds,
    }
